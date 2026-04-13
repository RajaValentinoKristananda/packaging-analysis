import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import plotly.express as px
from plotly.subplots import make_subplots
import os
import base64
import io
import warnings
warnings.filterwarnings('ignore')

# ─────────────────────────────────────────────────────────────────────────────
# PAGE CONFIG
# ─────────────────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Packaging Analysis – PT Güntner",
    page_icon=(
        open('assets/güntner_logo.png','rb').read()
        if os.path.exists('assets/güntner_logo.png')
        else open('assets/guntner_logo.png','rb').read()
        if os.path.exists('assets/guntner_logo.png')
        else "🏭"
    ),
    layout="wide",
    initial_sidebar_state="collapsed"
)

# ─────────────────────────────────────────────────────────────────────────────
# CONSTANTS
# ─────────────────────────────────────────────────────────────────────────────
MATERIAL_MAP    = {'W': 'Wood', 'B': 'Carton Box', 'Cb': 'Cardboard', 'Ct': 'Carton'}
MATERIAL_COLORS = {
    'Wood':       '#7C4A19',
    'Carton Box': '#E8A838',
    'Cardboard':  '#C8975A',
    'Carton':     '#D4845A',
}
MATERIAL_LIGHT  = {
    'Wood':       '#F5E6D8',
    'Carton Box': '#FEF8E6',
    'Cardboard':  '#FAF0E0',
    'Carton':     '#FAEADE',
}
DIM_LABELS  = ['XS (<1 m)', 'S (1–2 m)', 'M (2–4 m)', 'L (4–7 m)', 'XL (>7 m)']
GREEN_DARK  = '#166534'
GREEN_MID   = '#16a34a'
GREEN_LIGHT = '#dcfce7'
MONTH_NAMES = ['Jan','Feb','Mar','Apr','May','Jun','Jul','Aug','Sep','Oct','Nov','Dec']

# ── External Reference Data (hardcoded) ────────────────────────────────────
# YearMonth labels for 38 months: Jan 2023 – Feb 2026
_REF_MONTHS = (
    [f'2023-{m:02d}' for m in range(1, 13)] +
    [f'2024-{m:02d}' for m in range(1, 13)] +
    [f'2025-{m:02d}' for m in range(1, 13)] +
    ['2026-01', '2026-02']
)

METAL_USAGE_KG = [
    # 2023
    326350.676, 300586.885, 382580.881, 295315.365, 358011.615, 388379.773,
    378098.950, 438412.060, 390148.525, 358666.315, 411435.432, 320479.067,
    # 2024
    403345.054, 356375.684, 412142.068, 240726.695, 365928.380, 356302.556,
    422622.045, 425380.226, 402399.338, 490139.663, 562716.516, 484130.344,
    # 2025
    471976.042, 368059.406, 383077.216, 392793.658, 391747.843, 343565.470,
    429873.544, 405548.102, 400530.451, 435622.190, 405318.235, 516203.896,
    # 2026
    415352.575, 418697.965,
]

PRODUCTION_UNIT = [
    # 2023
    3634, 3450, 3595, 1944, 2474, 2262, 2808, 2469, 3191, 2481, 3168, 3298,
    # 2024
    3545, 2544, 3060, 2640, 2513, 2963, 3765, 2992, 2517, 3224, 3258, 3110,
    # 2025
    3803, 3701, 2790, 1729, 2640, 2572, 3428, 2927, 3322, 3172, 3142, 3401,
    # 2026
    4481, 4379,
]

# Build reference DataFrame once
REF_DF = pd.DataFrame({
    'YearMonth':    _REF_MONTHS,
    'MetalUsageKG': METAL_USAGE_KG,
    'ProductionUnit': PRODUCTION_UNIT,
})
REF_DF['Label'] = REF_DF['YearMonth'].apply(
    lambda ym: pd.Period(ym, freq='M').to_timestamp().strftime('%b %Y')
)

# ─────────────────────────────────────────────────────────────────────────────
# CSS — NO SIDEBAR, CLEAN UPLOAD
# ─────────────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
    .main { background-color: #f8fafc; }
    section[data-testid="stSidebar"] { display: none; }
    h1,h2,h3 { color:#1e293b; font-weight:700; font-family:'Arial',sans-serif; }
    .stTabs [data-baseweb="tab-list"] {
        gap:.4rem; background:white; border-radius:.75rem;
        padding:.4rem; box-shadow:0 1px 3px rgba(0,0,0,.08);
    }
    .stTabs [data-baseweb="tab"] {
        background:transparent; border-radius:.5rem;
        padding:.55rem 1.1rem; font-weight:600; color:#64748b; border:none; font-size:13.5px;
    }
    .stTabs [data-baseweb="tab"][aria-selected="true"] {
        background:#16a34a; color:white;
    }
    #MainMenu,footer { visibility:hidden; }
    .tooltip-box {
        background:#f0fdf4; border-left:4px solid #16a34a;
        padding:12px 16px; border-radius:8px; margin:10px 0;
        font-size:13.5px; line-height:1.7; color:#1e293b;
    }
    .insight-box {
        background:#fffbeb; border-left:4px solid #d97706;
        padding:12px 16px; border-radius:8px; margin:10px 0;
        font-size:13.5px; line-height:1.7; color:#1e293b;
    }
    .critical-box {
        background:#fef2f2; border-left:4px solid #dc2626;
        padding:12px 16px; border-radius:8px; margin:10px 0;
        font-size:13.5px; line-height:1.7; color:#1e293b;
    }
    .axis-note {
        background:#f8fafc; border:1px solid #e2e8f0;
        padding:6px 12px; border-radius:6px; margin-top:-4px; margin-bottom:8px;
        font-size:12px; color:#64748b;
    }
    .filter-row {
        display:flex; align-items:center; gap:12px; margin-bottom:4px;
    }
    div[data-testid="stFileUploader"] { padding:0; }
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────────────────────────────────────────
# HEADER
# ─────────────────────────────────────────────────────────────────────────────
logo_html = ''
for candidate in ['assets/güntner_logo.png', 'assets/guntner_logo.png']:
    if os.path.exists(candidate):
        with open(candidate, 'rb') as f:
            b64 = base64.b64encode(f.read()).decode()
        logo_html = f'<img src="data:image/png;base64,{b64}" style="height:80px;">'
        break

if not logo_html:
    logo_html = (
        '<div style="width:80px;height:80px;display:flex;align-items:center;justify-content:center;">'
        '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 64 64" width="80" height="80">'
        '<defs>'
        '<radialGradient id="leafGrad" cx="40%" cy="35%" r="65%">'
        '<stop offset="0%" stop-color="#4ade80"/>'
        '<stop offset="100%" stop-color="#166534"/>'
        '</radialGradient>'
        '</defs>'
        '<!-- Main leaf -->'
        '<path d="M32 8 C18 8 8 20 8 36 C8 50 18 58 32 58 C46 58 56 50 56 36 C56 20 46 8 32 8 Z" '
        'fill="url(#leafGrad)" opacity="0.15"/>'
        '<path d="M32 56 C32 56 10 44 10 28 C10 16 20 8 32 8 C32 8 14 22 20 38 C24 48 32 56 32 56 Z" '
        'fill="#166534" opacity="0.9"/>'
        '<path d="M32 56 C32 56 54 44 54 28 C54 16 44 8 32 8 C32 8 50 22 44 38 C40 48 32 56 32 56 Z" '
        'fill="#16a34a" opacity="0.85"/>'
        '<!-- Center vein -->'
        '<path d="M32 10 Q30 32 32 56" stroke="#4ade80" stroke-width="1.2" fill="none" opacity="0.7"/>'
        '<!-- Side veins -->'
        '<path d="M32 22 Q24 28 18 30" stroke="#4ade80" stroke-width="0.8" fill="none" opacity="0.5"/>'
        '<path d="M32 30 Q23 36 17 40" stroke="#4ade80" stroke-width="0.8" fill="none" opacity="0.5"/>'
        '<path d="M32 38 Q25 43 21 48" stroke="#4ade80" stroke-width="0.7" fill="none" opacity="0.4"/>'
        '</svg>'
        '</div>'
    )

st.markdown(
    f'<div style="background:white;border-bottom:1px solid #e5e7eb;padding:12px 28px;'
    f'display:flex;align-items:center;gap:16px;margin:-4rem -4rem 1.5rem -4rem;">'
    f'{logo_html}'
    f'<div>'
    f'<div style="font-size:32px;font-weight:700;color:#111827;line-height:1.2;">Packaging Analysis Dashboard</div>'
    f'<div style="font-size:16px;color:#9ca3af;margin-top:2px;">Supporting Pledge 6 · Make all our packaging fully sustainable · PT Güntner Indonesia</div>'
    f'</div>'
    f'<div style="margin-left:auto;background:linear-gradient(90deg,#166534,#16a34a);'
    f'color:white;padding:7px 18px;border-radius:20px;font-size:12px;font-weight:600;">'
    f'Pledge 6 Tracker</div>'
    f'</div>',
    unsafe_allow_html=True
)

# ─────────────────────────────────────────────────────────────────────────────
# DATA LOADING
# ─────────────────────────────────────────────────────────────────────────────
@st.cache_data(show_spinner=False)
def load_all_data(file_bytes):
    src = io.BytesIO(file_bytes)
    xl  = pd.ExcelFile(src)
    all_sheets = xl.sheet_names
    sheets = []

    for year in [2023, 2024, 2025, 2026]:
        sname = f'loadplan-{year}'
        if sname in all_sheets:
            try:
                src.seek(0)
                tmp = pd.read_excel(src, sheet_name=sname)
                tmp['Year'] = year
                sheets.append(tmp)
            except Exception as e:
                st.warning(f"Sheet {sname}: {e}")

    if not sheets:
        for sname in all_sheets:
            try:
                src.seek(0)
                tmp = pd.read_excel(src, sheet_name=sname)
                import re
                yr_match = re.search(r'(20\d{2})', sname)
                tmp['Year'] = int(yr_match.group(1)) if yr_match else 0
                sheets.append(tmp)
            except Exception as e:
                st.warning(f"Sheet {sname}: {e}")

    if not sheets:
        st.error("Tidak ada sheet yang berhasil dibaca.")
        st.stop()

    df = pd.concat(sheets, ignore_index=True)
    df['Date']      = pd.to_datetime(df['Date'], errors='coerce')
    df['Month']     = df['Date'].dt.to_period('M')
    df['MonthNum']  = df['Date'].dt.month
    df['Quarter']   = df['Date'].dt.to_period('Q')
    df['YearMonth'] = df['Date'].dt.strftime('%Y-%m')
    df['MonthName'] = df['Date'].dt.strftime('%b')

    df['Material_Clean'] = df['Material'].astype(str).str.strip().str.rstrip('`').str.strip()
    df['Material_Label'] = df['Material_Clean'].map(MATERIAL_MAP).fillna(df['Material_Clean'])

    for c in ['Length', 'Width', 'Height', 'Net Weight', 'Gross Wight', 'Pack', 'Pak']:
        if c in df.columns:
            df[c] = pd.to_numeric(df[c], errors='coerce').fillna(0)
        else:
            df[c] = 0

    df['Volume_mm3']    = df['Length'] * df['Width'] * df['Height']
    df['Volume_m3']     = df['Volume_mm3'] / 1e9
    df['Max_Dimension'] = df[['Length', 'Width', 'Height']].max(axis=1)
    df['Dim_Category']  = pd.cut(
        df['Max_Dimension'],
        bins=[0, 1000, 2000, 4000, 7000, float('inf')],
        labels=DIM_LABELS
    )

    if 'Model' in df.columns:
        df['Pack_Type']  = df['Model'].astype(str).str.extract(r'([PC])`')[0].map(
            {'P': 'Pallet (P)', 'C': 'Carton/Box (C)'}).fillna('Unknown')
        df['Pack_Count'] = pd.to_numeric(
            df['Model'].astype(str).str.extract(r'(\d+)')[0], errors='coerce').fillna(0)
    else:
        df['Pack_Type']  = 'Unknown'
        df['Pack_Count'] = 0

    df['Is_Wood'] = df['Material_Clean'] == 'W'
    df['Item']    = df['Unit'].astype(str).str.strip() if 'Unit' in df.columns else ''

    df['Overpack_Candidate'] = (
        df['Is_Wood'] &
        (df['Max_Dimension'] < 1500) &
        (df['Net Weight'] < 50) &
        (df['Max_Dimension'] > 0)
    )
    return df.dropna(subset=['Date']).reset_index(drop=True)


# ─────────────────────────────────────────────────────────────────────────────
# UPLOAD GATE — simple, centered, no sidebar clutter
# ─────────────────────────────────────────────────────────────────────────────
LOCAL_FILE = 'loadplan_2023-2026.xlsx'
_has_local = os.path.exists(LOCAL_FILE)

if _has_local and "uploaded_bytes" not in st.session_state:
    with open(LOCAL_FILE, 'rb') as f:
        st.session_state["uploaded_bytes"] = f.read()
        st.session_state["uploaded_name"]  = LOCAL_FILE

if "uploaded_bytes" not in st.session_state:
    # ── LANDING PAGE ──────────────────────────────────────────────────────
    _, col_c, _ = st.columns([1, 2, 1])
    with col_c:
        st.markdown("<br><br>", unsafe_allow_html=True)
        st.markdown(
            '<div style="text-align:center;padding:20px 0 32px;">'
            '<div style="font-size:56px;margin-bottom:12px;">📦</div>'
            '<div style="font-size:24px;font-weight:700;color:#166534;margin-bottom:6px;">'
            'Packaging Analysis Dashboard</div>'
            '<div style="font-size:14px;color:#64748b;line-height:1.7;">'
            'Analisis data packaging PT Güntner untuk mendukung<br>'
            '<b>Pledge 6: Make all our packaging fully sustainable</b></div>'
            '</div>',
            unsafe_allow_html=True
        )

        uploaded = st.file_uploader(
            "Upload file Excel Load Plan (.xlsx / .xls)",
            type=["xlsx", "xls"],
            label_visibility="collapsed",
            key="landing_uploader"
        )
        st.markdown(
            '<div style="text-align:center;font-size:12px;color:#9ca3af;margin-top:6px;">'
            'Format: .xlsx atau .xls &nbsp;·&nbsp; Sheet: loadplan-2023 s/d loadplan-2026</div>',
            unsafe_allow_html=True
        )

        if uploaded:
            st.session_state["uploaded_bytes"] = uploaded.read()
            st.session_state["uploaded_name"]  = uploaded.name
            st.rerun()
    st.stop()

# ── Load data ───────────────────────────────────────────────────────────────
with st.spinner('Memuat data…'):
    try:
        df_raw = load_all_data(st.session_state["uploaded_bytes"])
    except Exception as e:
        st.error(f"Gagal membaca file: {e}")
        if st.button("Upload Ulang"):
            del st.session_state["uploaded_bytes"]
            st.rerun()
        st.stop()

mat_opts   = ['Wood', 'Carton Box', 'Cardboard', 'Carton']
year_opts  = sorted(df_raw['Year'].unique().tolist())

# Tombol ganti file — kecil, di kanan atas setelah header
col_info, col_btn = st.columns([6, 1])
with col_info:
    st.caption(f"{st.session_state.get('uploaded_name','—')}  ·  {len(df_raw):,} rows  ·  Year {min(year_opts)}–{max(year_opts)}")
with col_btn:
    if st.button("Ganti File", use_container_width=True):
        del st.session_state["uploaded_bytes"]
        st.cache_data.clear()
        st.rerun()

# ─────────────────────────────────────────────────────────────────────────────
# HELPER FUNCTIONS
# ─────────────────────────────────────────────────────────────────────────────
def kpi_card(title, value, sub, color, icon="", pct=None):
    pct_html = (f'<span style="font-size:15px;font-weight:600;color:{color};opacity:.75;margin-left:6px;">({pct})</span>'
                if pct else "")
    icon_html = f'<span style="font-size:10px;font-weight:700;color:{color};opacity:.6;margin-right:4px;">{icon}</span>' if icon else ""
    return f"""
    <div style="background:white;border-radius:14px;padding:18px 20px;
    box-shadow:0 2px 10px rgba(0,0,0,.08);border-top:4px solid {color};height:100%;">
      <div style="font-size:11px;font-weight:600;color:#9ca3af;letter-spacing:.8px;
      text-transform:uppercase;margin-bottom:5px;">{icon_html}{title}</div>
      <div style="font-size:30px;font-weight:700;color:{color};line-height:1.1;">{value}{pct_html}</div>
      <div style="font-size:11px;color:#64748b;margin-top:4px;">{sub}</div>
    </div>"""

def trend_badge(pct, label="vs avg historis"):
    if pct is None: return ""
    col = "#16a34a" if pct < 0 else "#dc2626"
    bg  = "#f0fdf4" if pct < 0 else "#fef2f2"
    arr = "▼" if pct < 0 else "▲"
    return (f'<span style="background:{bg};color:{col};font-weight:700;font-size:12px;'
            f'padding:3px 9px;border-radius:20px;">{arr} {abs(pct):.1f}% {label}</span>')

def axis_note(x_label, y_label, note=""):
    extra = f" &nbsp;·&nbsp; {note}" if note else ""
    return (f'<div class="axis-note">Sumbu X: <b>{x_label}</b> &nbsp;|&nbsp; '
            f'Sumbu Y: <b>{y_label}</b>{extra}</div>')

def fmt_month_str(ym_str):
    try:
        return pd.Period(ym_str, freq='M').to_timestamp().strftime('%b %Y')
    except Exception:
        return str(ym_str)

def filter_row_year_type(tab_key):
    """Render compact filter row: Year dropdown + Yearly/Monthly toggle. Returns (year_val, view_type)."""
    col_spacer, col_yr, col_type = st.columns([5, 1, 1])
    with col_yr:
        year_val = st.selectbox("Year", ['All'] + [str(y) for y in year_opts], key=f'{tab_key}_year')
    with col_type:
        view_type = st.selectbox("View", ['Yearly', 'Monthly'], key=f'{tab_key}_view')
    return year_val, view_type

def apply_year_filter(df_in, year_val):
    if year_val == 'All':
        return df_in.copy()
    return df_in[df_in['Year'] == int(year_val)].copy()

# ─────────────────────────────────────────────────────────────────────────────
# XGBOOST FORECAST
# ─────────────────────────────────────────────────────────────────────────────
@st.cache_data(show_spinner=False)
def forecast_xgboost(material_label: str, periods: int = 12):
    try:
        from xgboost import XGBRegressor
    except ImportError:
        return None, None, None, None

    mat_df  = df_raw[df_raw['Material_Label'] == material_label].copy()
    monthly = mat_df.groupby('YearMonth').size().reset_index(name='Count')
    monthly = monthly.sort_values('YearMonth').reset_index(drop=True)
    if len(monthly) < 8:
        return None, None, None, None

    n = len(monthly)
    monthly['idx']   = monthly.index
    monthly['m_sin'] = np.sin(2 * np.pi * monthly.index / 12)
    monthly['m_cos'] = np.cos(2 * np.pi * monthly.index / 12)
    monthly['m_num'] = monthly['YearMonth'].apply(lambda x: int(x[5:7]))

    lags    = [l for l in [1, 2, 3, 6, 12] if l < n]
    windows = [w for w in [3, 6] if w <= n // 2]
    for lag in lags: monthly[f'lag_{lag}'] = monthly['Count'].shift(lag)
    for w in windows:
        monthly[f'rmean_{w}'] = monthly['Count'].rolling(w, min_periods=1).mean()
        monthly[f'rstd_{w}']  = monthly['Count'].rolling(w, min_periods=1).std().fillna(0)

    feat_cols = (['idx','m_sin','m_cos','m_num'] +
                 [f'lag_{l}' for l in lags] +
                 [f'rmean_{w}' for w in windows] +
                 [f'rstd_{w}' for w in windows])
    feat_cols = [c for c in feat_cols if c in monthly.columns]
    clean = monthly.dropna().copy()
    if len(clean) < 5: return None, None, None, None

    model = XGBRegressor(n_estimators=300, learning_rate=0.04, max_depth=4,
                         subsample=0.8, colsample_bytree=0.8,
                         min_child_weight=2, random_state=42, verbosity=0)
    model.fit(clean[feat_cols], clean['Count'])

    last_ym = monthly['YearMonth'].iloc[-1]
    last_y, last_m = int(last_ym[:4]), int(last_ym[5:7])
    future_yms = []
    for i in range(1, periods + 1):
        m_new = (last_m - 1 + i) % 12 + 1
        y_new = last_y + (last_m - 1 + i) // 12
        future_yms.append(f"{y_new}-{m_new:02d}")

    all_counts = list(monthly['Count'].values)
    predictions = []
    for i, fym in enumerate(future_yms):
        base_idx = n + i
        feat = {'idx': base_idx,
                'm_sin': np.sin(2 * np.pi * base_idx / 12),
                'm_cos': np.cos(2 * np.pi * base_idx / 12),
                'm_num': int(fym[5:7])}
        for lag in lags:
            feat[f'lag_{lag}'] = all_counts[-lag] if lag <= len(all_counts) else np.mean(all_counts)
        for w in windows:
            feat[f'rmean_{w}'] = np.mean(all_counts[-w:])
            feat[f'rstd_{w}']  = np.std(all_counts[-w:]) if len(all_counts) >= w else 0
        pred = max(0.0, float(model.predict(pd.DataFrame([feat])[feat_cols])[0]))
        predictions.append(pred)
        all_counts.append(pred)

    avg_hist = float(np.mean(monthly['Count'].values))
    avg_pred = float(np.mean(predictions))
    pct      = (avg_pred - avg_hist) / avg_hist * 100 if avg_hist > 0 else 0
    fc_df    = pd.DataFrame({'YearMonth': future_yms, 'Forecast': predictions})
    return monthly, fc_df, pct, model

# ─────────────────────────────────────────────────────────────────────────────
# TABS
# ─────────────────────────────────────────────────────────────────────────────
tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "Overview",
    "Distribution Analysis",
    "Deep Dive Analysis",
    "Prediction",
    "AI Recommendation",
])

# ═════════════════════════════════════════════════════════════════════════════
# ═════════════════════════════════════════════════════════════════════════════
# TAB 1 — OVERVIEW  (Executive Summary — Presentasi Manajemen)
# ═════════════════════════════════════════════════════════════════════════════
with tab1:

    # ─────────────────────────────────────────────────────────────────────────
    # FILTER PERIODE — di atas semua konten
    # ─────────────────────────────────────────────────────────────────────────
    _fspc, _ftype, _fyr = st.columns([3, 1, 1])

    with _ftype:
        ov_view = st.selectbox(
            "Period Type",
            ["Yearly", "Monthly"],
            key="ov_view"
        )

    with _fyr:
        if ov_view == "Yearly":
            ov_year = st.selectbox(
                "Select Year",
                ["All"] + [str(y) for y in sorted(year_opts)],
                key="ov_year"
            )
            ov_label = f"{min(year_opts)}–2026" if ov_year == "All" else f"Year {ov_year}"
        else:
            _all_yms_ov   = sorted(df_raw['YearMonth'].unique(), reverse=True)
            _ym_labels_ov = [fmt_month_str(ym) for ym in _all_yms_ov]
            _ym_map_ov    = dict(zip(_ym_labels_ov, _all_yms_ov))
            ov_month_label = st.selectbox("Select Month", _ym_labels_ov, key="ov_month")
            ov_ym_sel = _ym_map_ov[ov_month_label]
            ov_label  = ov_month_label

    # ── FILTER DATAFRAME SESUAI PILIHAN ──────────────────────────────────────
    if ov_view == "Yearly" and ov_year != "All":
        df_all = df_raw[df_raw['Year'] == int(ov_year)].copy()
    elif ov_view == "Monthly":
        df_all = df_raw[df_raw['YearMonth'] == ov_ym_sel].copy()
    else:
        df_all = df_raw.copy()




    # ── KALKULASI STATISTIK (semua ikut filter df_all) ───────────────────────
    total_l          = len(df_all)
    wood_n           = int((df_all['Material_Clean'] == 'W').sum())
    wood_pct_all     = wood_n / total_l * 100 if total_l else 0
    cb_n             = int((df_all['Material_Clean'] == 'B').sum())
    cb_pct_all       = cb_n / total_l * 100 if total_l else 0
    cbd_n_all        = int((df_all['Material_Clean'] == 'Cb').sum())
    ct_n_all         = int((df_all['Material_Clean'] == 'Ct').sum())
    non_wood_n       = cb_n + cbd_n_all + ct_n_all
    non_wood_pct_all = non_wood_n / total_l * 100 if total_l else 0
    overpack_n_all   = int(df_all['Overpack_Candidate'].sum())
    overpack_pct_all = overpack_n_all / wood_n * 100 if wood_n else 0

    xs_wood_n  = df_all[(df_all['Dim_Category'] == 'XS (<1 m)') & df_all['Is_Wood']].shape[0]
    xs_total_n = df_all[df_all['Dim_Category'] == 'XS (<1 m)'].shape[0]
    s_wood_n   = df_all[(df_all['Dim_Category'] == 'S (1\u20132 m)') & df_all['Is_Wood']].shape[0]
    s_total_n  = df_all[df_all['Dim_Category'] == 'S (1\u20132 m)'].shape[0]
    xs_pct_val = xs_wood_n / xs_total_n * 100 if xs_total_n else 0
    s_pct_val  = s_wood_n  / s_total_n  * 100 if s_total_n  else 0

    top_cust_all  = df_all[df_all['Is_Wood']].groupby('Name').size().sort_values(ascending=False)
    top_cust_name = top_cust_all.index[0] if len(top_cust_all) else '\u2014'

    overpack_items = (df_all[df_all['Overpack_Candidate']]
                      .groupby('Item').size().sort_values(ascending=False)
                      .reset_index(name='Lines'))

    cust_wood_pct = (df_all.groupby('Name')
                     .apply(lambda g: pd.Series({
                         'Lines':    len(g),
                         'Wood':     g['Is_Wood'].sum(),
                         'Pct_Wood': g['Is_Wood'].mean() * 100
                     })).reset_index()
                     .sort_values('Lines', ascending=False)
                     .head(10))

    items_wood_rank = (df_all[df_all['Is_Wood']].groupby('Item').size()
                       .sort_values(ascending=False).reset_index(name='Wood Shipments'))

    n_sm         = int(df_all[df_all['Dim_Category'].isin(['S (1\u20132 m)', 'M (2\u20134 m)'])]['Is_Wood'].sum())
    progress_pct = max(0.0, min(100.0, 100.0 - wood_pct_all))

    # Tren Wood — dihitung dinamis sesuai mode filter
    wood_2023 = int(((df_raw['Material_Clean'] == 'W') & (df_raw['Year'] == 2023)).sum())
    wood_2025 = int(((df_raw['Material_Clean'] == 'W') & (df_raw['Year'] == 2025)).sum())
    wood_yoy  = (wood_2025 - wood_2023) / wood_2023 * 100 if wood_2023 else 0

    # Tren Yearly: vs tahun sebelumnya
    if ov_view == "Yearly" and ov_year != "All":
        _yr_curr = int(ov_year)
        _yr_prev = _yr_curr - 1
        _wood_curr = int(((df_raw['Material_Clean'] == 'W') & (df_raw['Year'] == _yr_curr)).sum())
        _tot_curr  = int((df_raw['Year'] == _yr_curr).sum())
        _wood_prev = int(((df_raw['Material_Clean'] == 'W') & (df_raw['Year'] == _yr_prev)).sum())
        _tot_prev  = int((df_raw['Year'] == _yr_prev).sum())
        _pct_curr  = _wood_curr / _tot_curr  * 100 if _tot_curr  else 0
        _pct_prev  = _wood_prev / _tot_prev  * 100 if _tot_prev  else 0
        _tren_yr   = _pct_curr - _pct_prev
        _has_prev_yr = _yr_prev in df_raw['Year'].unique()
    else:
        _has_prev_yr = False
        _tren_yr = 0
        _yr_prev = None
        _pct_curr = 0
        _pct_prev = 0
        _yr_curr = 0

    # Tren Monthly: vs bulan sebelumnya
    if ov_view == "Monthly":
        import calendar
        _ym_curr = ov_ym_sel  # format 'YYYY-MM'
        _y, _m   = int(_ym_curr[:4]), int(_ym_curr[5:7])
        _m_prev  = _m - 1 if _m > 1 else 12
        _y_prev  = _y if _m > 1 else _y - 1
        _ym_prev = f"{_y_prev}-{_m_prev:02d}"
        _wood_m_curr = int(((df_raw['Material_Clean'] == 'W') & (df_raw['YearMonth'] == _ym_curr)).sum())
        _tot_m_curr  = int((df_raw['YearMonth'] == _ym_curr).sum())
        _wood_m_prev = int(((df_raw['Material_Clean'] == 'W') & (df_raw['YearMonth'] == _ym_prev)).sum())
        _tot_m_prev  = int((df_raw['YearMonth'] == _ym_prev).sum())
        _pct_m_curr  = _wood_m_curr / _tot_m_curr * 100 if _tot_m_curr else 0
        _pct_m_prev  = _wood_m_prev / _tot_m_prev * 100 if _tot_m_prev else 0
        _tren_mo     = _pct_m_curr - _pct_m_prev
        _has_prev_mo = _ym_prev in df_raw['YearMonth'].values
        _prev_mo_label = fmt_month_str(_ym_prev)
    else:
        _has_prev_mo = False
        _tren_mo = 0
        _prev_mo_label = ""

    # ─────────────────────────────────────────────────────────────────────────
    # BAGIAN 1 — KPI SCORECARD
    # ─────────────────────────────────────────────────────────────────────────
    st.markdown(
        f"<div style='font-size:13px;color:#64748b;margin-bottom:8px;'>"
        f"Periode: <b>{ov_label}</b> &nbsp;&middot;&nbsp; {total_l:,} pengiriman</div>",
        unsafe_allow_html=True)

    k1, k2, k3, k4 = st.columns(4)

    with k1:
        st.markdown(kpi_card(
            "Total Shipments", f"{total_l:,}",
            f"Periode: {ov_label}", "#2563eb"),
            unsafe_allow_html=True)

    with k2:
        st.markdown(kpi_card(
            "Wood Shipments", f"{wood_n:,}",
            f"dari {total_l:,} total pengiriman", "#7C4A19",
            pct=f"{wood_pct_all:.1f}%"),
            unsafe_allow_html=True)

    with k3:
        st.markdown(kpi_card(
            "Switchable to Carton", f"{overpack_n_all:,}",
            f"dari {wood_n:,} Wood · dimensi ≤2m & berat ≤75kg", "#d97706",
            pct=f"{overpack_pct_all:.1f}%"),
            unsafe_allow_html=True)

    with k4:
        if ov_view == "Yearly" and ov_year == "All" and wood_2023 > 0:
            # All years — tren 2023 vs 2025
            arrow     = "\u25bc" if wood_yoy < 0 else "\u25b2"
            color_yoy = GREEN_MID if wood_yoy < 0 else "#dc2626"
            st.markdown(kpi_card(
                "Wood Trend 2023→2025", f"{arrow} {abs(wood_yoy):.1f}%",
                "▼ Turun = menuju Pledge 6  ·  ▲ Naik = perlu intervensi",
                color_yoy),
                unsafe_allow_html=True)
        elif ov_view == "Yearly" and ov_year != "All" and _has_prev_yr:
            # Specific year — tren vs tahun sebelumnya
            arrow     = "\u25bc" if _tren_yr < 0 else "\u25b2"
            color_tyr = GREEN_MID if _tren_yr < 0 else "#dc2626"
            st.markdown(kpi_card(
                f"Wood Trend vs {_yr_prev}", f"{arrow} {abs(_tren_yr):.1f}%",
                f"{_yr_prev}: {_pct_prev:.1f}% → {_yr_curr}: {_pct_curr:.1f}% Wood",
                color_tyr),
                unsafe_allow_html=True)
        elif ov_view == "Monthly" and _has_prev_mo:
            # Specific month — tren vs bulan sebelumnya
            arrow     = "\u25bc" if _tren_mo < 0 else "\u25b2"
            color_tmo = GREEN_MID if _tren_mo < 0 else "#dc2626"
            st.markdown(kpi_card(
                f"Wood Trend vs {_prev_mo_label}", f"{arrow} {abs(_tren_mo):.1f}%",
                f"{_prev_mo_label}: {_pct_m_prev:.1f}% → {fmt_month_str(ov_ym_sel)}: {_pct_m_curr:.1f}% Wood",
                color_tmo),
                unsafe_allow_html=True)
        else:
            # Fallback — tidak ada data pembanding
            st.markdown(kpi_card(
                "Sustainable Packaging", f"{non_wood_n:,}",
                f"dari {total_l:,} · Carton Box + Cardboard + Carton",
                "#16a34a",
                pct=f"{non_wood_pct_all:.1f}%"),
                unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    # Pledge 6 Progress Bar
    st.markdown(f"""
    <div style="background:white;border-radius:14px;padding:18px 24px;
    box-shadow:0 2px 10px rgba(0,0,0,.07);border:1px solid #e5e7eb;margin-bottom:8px;">
      <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:10px;">
        <div style="font-size:14px;font-weight:700;color:#166534;">
          Pledge 6 Progress &mdash; Target: 0% Wood Packaging
        </div>
        <div style="font-size:12.5px;color:#64748b;">
          Periode: <b>{ov_label}</b>
          &nbsp;&middot;&nbsp; Saat ini: <b style="color:#dc2626;">{wood_pct_all:.1f}% Wood</b>
          &nbsp;&middot;&nbsp; Target: <b style="color:#16a34a;">0% Wood</b>
        </div>
      </div>
      <div style="background:#f1f5f9;border-radius:10px;height:22px;overflow:hidden;position:relative;">
        <div style="background:linear-gradient(90deg,#16a34a,#4ade80);height:100%;
        width:{progress_pct:.1f}%;border-radius:10px;"></div>
        <div style="position:absolute;top:3px;left:50%;transform:translateX(-50%);
        font-size:12px;font-weight:700;color:#166534;">{progress_pct:.1f}% tercapai</div>
      </div>
      <div style="font-size:12px;color:#64748b;margin-top:8px;display:flex;gap:24px;">
        <span>Quick win: konversi <b>{overpack_n_all:,} item kecil</b> dari Wood ke Carton</span>
        <span>&rarr; tambah <b>~{overpack_pct_all:.1f}%</b> kemajuan</span>
        <span>&rarr; sisa gap: <b>{wood_pct_all:.1f}%</b> masih perlu direduksi</span>
      </div>
    </div>""", unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    # ─────────────────────────────────────────────────────────────────────────
    # BAGIAN 2 — KOMPOSISI & TREN
    # ─────────────────────────────────────────────────────────────────────────
    st.markdown("### Packaging Composition & Trend")

    col_a, col_b = st.columns([1, 1.8])

    with col_a:
        mat_dist = df_all['Material_Label'].value_counts().reset_index()
        mat_dist.columns = ['Material', 'Count']
        fig_donut = go.Figure(go.Pie(
            labels=mat_dist['Material'], values=mat_dist['Count'], hole=0.58,
            marker=dict(colors=[MATERIAL_COLORS.get(m, '#aaa') for m in mat_dist['Material']],
                        line=dict(color='white', width=2)),
            textinfo='percent+label', textfont=dict(size=11),
            hovertemplate='<b>%{label}</b><br>%{value:,} item (%{percent})<extra></extra>'
        ))
        fig_donut.add_annotation(
            text=f"<b>{wood_pct_all:.0f}%</b><br><span style='font-size:10px'>Wood</span>",
            x=0.5, y=0.5, font=dict(size=16, color='#7C4A19'), showarrow=False)
        fig_donut.update_layout(
            title=f"Komposisi Kemasan \u2014 {ov_label}",
            paper_bgcolor='white', height=300,
            margin=dict(l=10, r=10, t=40, b=10),
            showlegend=True,
            legend=dict(orientation='v', x=1.02, y=0.5, font=dict(size=11)))
        st.plotly_chart(fig_donut, use_container_width=True, config={'displayModeBar': False})

    with col_b:
        if ov_view == "Monthly":
            fig_trend = go.Figure(go.Bar(
                x=mat_dist['Material'],
                y=mat_dist['Count'],
                marker_color=[MATERIAL_COLORS.get(m, '#aaa') for m in mat_dist['Material']],
                text=mat_dist['Count'].apply(lambda v: f"{v:,}"),
                textposition='outside',
                hovertemplate='<b>%{x}</b><br>%{y:,} lines<extra></extra>'
            ))
            fig_trend.update_layout(
                title=f"Distribusi Material \u2014 {ov_label}",
                paper_bgcolor='white', plot_bgcolor='white', height=300,
                margin=dict(l=20, r=20, t=40, b=30),
                xaxis=dict(showgrid=False),
                yaxis=dict(showgrid=True, gridcolor='#f1f5f9', title='Shipments'))
        else:
            mn_ov = df_all.groupby(['YearMonth', 'Material_Label']).size().reset_index(name='Count')
            mn_ov['Label'] = mn_ov['YearMonth'].apply(fmt_month_str)
            mn_ov = mn_ov.sort_values('YearMonth')
            fig_trend = go.Figure()
            for mat in mat_opts:
                sub = mn_ov[mn_ov['Material_Label'] == mat]
                fig_trend.add_trace(go.Scatter(
                    x=sub['Label'], y=sub['Count'], name=mat,
                    mode='lines+markers',
                    line=dict(color=MATERIAL_COLORS.get(mat), width=2),
                    marker=dict(size=3.5)))
            fig_trend.update_layout(
                title=f"Monthly Trend by Material — {ov_label}",
                paper_bgcolor='white', plot_bgcolor='white', height=300,
                margin=dict(l=20, r=20, t=40, b=55),
                xaxis=dict(showgrid=False, tickangle=-45, tickfont=dict(size=8), title='Month'),
                yaxis=dict(showgrid=True, gridcolor='#f1f5f9', title='Shipments'),
                legend=dict(orientation='h', y=1.12, x=1, xanchor='right'),
                hovermode='x unified')
        st.plotly_chart(fig_trend, use_container_width=True, config={'displayModeBar': False})

    cb_sig = "hanya signifikan di kategori kecil (XS & S)" if cb_pct_all < 30 else f"cukup signifikan ({cb_pct_all:.1f}%)"
    st.markdown(
        f'<div class="tooltip-box"><b>Insight ({ov_label}):</b> Wood mendominasi '
        f'<b>{wood_pct_all:.1f}%</b> dari total shipment. Carton Box {cb_sig}. '
        f'Peluang terbesar ada di kategori ukuran kecil '
        f'(XS: {xs_pct_val:.0f}% Wood, S: {s_pct_val:.0f}% Wood).</div>',
        unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    # Tabel volume per tahun — hanya tampil saat mode All Years
    if ov_view == "Yearly" and ov_year == "All":
        st.markdown("**Volume Pengiriman per Tahun per Packaging Type**")
        yr_totals     = {yr: int((df_raw['Year'] == yr).sum()) for yr in year_opts}
        total_raw     = len(df_raw)
        summary_rows2 = []
        for mat in mat_opts:
            row2 = {'Kemasan': mat}
            total_mat2 = 0
            for yr in year_opts:
                cnt2   = int(((df_raw['Material_Label'] == mat) & (df_raw['Year'] == yr)).sum())
                yr_tot = yr_totals[yr]
                pct2   = cnt2 / yr_tot * 100 if yr_tot else 0
                row2[str(yr)] = f"{cnt2:,} ({pct2:.0f}%)"
                total_mat2   += cnt2
            row2['Total'] = f"{total_mat2:,}"
            row2['%']     = f"{total_mat2 / total_raw * 100:.1f}%"
            summary_rows2.append(row2)
        st.dataframe(pd.DataFrame(summary_rows2), use_container_width=True, hide_index=True, height=195)
        st.markdown("<br>", unsafe_allow_html=True)

    # ─────────────────────────────────────────────────────────────────────────
    # BAGIAN 2B — GRAFIK TREN PER MATERIAL (MASING-MASING TERPISAH)
    # ─────────────────────────────────────────────────────────────────────────
    st.markdown("### Monthly Trend per Packaging Type")

    # Build monthly counts per material for the filtered period
    _mn_sep = df_all.groupby(['YearMonth', 'Material_Label']).size().reset_index(name='Count')
    _mn_sep['Label'] = _mn_sep['YearMonth'].apply(fmt_month_str)
    _mn_sep = _mn_sep.sort_values('YearMonth')

    # If Monthly view only one month — show a note
    if ov_view == "Monthly":
        st.info("Grafik tren per material membutuhkan lebih dari satu bulan. "
                "Pilih 'Yearly' atau 'All' untuk melihat tren.")
    else:
        # ── Row 1: Wood + Carton Box ──────────────────────────────────────
        _col_w, _col_cb = st.columns(2)

        for _col, _mat in zip([_col_w, _col_cb], ['Wood', 'Carton Box']):
            with _col:
                _sub = _mn_sep[_mn_sep['Material_Label'] == _mat].copy()
                _color = MATERIAL_COLORS.get(_mat, '#888')
                _fig = go.Figure()
                _fig.add_trace(go.Scatter(
                    x=_sub['Label'], y=_sub['Count'],
                    name=_mat,
                    mode='lines+markers',
                    fill='tozeroy',
                    fillcolor=MATERIAL_LIGHT.get(_mat, '#eee'),
                    line=dict(color=_color, width=2.5),
                    marker=dict(size=5, color=_color),
                    hovertemplate='<b>%{x}</b><br>%{y:,} lines<extra></extra>'
                ))
                _avg = _sub['Count'].mean() if not _sub.empty else 0
                _fig.add_hline(y=_avg, line_dash='dot', line_color='#94a3b8',
                               annotation_text=f"Avg: {_avg:,.0f}",
                               annotation_position='bottom right',
                               annotation_font_size=10)
                _fig.update_layout(
                    title=dict(text=f"{_mat} — Monthly Trend ({ov_label})",
                               font=dict(size=13, color=_color)),
                    paper_bgcolor='white', plot_bgcolor='white', height=260,
                    margin=dict(l=20, r=20, t=45, b=55),
                    xaxis=dict(showgrid=False, tickangle=-45,
                               tickfont=dict(size=8), title='Month'),
                    yaxis=dict(showgrid=True, gridcolor='#f1f5f9',
                               title='Shipments', rangemode='tozero'),
                    showlegend=False,
                    hovermode='x unified')
                st.plotly_chart(_fig, use_container_width=True,
                                config={'displayModeBar': False})

        # ── Row 2: Cardboard + Carton ────────────────────────────────────
        _col_cbd, _col_ct = st.columns(2)

        for _col, _mat in zip([_col_cbd, _col_ct], ['Cardboard', 'Carton']):
            with _col:
                _sub = _mn_sep[_mn_sep['Material_Label'] == _mat].copy()
                _color = MATERIAL_COLORS.get(_mat, '#888')
                _fig = go.Figure()
                _fig.add_trace(go.Scatter(
                    x=_sub['Label'], y=_sub['Count'],
                    name=_mat,
                    mode='lines+markers',
                    fill='tozeroy',
                    fillcolor=MATERIAL_LIGHT.get(_mat, '#eee'),
                    line=dict(color=_color, width=2.5),
                    marker=dict(size=5, color=_color),
                    hovertemplate='<b>%{x}</b><br>%{y:,} lines<extra></extra>'
                ))
                _avg = _sub['Count'].mean() if not _sub.empty else 0
                _fig.add_hline(y=_avg, line_dash='dot', line_color='#94a3b8',
                               annotation_text=f"Avg: {_avg:,.0f}",
                               annotation_position='bottom right',
                               annotation_font_size=10)
                _fig.update_layout(
                    title=dict(text=f"{_mat} — Monthly Trend ({ov_label})",
                               font=dict(size=13, color=_color)),
                    paper_bgcolor='white', plot_bgcolor='white', height=260,
                    margin=dict(l=20, r=20, t=45, b=55),
                    xaxis=dict(showgrid=False, tickangle=-45,
                               tickfont=dict(size=8), title='Month'),
                    yaxis=dict(showgrid=True, gridcolor='#f1f5f9',
                               title='Shipments', rangemode='tozero'),
                    showlegend=False,
                    hovermode='x unified')
                st.plotly_chart(_fig, use_container_width=True,
                                config={'displayModeBar': False})

    st.markdown("<br>", unsafe_allow_html=True)

    # ─────────────────────────────────────────────────────────────────────────
    # BAGIAN 2C — ANALISIS KORELASI: WOOD vs METAL USAGE & PRODUCTION OUTPUT
    # ─────────────────────────────────────────────────────────────────────────
    with st.expander("Correlation Analysis: Wood Shipments vs Metal Usage & Production Output", expanded=True):

        # ── Build base data ───────────────────────────────────────────────
        _wood_monthly = (
            df_raw[df_raw['Material_Clean'] == 'W']
            .groupby('YearMonth').size()
            .reset_index(name='WoodLines')
            .sort_values('YearMonth')
        )
        _wood_monthly['Label'] = _wood_monthly['YearMonth'].apply(fmt_month_str)
        _wood_ref = pd.merge(_wood_monthly, REF_DF, on='YearMonth', how='outer').sort_values('YearMonth')
        _wood_ref['WoodLines']    = _wood_ref['WoodLines'].fillna(0)
        _wood_ref['DisplayLabel'] = _wood_ref['Label_x'].fillna(_wood_ref['Label_y'])

        # Apply year filter
        if ov_view == "Yearly" and ov_year != "All":
            _wrp = _wood_ref[_wood_ref['YearMonth'].str.startswith(str(ov_year))].copy()
        else:
            _wrp = _wood_ref.copy()

        _wrp_clean = _wrp.dropna(subset=['WoodLines', 'MetalUsageKG', 'ProductionUnit']).copy()
        _wrp_clean = _wrp_clean[_wrp_clean['WoodLines'] > 0].reset_index(drop=True)

        # ── Compute correlations ──────────────────────────────────────────
        _corr_metal = _wrp_clean['WoodLines'].corr(_wrp_clean['MetalUsageKG']) if len(_wrp_clean) >= 3 else 0
        _corr_prod  = _wrp_clean['WoodLines'].corr(_wrp_clean['ProductionUnit']) if len(_wrp_clean) >= 3 else 0

        def _interp_r(r):
            a = abs(r)
            direction = "positif" if r >= 0 else "negatif"
            if a >= 0.7:   strength, badge = "STRONG", "🟢"
            elif a >= 0.4: strength, badge = "MODERATE", "🟡"
            else:          strength, badge = "WEAK", "🔴"
            return direction, strength, badge

        _dir_m, _str_m, _badge_m = _interp_r(_corr_metal)
        _dir_p, _str_p, _badge_p = _interp_r(_corr_prod)

        # ── Direction Agreement (naik/turun bareng) ───────────────────────
        if len(_wrp_clean) >= 3:
            _da = _wrp_clean.copy()
            _da['dW'] = _da['WoodLines'].diff()
            _da['dM'] = _da['MetalUsageKG'].diff()
            _da['dP'] = _da['ProductionUnit'].diff()
            _da = _da.dropna(subset=['dW','dM','dP'])
            _agree_m = int((_da['dW'].apply(np.sign) == _da['dM'].apply(np.sign)).sum())
            _agree_p = int((_da['dW'].apply(np.sign) == _da['dP'].apply(np.sign)).sum())
            _total_m = len(_da)
            _pct_agree_m = _agree_m / _total_m * 100
            _pct_agree_p = _agree_p / _total_m * 100
        else:
            _agree_m = _agree_p = _total_m = 0
            _pct_agree_m = _pct_agree_p = 0

        # ── KPI Summary Cards ─────────────────────────────────────────────
        st.markdown("##### Wood vs Metal & Production — Trend Summary")

        _kk1, _kk2, _kk3, _kk4 = st.columns(4)
        with _kk1:
            _cm_col = '#16a34a' if _pct_agree_m >= 60 else ('#d97706' if _pct_agree_m >= 45 else '#dc2626')
            st.markdown(f"""
            <div style="background:white;border-radius:12px;padding:16px 18px;
            box-shadow:0 2px 8px rgba(0,0,0,.08);border-top:4px solid {_cm_col};text-align:center;">
              <div style="font-size:11px;color:#9ca3af;font-weight:600;text-transform:uppercase;
              letter-spacing:.7px;margin-bottom:6px;">Wood &amp; Metal Trend</div>
              <div style="font-size:32px;font-weight:800;color:{_cm_col};">{_pct_agree_m:.0f}%</div>
              <div style="font-size:12px;color:#64748b;margin-top:4px;">
                {_agree_m} dari {_total_m} bulan<br>Wood &amp; Metal bergerak searah</div>
            </div>""", unsafe_allow_html=True)
        with _kk2:
            _cp_col = '#16a34a' if _pct_agree_p >= 60 else ('#d97706' if _pct_agree_p >= 45 else '#dc2626')
            st.markdown(f"""
            <div style="background:white;border-radius:12px;padding:16px 18px;
            box-shadow:0 2px 8px rgba(0,0,0,.08);border-top:4px solid {_cp_col};text-align:center;">
              <div style="font-size:11px;color:#9ca3af;font-weight:600;text-transform:uppercase;
              letter-spacing:.7px;margin-bottom:6px;">Wood &amp; Production Trend</div>
              <div style="font-size:32px;font-weight:800;color:{_cp_col};">{_pct_agree_p:.0f}%</div>
              <div style="font-size:12px;color:#64748b;margin-top:4px;">
                {_agree_p} dari {_total_m} bulan<br>Wood &amp; Production bergerak searah</div>
            </div>""", unsafe_allow_html=True)
        with _kk3:
            if len(_wrp_clean):
                _avg_metal_kg = _wrp_clean['MetalUsageKG'].mean()
                _avg_prod     = _wrp_clean['ProductionUnit'].mean()
                st.markdown(f"""
                <div style="background:white;border-radius:12px;padding:16px 18px;
                box-shadow:0 2px 8px rgba(0,0,0,.08);border-top:4px solid #2563eb;text-align:center;">
                  <div style="font-size:11px;color:#9ca3af;font-weight:600;text-transform:uppercase;
                  letter-spacing:.7px;margin-bottom:6px;">Avg Metal / Unit</div>
                  <div style="font-size:28px;font-weight:800;color:#2563eb;">{_avg_metal_kg/_avg_prod:,.1f} KG</div>
                  <div style="font-size:12px;color:#64748b;margin-top:4px;">
                    per unit produksi<br>avg {_avg_metal_kg/1000:,.0f}t metal · {_avg_prod:,.0f} unit/bln</div>
                </div>""", unsafe_allow_html=True)
        with _kk4:
            if len(_wrp_clean):
                _avg_wood = _wrp_clean['WoodLines'].mean()
                st.markdown(f"""
                <div style="background:white;border-radius:12px;padding:16px 18px;
                box-shadow:0 2px 8px rgba(0,0,0,.08);border-top:4px solid #7C4A19;text-align:center;">
                  <div style="font-size:11px;color:#9ca3af;font-weight:600;text-transform:uppercase;
                  letter-spacing:.7px;margin-bottom:6px;">Avg Wood Shipments / Month</div>
                  <div style="font-size:28px;font-weight:800;color:#7C4A19;">{_avg_wood:,.0f}</div>
                  <div style="font-size:12px;color:#64748b;margin-top:4px;">
                    pengiriman per bulan<br>dari {len(_wrp_clean)} bulan data</div>
                </div>""", unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)

        # ════════════════════════════════════════════════════════════════
        # GRAFIK A — NILAI ASLI (dual-axis: bar + lines)
        # ════════════════════════════════════════════════════════════════
        st.markdown("#### Raw Values (Individual Scale)")
        st.markdown(
            '<div class="axis-note">'
            'Sumbu kiri (coklat): <b>Wood Packaging Shipments</b> · '
            'Sumbu kanan (biru & hijau): <b>Metal Usage (KG)</b> dan <b>Production Output (Unit)</b>. '
            'Arahkan kursor ke tiap titik untuk melihat nilai aktual.</div>',
            unsafe_allow_html=True)

        if len(_wrp_clean) >= 2:
            _fig_raw = make_subplots(specs=[[{"secondary_y": True}]])
            _fig_raw.add_trace(go.Bar(
                x=_wrp_clean['DisplayLabel'], y=_wrp_clean['WoodLines'],
                name='Wood Shipments', marker_color='#7C4A19', opacity=0.70,
                hovertemplate='<b>%{x}</b><br>Wood: <b>%{y:,} shipments</b><extra></extra>'
            ), secondary_y=False)
            _fig_raw.add_trace(go.Scatter(
                x=_wrp_clean['DisplayLabel'], y=_wrp_clean['MetalUsageKG'],
                name='Metal Usage (KG)', mode='lines+markers',
                line=dict(color='#2563eb', width=2.5),
                marker=dict(size=5),
                hovertemplate='<b>%{x}</b><br>Metal: <b>%{y:,.0f} KG</b><extra></extra>'
            ), secondary_y=True)
            _fig_raw.add_trace(go.Scatter(
                x=_wrp_clean['DisplayLabel'], y=_wrp_clean['ProductionUnit'],
                name='Production (Unit)', mode='lines+markers',
                line=dict(color='#16a34a', width=2.5, dash='dot'),
                marker=dict(size=5, symbol='diamond'),
                hovertemplate='<b>%{x}</b><br>Production: <b>%{y:,} units</b><extra></extra>'
            ), secondary_y=True)
            _fig_raw.update_yaxes(title_text='Wood Packaging Lines',
                                  showgrid=True, gridcolor='#f1f5f9', secondary_y=False)
            _fig_raw.update_yaxes(title_text='Metal (KG)  /  Produksi (Unit)',
                                  showgrid=False, secondary_y=True)
            _fig_raw.update_layout(
                paper_bgcolor='white', plot_bgcolor='white', height=340,
                margin=dict(l=20, r=20, t=20, b=65),
                xaxis=dict(showgrid=False, tickangle=-45, tickfont=dict(size=8.5), title='Month'),
                legend=dict(orientation='h', y=1.10, x=0.5, xanchor='center', font=dict(size=11)),
                hovermode='x unified', barmode='overlay')
            st.plotly_chart(_fig_raw, use_container_width=True, config={'displayModeBar': False})

        # ════════════════════════════════════════════════════════════════
        # GRAFIK B — INDEKS (Base = 100)
        # ════════════════════════════════════════════════════════════════
        st.markdown("#### Movement Index (Base = 100 at first month)")
        st.markdown(
            '<div class="axis-note">'
            'Nilai asli disetarakan ke skala yang sama: bulan pertama = 100. '
            'Di atas 100 = lebih tinggi dari baseline · di bawah 100 = lebih rendah. '
            '<b>Jika ketiga garis bergerak bersamaan → Wood Shipments mengikuti pola Production & Metal Usage.</b>'
            '</div>',
            unsafe_allow_html=True)

        if len(_wrp_clean) >= 2:
            _idx = _wrp_clean.copy().reset_index(drop=True)
            _base_wood  = _idx['WoodLines'].iloc[0]    or 1
            _base_metal = _idx['MetalUsageKG'].iloc[0]  or 1
            _base_prod  = _idx['ProductionUnit'].iloc[0] or 1
            _idx['idx_wood']  = _idx['WoodLines']      / _base_wood  * 100
            _idx['idx_metal'] = _idx['MetalUsageKG']   / _base_metal * 100
            _idx['idx_prod']  = _idx['ProductionUnit']  / _base_prod  * 100

            _fig_idx = go.Figure()
            _fig_idx.add_trace(go.Scatter(
                x=_idx['DisplayLabel'], y=_idx['idx_wood'],
                name='Wood Shipments', mode='lines+markers',
                line=dict(color='#7C4A19', width=3),
                marker=dict(size=7, color='#7C4A19'),
                hovertemplate=(
                    '<b>%{x}</b><br>'
                    'Wood Index: <b>%{y:.1f}</b><br>'
                    '%{customdata}<extra></extra>'),
                customdata=[
                    f"{'naik' if v > 100 else 'turun'} {abs(v-100):.1f}% dari baseline "
                    f"({int(r):,} shipments)"
                    for v, r in zip(_idx['idx_wood'], _idx['WoodLines'])]))
            _fig_idx.add_trace(go.Scatter(
                x=_idx['DisplayLabel'], y=_idx['idx_metal'],
                name='Metal Usage', mode='lines+markers',
                line=dict(color='#2563eb', width=2.5),
                marker=dict(size=5, color='#2563eb'),
                hovertemplate=(
                    '<b>%{x}</b><br>'
                    'Metal Index: <b>%{y:.1f}</b><br>'
                    '%{customdata}<extra></extra>'),
                customdata=[
                    f"{'naik' if v > 100 else 'turun'} {abs(v-100):.1f}% dari baseline "
                    f"({r:,.0f} KG)"
                    for v, r in zip(_idx['idx_metal'], _idx['MetalUsageKG'])]))
            _fig_idx.add_trace(go.Scatter(
                x=_idx['DisplayLabel'], y=_idx['idx_prod'],
                name='Production Output', mode='lines+markers',
                line=dict(color='#16a34a', width=2.5, dash='dot'),
                marker=dict(size=5, symbol='diamond', color='#16a34a'),
                hovertemplate=(
                    '<b>%{x}</b><br>'
                    'Production Index: <b>%{y:.1f}</b><br>'
                    '%{customdata}<extra></extra>'),
                customdata=[
                    f"{'naik' if v > 100 else 'turun'} {abs(v-100):.1f}% dari baseline "
                    f"({int(r):,} unit)"
                    for v, r in zip(_idx['idx_prod'], _idx['ProductionUnit'])]))
            _fig_idx.add_hline(y=100, line_dash='dot', line_color='#cbd5e1',
                               annotation_text="Baseline (100)",
                               annotation_position='bottom left',
                               annotation_font=dict(size=10, color='#94a3b8'))
            _fig_idx.update_layout(
                paper_bgcolor='white', plot_bgcolor='white', height=340,
                margin=dict(l=20, r=20, t=20, b=65),
                xaxis=dict(showgrid=False, tickangle=-45, tickfont=dict(size=8.5), title='Month'),
                yaxis=dict(showgrid=True, gridcolor='#f1f5f9',
                           title='Indeks  (bulan pertama = 100)'),
                legend=dict(orientation='h', y=1.10, x=0.5, xanchor='center', font=dict(size=11)),
                hovermode='x unified')
            st.plotly_chart(_fig_idx, use_container_width=True, config={'displayModeBar': False})
        else:
            st.info("Data tidak cukup untuk grafik indeks pada periode ini.")

        # ════════════════════════════════════════════════════════════════
        # GRAFIK 2 — KUADRAN SEARAH vs BERLAWANAN (% perubahan, label ramah)
        # ════════════════════════════════════════════════════════════════
        st.markdown("#### Wood Packaging vs Metal & Production — Monthly Direction")
        st.markdown(
            '<div class="axis-note">'
            'Setiap titik mewakili satu bulan. Garis tengah = tidak ada perubahan dari bulan sebelumnya. '
            '<b>Area hijau</b>: Wood dan Metal/Production sama-sama naik atau turun → <b>bergerak searah</b>. '
            '<b>Area merah</b>: salah satu naik, yang lain turun → <b>berlawanan arah</b>. '
            'Arahkan kursor ke titik untuk melihat nilai aktual bulan tersebut.'
            '</div>',
            unsafe_allow_html=True)

        if len(_wrp_clean) >= 3:
            _da2 = _wrp_clean.copy().reset_index(drop=True)
            _da2['pct_W'] = _da2['WoodLines'].pct_change() * 100
            _da2['pct_M'] = _da2['MetalUsageKG'].pct_change() * 100
            _da2['pct_P'] = _da2['ProductionUnit'].pct_change() * 100
            _da2 = _da2.dropna(subset=['pct_W','pct_M','pct_P']).reset_index(drop=True)

            _col_m  = ['#16a34a' if np.sign(w)==np.sign(m) else '#dc2626'
                       for w,m in zip(_da2['pct_W'], _da2['pct_M'])]
            _col_p  = ['#16a34a' if np.sign(w)==np.sign(p) else '#dc2626'
                       for w,p in zip(_da2['pct_W'], _da2['pct_P'])]
            _col_mp = ['#16a34a' if np.sign(m)==np.sign(p) else '#dc2626'
                       for m,p in zip(_da2['pct_M'], _da2['pct_P'])]

            _agree_mp     = sum(1 for c in _col_mp if c == '#16a34a')
            _total_mp     = len(_col_mp)
            _pct_agree_mp = _agree_mp / _total_mp * 100 if _total_mp else 0

            _sc1, _sc2, _sc3 = st.columns(3)

            # ── Loop untuk scatter plot kolom 1 & 2 (x = Wood) ──────────
            for _col, _var_name, _y_pct, _colors_plot, _y_actual, _y_unit in [
                (_sc1, 'Metal Usage',       _da2['pct_M'], _col_m, _da2['MetalUsageKG'],    'KG'),
                (_sc2, 'Production Output', _da2['pct_P'], _col_p, _da2['ProductionUnit'],  'unit'),
            ]:
                with _col:
                    _xmax = max(abs(_da2['pct_W']).max(), 5) * 1.25
                    _ymax = max(abs(_y_pct).max(), 5) * 1.25

                    _fig_sc = go.Figure()

                    for _qx, _qy, _qc in [
                        ([0,  _xmax], [0,  _ymax], 'rgba(220,252,231,0.55)'),
                        ([-_xmax, 0], [-_ymax, 0], 'rgba(220,252,231,0.55)'),
                        ([-_xmax, 0], [0,  _ymax], 'rgba(254,226,226,0.55)'),
                        ([0,  _xmax], [-_ymax, 0], 'rgba(254,226,226,0.55)'),
                    ]:
                        _fig_sc.add_shape(type='rect',
                            x0=_qx[0], x1=_qx[1], y0=_qy[0], y1=_qy[1],
                            fillcolor=_qc, line_width=0, layer='below')

                    _fig_sc.add_hline(y=0, line_color='#94a3b8', line_width=1.5)
                    _fig_sc.add_vline(x=0, line_color='#94a3b8', line_width=1.5)

                    _vn_short = 'Metal' if 'Metal' in _var_name else 'Produksi'
                    for _ax, _ay, _txt, _tc in [
                        ( _xmax*0.55,  _ymax*0.75, f'Same Direction\nWood ↑  {_vn_short} ↑', '#166534'),
                        (-_xmax*0.55, -_ymax*0.75, f'Same Direction\nWood ↓  {_vn_short} ↓', '#166534'),
                        (-_xmax*0.55,  _ymax*0.75, f'Opposite\nWood ↓  {_vn_short} ↑', '#991b1b'),
                        ( _xmax*0.55, -_ymax*0.75, f'Opposite\nWood ↑  {_vn_short} ↓', '#991b1b'),
                    ]:
                        _fig_sc.add_annotation(x=_ax, y=_ay, text=_txt,
                            showarrow=False, font=dict(size=10, color=_tc),
                            bgcolor='rgba(255,255,255,0.78)', borderpad=4,
                            align='center')

                    _pct_col = 'pct_M' if 'Metal' in _var_name else 'pct_P'
                    _hover_texts = []
                    for i, row in _da2.iterrows():
                        _orig_idx  = _wrp_clean.index[_wrp_clean['DisplayLabel'] == row['DisplayLabel']]
                        _orig_i    = _orig_idx[0] if len(_orig_idx) else None
                        _w_prev    = _wrp_clean.loc[_orig_i - 1, 'WoodLines']       if (_orig_i and _orig_i > 0) else None
                        _y_prev    = _wrp_clean.loc[_orig_i - 1, _y_actual.name]    if (_orig_i and _orig_i > 0) else None

                        _pct_w_val = row['pct_W']
                        _pct_y_val = row[_pct_col]
                        _dir_w     = '▲ naik' if _pct_w_val >= 0 else '▼ turun'
                        _dir_y     = '▲ naik' if _pct_y_val >= 0 else '▼ turun'
                        _status    = 'Searah' if np.sign(_pct_w_val)==np.sign(_pct_y_val) else 'Berlawanan'

                        _w_prev_str = f'bln lalu: {int(_w_prev):,} shipments' if _w_prev is not None else ''
                        _y_prev_str = f'bln lalu: {_y_prev:,.0f} {_y_unit}'   if _y_prev is not None else ''

                        _hover_texts.append(
                            f"<b>{row['DisplayLabel']}</b><br>"
                            f"Wood: <b>{int(row['WoodLines']):,} shipments</b>  "
                            f"{_dir_w} <b>{abs(_pct_w_val):.1f}%</b>  <i>({_w_prev_str})</i><br>"
                            f"{'Metal' if 'Metal' in _var_name else 'Production'} ({_vn_short}): "
                            f"<b>{row[_y_actual.name]:,.0f} {_y_unit}</b>  "
                            f"{_dir_y} <b>{abs(_pct_y_val):.1f}%</b>  <i>({_y_prev_str})</i><br>"
                            f"{_status}"
                        )

                    _fig_sc.add_trace(go.Scatter(
                        x=list(_da2['pct_W']),
                        y=list(_y_pct),
                        mode='markers+text',
                        marker=dict(color=_colors_plot, size=12,
                                    line=dict(color='white', width=1.5)),
                        text=_da2['DisplayLabel'].str.replace(' ', '<br>'),
                        textposition='top center',
                        textfont=dict(size=7.5, color='#475569'),
                        hovertemplate='%{customdata}<extra></extra>',
                        customdata=_hover_texts
                    ))

                    _agree_n = sum(1 for c in _colors_plot if c == '#16a34a')
                    _total_n = len(_colors_plot)
                    _pct_agr = _agree_n / _total_n * 100 if _total_n else 0

                    _fig_sc.update_layout(
                        title=dict(
                            text=f'Wood vs {_var_name}  —  '
                                 f'<span style="color:#16a34a;font-weight:700;">{_agree_n} bulan searah</span>'
                                 f' / <span style="color:#dc2626;">{_total_n-_agree_n} bulan berlawanan</span>'
                                 f'  <span style="color:#64748b;font-size:11px;">({_pct_agr:.0f}% searah)</span>',
                            font=dict(size=12.5), x=0.5, xanchor='center'),
                        paper_bgcolor='white', plot_bgcolor='white', height=400,
                        margin=dict(l=20, r=20, t=65, b=55),
                        xaxis=dict(showgrid=False, zeroline=False,
                                   title='Wood packaging — change from previous month (%)',
                                   range=[-_xmax, _xmax], ticksuffix='%'),
                        yaxis=dict(showgrid=False, zeroline=False,
                                   title=f'{_var_name} — change from previous month (%)',
                                   range=[-_ymax, _ymax], ticksuffix='%'),
                        showlegend=False)
                    st.plotly_chart(_fig_sc, use_container_width=True,
                                    config={'displayModeBar': False})

            # ── Kolom 3: Metal Usage vs Production Unit — scatter plot ────
            with _sc3:
                _xmax3 = max(abs(_da2['pct_M']).max(), 5) * 1.25
                _ymax3 = max(abs(_da2['pct_P']).max(), 5) * 1.25

                _fig_mp = go.Figure()

                for _qx, _qy, _qc in [
                    ([0,  _xmax3], [0,  _ymax3], 'rgba(220,252,231,0.55)'),
                    ([-_xmax3, 0], [-_ymax3, 0], 'rgba(220,252,231,0.55)'),
                    ([-_xmax3, 0], [0,  _ymax3], 'rgba(254,226,226,0.55)'),
                    ([0,  _xmax3], [-_ymax3, 0], 'rgba(254,226,226,0.55)'),
                ]:
                    _fig_mp.add_shape(type='rect',
                        x0=_qx[0], x1=_qx[1], y0=_qy[0], y1=_qy[1],
                        fillcolor=_qc, line_width=0, layer='below')

                _fig_mp.add_hline(y=0, line_color='#94a3b8', line_width=1.5)
                _fig_mp.add_vline(x=0, line_color='#94a3b8', line_width=1.5)

                for _ax, _ay, _txt, _tc in [
                    ( _xmax3*0.55,  _ymax3*0.75, 'Same Direction\nMetal ↑  Produksi ↑', '#166534'),
                    (-_xmax3*0.55, -_ymax3*0.75, 'Same Direction\nMetal ↓  Produksi ↓', '#166534'),
                    (-_xmax3*0.55,  _ymax3*0.75, 'Opposite\nMetal ↓  Produksi ↑',       '#991b1b'),
                    ( _xmax3*0.55, -_ymax3*0.75, 'Opposite\nMetal ↑  Produksi ↓',       '#991b1b'),
                ]:
                    _fig_mp.add_annotation(x=_ax, y=_ay, text=_txt,
                        showarrow=False, font=dict(size=10, color=_tc),
                        bgcolor='rgba(255,255,255,0.78)', borderpad=4,
                        align='center')

                _hover_mp = []
                for i, row in _da2.iterrows():
                    _orig_idx = _wrp_clean.index[_wrp_clean['DisplayLabel'] == row['DisplayLabel']]
                    _orig_i   = _orig_idx[0] if len(_orig_idx) else None
                    _m_prev   = _wrp_clean.loc[_orig_i - 1, 'MetalUsageKG']    if (_orig_i and _orig_i > 0) else None
                    _p_prev   = _wrp_clean.loc[_orig_i - 1, 'ProductionUnit']  if (_orig_i and _orig_i > 0) else None

                    _pct_m_val = row['pct_M']
                    _pct_p_val = row['pct_P']
                    _dir_m2    = '▲ naik' if _pct_m_val >= 0 else '▼ turun'
                    _dir_p2    = '▲ naik' if _pct_p_val >= 0 else '▼ turun'
                    _status2   = 'Searah' if np.sign(_pct_m_val)==np.sign(_pct_p_val) else 'Berlawanan'

                    _m_prev_str = f'bln lalu: {_m_prev:,.0f} KG'   if _m_prev is not None else ''
                    _p_prev_str = f'bln lalu: {int(_p_prev):,} unit' if _p_prev is not None else ''

                    _hover_mp.append(
                        f"<b>{row['DisplayLabel']}</b><br>"
                        f"Metal: <b>{row['MetalUsageKG']:,.0f} KG</b>  "
                        f"{_dir_m2} <b>{abs(_pct_m_val):.1f}%</b>  <i>({_m_prev_str})</i><br>"
                        f"Production: <b>{int(row['ProductionUnit']):,} unit</b>  "
                        f"{_dir_p2} <b>{abs(_pct_p_val):.1f}%</b>  <i>({_p_prev_str})</i><br>"
                        f"{_status2}"
                    )

                _fig_mp.add_trace(go.Scatter(
                    x=list(_da2['pct_M']),
                    y=list(_da2['pct_P']),
                    mode='markers+text',
                    marker=dict(color=_col_mp, size=12,
                                line=dict(color='white', width=1.5)),
                    text=_da2['DisplayLabel'].str.replace(' ', '<br>'),
                    textposition='top center',
                    textfont=dict(size=7.5, color='#475569'),
                    hovertemplate='%{customdata}<extra></extra>',
                    customdata=_hover_mp
                ))

                _fig_mp.update_layout(
                    title=dict(
                        text=f'Metal Usage vs Production Unit  —  '
                             f'<span style="color:#16a34a;font-weight:700;">{_agree_mp} bulan searah</span>'
                             f' / <span style="color:#dc2626;">{_total_mp - _agree_mp} bulan berlawanan</span>'
                             f'  <span style="color:#64748b;font-size:11px;">({_pct_agree_mp:.0f}% searah)</span>',
                        font=dict(size=12.5), x=0.5, xanchor='center'),
                    paper_bgcolor='white', plot_bgcolor='white', height=400,
                    margin=dict(l=20, r=20, t=65, b=55),
                    xaxis=dict(showgrid=False, zeroline=False,
                               title='Metal Usage — change from previous month (%)',
                               range=[-_xmax3, _xmax3], ticksuffix='%'),
                    yaxis=dict(showgrid=False, zeroline=False,
                               title='Production Unit — change from previous month (%)',
                               range=[-_ymax3, _ymax3], ticksuffix='%'),
                    showlegend=False)
                st.plotly_chart(_fig_mp, use_container_width=True,
                                config={'displayModeBar': False})

            # ── Summary insight ───────────────────────────────────────
            _verdict_m = (
                "✅ Wood Shipments <b>cenderung mengikuti</b> pergerakan Metal Usage."
                if _pct_agree_m >= 60 else
                "⚠️ Wood Shipments <b>tidak konsisten mengikuti</b> Metal Usage — faktor lain lebih dominan."
                if _pct_agree_m >= 45 else
                "🚨 Wood Shipments <b>bergerak relatif independen</b> dari Metal Usage — banyak bulan berlawanan arah."
            )
            _verdict_p = (
                "✅ Wood Shipments <b>cenderung mengikuti</b> pergerakan Production Output."
                if _pct_agree_p >= 60 else
                "⚠️ Wood Shipments <b>tidak konsisten mengikuti</b> Production Output."
                if _pct_agree_p >= 45 else
                "🚨 Wood Shipments <b>bergerak relatif independen</b> dari output Produksi."
            )
            _verdict_mp = (
                "✅ Metal Usage <b>cenderung bergerak searah</b> dengan Production Unit."
                if _pct_agree_mp >= 60 else
                "⚠️ Metal Usage <b>tidak selalu searah</b> dengan Production Unit — ada faktor lain."
                if _pct_agree_mp >= 45 else
                "🚨 Metal Usage <b>bergerak relatif independen</b> dari Production Unit."
            )

            _si1, _si2, _si3 = st.columns(3)
            with _si1:
                _bc  = '#f0fdf4' if _pct_agree_m  >= 60 else ('#fffbeb' if _pct_agree_m  >= 45 else '#fef2f2')
                _bl  = '#16a34a' if _pct_agree_m  >= 60 else ('#d97706' if _pct_agree_m  >= 45 else '#dc2626')
                st.markdown(
                    f'<div style="background:{_bc};border-left:4px solid {_bl};'
                    f'padding:12px 16px;border-radius:8px;font-size:13px;line-height:1.7;">'
                    f'<b>Wood ↔ Metal Usage:</b><br>'
                    f'{_badge_m} Searah dalam <b>{_agree_m}/{_total_m} bulan ({_pct_agree_m:.0f}%)</b><br>'
                    f'{_verdict_m}</div>',
                    unsafe_allow_html=True)
            with _si2:
                _bc2 = '#f0fdf4' if _pct_agree_p  >= 60 else ('#fffbeb' if _pct_agree_p  >= 45 else '#fef2f2')
                _bl2 = '#16a34a' if _pct_agree_p  >= 60 else ('#d97706' if _pct_agree_p  >= 45 else '#dc2626')
                st.markdown(
                    f'<div style="background:{_bc2};border-left:4px solid {_bl2};'
                    f'padding:12px 16px;border-radius:8px;font-size:13px;line-height:1.7;">'
                    f'<b>Wood ↔ Production Output:</b><br>'
                    f'{_badge_p} Searah dalam <b>{_agree_p}/{_total_m} bulan ({_pct_agree_p:.0f}%)</b><br>'
                    f'{_verdict_p}</div>',
                    unsafe_allow_html=True)
            with _si3:
                _bc3 = '#f0fdf4' if _pct_agree_mp >= 60 else ('#fffbeb' if _pct_agree_mp >= 45 else '#fef2f2')
                _bl3 = '#16a34a' if _pct_agree_mp >= 60 else ('#d97706' if _pct_agree_mp >= 45 else '#dc2626')
                _badge_mp2 = '🟢' if _pct_agree_mp >= 70 else ('🟡' if _pct_agree_mp >= 45 else '🔴')
                st.markdown(
                    f'<div style="background:{_bc3};border-left:4px solid {_bl3};'
                    f'padding:12px 16px;border-radius:8px;font-size:13px;line-height:1.7;">'
                    f'<b>Metal Usage ↔ Production Unit:</b><br>'
                    f'{_badge_mp2} Searah dalam <b>{_agree_mp}/{_total_mp} bulan ({_pct_agree_mp:.0f}%)</b><br>'
                    f'{_verdict_mp}</div>',
                    unsafe_allow_html=True)

        else:
            st.info("Data tidak cukup untuk analisis arah pergerakan.")

    st.markdown("<br>", unsafe_allow_html=True)

    # ─────────────────────────────────────────────────────────────────────────
    # BAGIAN 3 — TEMUAN DIMENSI & BERAT
    # ─────────────────────────────────────────────────────────────────────────
    with st.expander("\U0001f4d0 Temuan Utama: Dimensi & Berat \u2192 Pilihan Kemasan", expanded=True):

        st.markdown("**Profil Kemasan per Kategori Ukuran**")
        dim_profile_rows = []
        for cat in DIM_LABELS:
            sub_cat = df_all[df_all['Dim_Category'] == cat]
            if sub_cat.empty:
                continue
            n_total  = len(sub_cat)
            w_n      = int((sub_cat['Material_Clean'] == 'W').sum())
            cb_n     = int((sub_cat['Material_Clean'] == 'B').sum())
            cbd_n    = int((sub_cat['Material_Clean'] == 'Cb').sum())
            ct_n     = int((sub_cat['Material_Clean'] == 'Ct').sum())
            w_pct    = w_n   / n_total * 100 if n_total else 0
            cb_pct   = cb_n  / n_total * 100 if n_total else 0
            cbd_pct  = cbd_n / n_total * 100 if n_total else 0
            ct_pct   = ct_n  / n_total * 100 if n_total else 0
            avg_wt   = sub_cat[sub_cat['Net Weight'] > 0]['Net Weight'].mean()
            avg_vol  = sub_cat[sub_cat['Volume_m3'] > 0]['Volume_m3'].mean()
            if w_pct > 90:
                status = "\U0001f534 Wajib Kayu"
            elif w_pct >= 40:
                status = "\U0001f7e1 Perlu Dikaji"
            else:
                status = "\U0001f7e2 Potensial Diganti"
            dim_profile_rows.append({
                'Kategori':             cat,
                'Jumlah Items':         f"{n_total:,}",
                'Wood':                 f"{w_n:,} ({w_pct:.1f}%)",
                'Carton Box':           f"{cb_n:,} ({cb_pct:.1f}%)",
                'Cardboard':            f"{cbd_n:,} ({cbd_pct:.1f}%)",
                'Carton':               f"{ct_n:,} ({ct_pct:.1f}%)",
                'Rata-rata Berat (kg)': f"{avg_wt:.1f}" if not pd.isna(avg_wt) else "\u2014",
                'Rata-rata Volume (m³)':f"{avg_vol:.4f}" if not pd.isna(avg_vol) else "\u2014",
                'Status':               status,
            })
        st.dataframe(pd.DataFrame(dim_profile_rows), use_container_width=True, hide_index=True)

        st.markdown(
            f'<div class="insight-box">\u26a0\ufe0f <b>Temuan Kritis:</b> Item kategori '
            f'<b>M (2\u20134m), L (4\u20137m), dan XL (&gt;7m)</b> menggunakan Wood hampir 100% '
            f'karena alasan struktural. Peluang penggantian <b>terfokus pada XS dan S</b> \u2014 '
            f'XS: {xs_pct_val:.0f}% Wood, S: {s_pct_val:.0f}% Wood pada periode {ov_label}.</div>',
            unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown("**Panduan Pemilihan Kemasan \u2014 Ketetapan Berbasis Data**")
        st.markdown("""
        <div style="background:white;border-radius:12px;padding:16px 20px;border:1px solid #e5e7eb;">
          <table style="width:100%;border-collapse:collapse;font-size:13px;">
            <thead>
              <tr style="background:#f0f6ff;">
                <th style="padding:8px 12px;text-align:left;color:#1e3a5f;font-weight:600;
                border-bottom:2px solid #dbeafe;">Ukuran &amp; Berat Item</th>
                <th style="padding:8px 12px;text-align:left;color:#1e3a5f;font-weight:600;
                border-bottom:2px solid #dbeafe;">Kemasan yang Tepat</th>
                <th style="padding:8px 12px;text-align:left;color:#1e3a5f;font-weight:600;
                border-bottom:2px solid #dbeafe;">Alasan</th>
              </tr>
            </thead>
            <tbody>
              <tr>
                <td style="padding:9px 12px;border-bottom:1px solid #f1f5f9;">
                  Panjang/Lebar/Tinggi &gt; 4 m <b>atau</b> Berat &gt; 200 kg</td>
                <td style="padding:9px 12px;border-bottom:1px solid #f1f5f9;">
                  <span style="background:#fef2f2;color:#dc2626;padding:3px 10px;
                  border-radius:6px;font-weight:700;">&#128308; Wajib Kayu</span></td>
                <td style="padding:9px 12px;border-bottom:1px solid #f1f5f9;color:#64748b;">
                  Item besar/berat butuh palet kayu untuk kekuatan struktural</td>
              </tr>
              <tr style="background:#fafafa;">
                <td style="padding:9px 12px;border-bottom:1px solid #f1f5f9;">
                  Ukuran 1,5&ndash;4 m <b>dan</b> Berat 50&ndash;200 kg</td>
                <td style="padding:9px 12px;border-bottom:1px solid #f1f5f9;">
                  <span style="background:#fffbeb;color:#d97706;padding:3px 10px;
                  border-radius:6px;font-weight:700;">&#128993; Perlu Dikaji</span></td>
                <td style="padding:9px 12px;border-bottom:1px solid #f1f5f9;color:#64748b;">
                  Tergantung bentuk item &mdash; harus dievaluasi satu per satu</td>
              </tr>
              <tr>
                <td style="padding:9px 12px;">
                  Semua sisi &lt; 1,5 m <b>dan</b> Berat &lt; 50 kg</td>
                <td style="padding:9px 12px;">
                  <span style="background:#eff6ff;color:#2563eb;padding:3px 10px;
                  border-radius:6px;font-weight:700;">&#128994; Ganti Karton</span></td>
                <td style="padding:9px 12px;color:#64748b;">
                  Item kecil &amp; ringan cukup dilindungi karton &mdash; lebih hemat &amp; ramah lingkungan</td>
              </tr>
            </tbody>
          </table>
        </div>""", unsafe_allow_html=True)

    # ─────────────────────────────────────────────────────────────────────────
    # BAGIAN 4 — PELUANG PENGGANTIAN / OVERPACKAGED
    # ─────────────────────────────────────────────────────────────────────────
    with st.expander("\u267b\ufe0f Peluang Konversi: Wood \u2192 Material Alternatif", expanded=True):

        o1, o2, o3 = st.columns(3)
        o1.metric("Item Wood 'Overpackaged'",
                  f"{overpack_n_all:,}",
                  "Dimensi <1.5m & berat <50kg \u2014 tapi pakai kayu")
        o2.metric("% dari Total Wood",
                  f"{overpack_pct_all:.1f}%",
                  "Quick win langsung untuk Pledge 6")
        o3.metric("Target Konversi",
                  f"{overpack_n_all:,} item",
                  f"\u2192 hemat ~{overpack_pct_all:.1f}% Wood dari total pengiriman")

        if not overpack_items.empty:
            top10_over = overpack_items.head(10)
            fig_over_ov = go.Figure(go.Bar(
                x=top10_over['Lines'],
                y=top10_over['Item'],
                orientation='h',
                marker=dict(color=top10_over['Lines'],
                            colorscale=[[0, '#FEF2F2'], [1, '#dc2626']], showscale=False),
                text=top10_over['Lines'].apply(lambda v: f"{v:,} lines"),
                textposition='outside',
                hovertemplate='<b>%{y}</b><br>%{x:,} shipment lines overpackaged<extra></extra>'
            ))
            fig_over_ov.update_layout(
                title=f'Top 10 Item Overpackaged \u2014 {ov_label}',
                paper_bgcolor='white', plot_bgcolor='white', height=340,
                margin=dict(l=10, r=60, t=50, b=10),
                xaxis=dict(showgrid=True, gridcolor='#f1f5f9', title='Shipments'),
                yaxis=dict(showgrid=False, autorange='reversed', tickfont=dict(size=10.5)))
            st.plotly_chart(fig_over_ov, use_container_width=True, config={'displayModeBar': False})
        else:
            st.info(f"Tidak ada item overpackaged pada periode {ov_label}.")

        st.markdown(
            f'<div class="critical-box">\U0001f6a8 <b>Quick Win Terbesar ({ov_label}):</b> Terdapat '
            f'<b>{overpack_n_all:,} item ({overpack_pct_all:.1f}% dari total Wood)</b> '
            f'yang berukuran kecil (&lt;1.500 mm) dan ringan (&lt;50 kg) namun masih pakai kayu. '
            f'Mengganti ke Carton Box adalah langkah tercepat untuk Pledge 6.</div>',
            unsafe_allow_html=True)

        with st.expander(f"\U0001f4cb Lihat daftar lengkap {len(overpack_items):,} item prioritas konversi"):
            st.dataframe(
                overpack_items.rename(columns={'Item': 'Item Name', 'Lines': 'Shipments'}),
                use_container_width=True, hide_index=True)

    # ─────────────────────────────────────────────────────────────────────────
    # BAGIAN 5 — POLA CUSTOMER
    # ─────────────────────────────────────────────────────────────────────────
    with st.expander("\U0001f465 Pola Customer & Penggunaan Wood", expanded=True):

        if cust_wood_pct.empty:
            st.info(f"Tidak ada data customer pada periode {ov_label}.")
        else:
            fig_cust_ov = go.Figure(go.Bar(
                x=cust_wood_pct['Name'],
                y=cust_wood_pct['Pct_Wood'],
                marker=dict(color=cust_wood_pct['Pct_Wood'],
                            colorscale=[[0, '#FEF9F5'], [1, '#7C4A19']], showscale=False),
                text=cust_wood_pct['Pct_Wood'].apply(lambda v: f'{v:.1f}%'),
                textposition='outside',
                hovertemplate='<b>%{x}</b><br>%{y:.1f}% pengiriman menggunakan Wood<extra></extra>'
            ))
            fig_cust_ov.update_layout(
                title=f'% Penggunaan Wood per Customer (Top 10) \u2014 {ov_label}',
                paper_bgcolor='white', plot_bgcolor='white', height=340,
                margin=dict(l=20, r=20, t=50, b=80),
                xaxis=dict(showgrid=False, tickangle=-30, tickfont=dict(size=9.5), title='Customer'),
                yaxis=dict(showgrid=True, gridcolor='#f1f5f9', title='% Wood', range=[0, 115]))
            st.plotly_chart(fig_cust_ov, use_container_width=True, config={'displayModeBar': False})

            cust_tbl = cust_wood_pct[['Name', 'Lines', 'Pct_Wood']].copy()
            def _cust_cat(p):
                if p > 90:    return "\U0001f534 Prioritas Kolaborasi"
                elif p >= 50: return "\U0001f7e1 Perlu Review"
                else:         return "\U0001f7e2 Relatif Baik"
            cust_tbl['Kategori']     = cust_tbl['Pct_Wood'].apply(_cust_cat)
            cust_tbl['% Pakai Kayu'] = cust_tbl['Pct_Wood'].apply(lambda v: f"{v:.1f}%")
            cust_tbl = cust_tbl.rename(columns={'Name': 'Pelanggan', 'Lines': 'Total Pengiriman'})
            st.dataframe(cust_tbl[['Pelanggan', 'Total Pengiriman', '% Pakai Kayu', 'Kategori']],
                         use_container_width=True, hide_index=True)

            top_cust_pct = float(cust_wood_pct['Pct_Wood'].iloc[0])
            st.markdown(
                f'<div class="insight-box">\U0001f465 <b>Insight ({ov_label}):</b> Customer '
                f'<b>{top_cust_name}</b> adalah pengguna Wood terbesar di top 10 '
                f'({top_cust_pct:.1f}% pengirimannya menggunakan kayu). '
                f'Customer dengan % Wood tinggi adalah kandidat utama kolaborasi Pledge 6.</div>',
                unsafe_allow_html=True)

    # ─────────────────────────────────────────────────────────────────────────
    # BAGIAN 6 — PREDIKSI TREN
    # ─────────────────────────────────────────────────────────────────────────
    with st.expander("\U0001f52e Proyeksi Penggunaan Wood (12 Bulan ke Depan \u2014 XGBoost)", expanded=True):

        if ov_view == "Monthly":
            st.info("Proyeksi prediksi tidak tersedia untuk tampilan Monthly. "
                    "Pilih 'All' atau 'Yearly' untuk melihat proyeksi.")
        else:
            monthly_fc, fc_df, fc_pct, _fc_model = forecast_xgboost('Wood', periods=12)

            if monthly_fc is None or fc_df is None:
                st.info("Data historis belum cukup untuk prediksi (minimal 8 bulan).")
            else:
                st.session_state['fc_wood_pct'] = fc_pct

                hist_labels = [fmt_month_str(ym) for ym in monthly_fc['YearMonth']]
                pred_labels = [fmt_month_str(ym) for ym in fc_df['YearMonth']]

                # Historis: selalu tampilkan SEMUA data (agar grafik tidak kosong)
                # Saat Yearly dipilih, highlight saja tapi tetap tunjukkan full history
                hist_plot        = monthly_fc  # selalu full
                hist_labels_plot = hist_labels  # semua label


                # ── Sambungkan historis → prediksi seperti gambar 2 ──────────
                last_hist_ym_ov    = monthly_fc['YearMonth'].iloc[-1]
                last_hist_count_ov = float(monthly_fc['Count'].iloc[-1])
                # Bridge: titik terakhir historis + semua prediksi
                bridge_x_ov = [fmt_month_str(last_hist_ym_ov)] + pred_labels
                bridge_y_ov = [last_hist_count_ov] + fc_df['Forecast'].tolist()

                fig_fc_ov = go.Figure()
                # Garis historis — solid abu-abu/cokelat tua
                fig_fc_ov.add_trace(go.Scatter(
                    x=hist_labels_plot, y=hist_plot['Count'],
                    name='Historis (Aktual)', mode='lines+markers',
                    line=dict(color='#7C4A19', width=2.5),
                    marker=dict(size=4, color='#7C4A19'),
                    hovertemplate='<b>%{x}</b><br>Aktual: %{y:,}<extra></extra>'
                ))
                # Garis prediksi — disambung dari titik terakhir, warna amber, putus-putus
                fig_fc_ov.add_trace(go.Scatter(
                    x=bridge_x_ov, y=bridge_y_ov,
                    name='Prediksi 12 Bulan', mode='lines+markers',
                    line=dict(color='#f59e0b', width=2.5, dash='dot'),
                    marker=dict(size=6, color='#f59e0b', symbol='diamond'),
                    hovertemplate='<b>%{x}</b><br>Prediksi: %{y:,.0f}<extra></extra>'
                ))
                avg_hist_fc = float(monthly_fc['Count'].mean())
                fig_fc_ov.add_hline(
                    y=avg_hist_fc, line_dash='dash', line_color='#94a3b8',
                    annotation_text=f'Avg historis {avg_hist_fc:.0f}',
                    annotation_font_size=9, annotation_font_color='#94a3b8'
                )
                # Separator garis vertikal pakai shapes (categorical x-axis)
                all_x_ov   = hist_labels_plot + bridge_x_ov[1:]
                split_lbl  = fmt_month_str(last_hist_ym_ov)
                split_pos_ov = all_x_ov.index(split_lbl) if split_lbl in all_x_ov else len(hist_labels_plot) - 1
                fig_fc_ov.update_layout(
                    title='Prediksi Penggunaan Wood — 12 Bulan ke Depan (XGBoost)',
                    paper_bgcolor='white', plot_bgcolor='white', height=320,
                    margin=dict(l=20, r=20, t=50, b=60),
                    xaxis=dict(showgrid=False, tickangle=-45, tickfont=dict(size=8), title='Month'),
                    yaxis=dict(showgrid=True, gridcolor='#f1f5f9', title='Wood Shipments'),
                    legend=dict(orientation='h', y=1.12, x=1, xanchor='right'),
                    hovermode='x unified',
                    shapes=[dict(
                        type='line', xref='x', yref='paper',
                        x0=split_pos_ov, x1=split_pos_ov, y0=0, y1=1,
                        line=dict(color='#64748b', width=1.2, dash='dot')
                    )],
                    annotations=[dict(
                        xref='x', yref='paper',
                        x=split_pos_ov + 0.3, y=0.97,
                        text='← Historis  |  Prediksi →',
                        showarrow=False,
                        font=dict(size=8, color='#64748b'),
                        xanchor='left',
                        bgcolor='rgba(255,255,255,0.85)',
                        bordercolor='#e2e8f0', borderwidth=1
                    )]
                )
                st.plotly_chart(fig_fc_ov, use_container_width=True, config={'displayModeBar': False})

                # ── Tabel: 3 bulan historis + 12 bulan prediksi ──────────────
                st.markdown(
                    '<div class="axis-note">'
                    '📋 Tabel berikut menampilkan <b>3 bulan historis terakhir</b> '
                    '(konteks aktual) diikuti <b>12 bulan prediksi ke depan</b>.'
                    '</div>', unsafe_allow_html=True
                )
                last3_hist = monthly_fc.tail(3)
                tbl_rows = []
                for _, hrow in last3_hist.iterrows():
                    tbl_rows.append({
                        'Status': '📋 Historis',
                        'Bulan': fmt_month_str(hrow['YearMonth']),
                        'Lines (Aktual)': f"{int(hrow['Count']):,}",
                        'vs Rata-rata Historis': '',
                    })
                for _, frow in fc_df.iterrows():
                    vs     = frow['Forecast'] - avg_hist_fc
                    vs_str = f"+{vs:,.0f}" if vs >= 0 else f"{vs:,.0f}"
                    tbl_rows.append({
                        'Status': '🔮 Prediksi',
                        'Bulan': fmt_month_str(frow['YearMonth']),
                        'Lines (Aktual)': f"{frow['Forecast']:,.0f}",
                        'vs Rata-rata Historis': vs_str,
                    })
                st.dataframe(
                    pd.DataFrame(tbl_rows),
                    use_container_width=True, hide_index=True, height=500
                )


                if fc_pct > 5:
                    fc_cls, fc_icon = "critical-box", "\u26a0\ufe0f"
                    fc_msg = (f"Prediksi menunjukkan <b>peningkatan Wood {fc_pct:+.1f}%</b>. "
                              f"Intervensi segera diperlukan. "
                              f"Prioritaskan konversi {overpack_n_all:,} item overpackaged.")
                elif fc_pct < -5:
                    fc_cls, fc_icon = "tooltip-box", "\u2705"
                    fc_msg = (f"Prediksi menunjukkan <b>pengurangan Wood {abs(fc_pct):.1f}%</b>. "
                              f"Pertahankan momentum dan percepat program konversi.")
                else:
                    fc_cls, fc_icon = "insight-box", "\U0001f4ca"
                    fc_msg = (f"Wood diprediksi <b>relatif stabil ({fc_pct:+.1f}%)</b>. "
                              f"Diperlukan inisiatif aktif untuk mendorong perubahan ke Pledge 6.")
                st.markdown(
                    f'<div class="{fc_cls}">{fc_icon} <b>Interpretasi:</b> {fc_msg}</div>',
                    unsafe_allow_html=True)

    # ─────────────────────────────────────────────────────────────────────────
    # BAGIAN 7 — ROADMAP AKSI
    # ─────────────────────────────────────────────────────────────────────────
    st.markdown("<br>", unsafe_allow_html=True)
    st.markdown("### \U0001f5fa\ufe0f Roadmap Aksi menuju Pledge 6")

    a1, a2, a3 = st.columns(3)

    with a1:
        st.markdown(f"""
        <div style="background:white;border-radius:14px;padding:20px;
        box-shadow:0 2px 12px rgba(0,0,0,.08);border-top:5px solid #16a34a;">
          <div style="font-size:12px;font-weight:700;color:#16a34a;text-transform:uppercase;
          letter-spacing:.6px;margin-bottom:8px;">&#9889; SEGERA &mdash; 0 sampai 3 Bulan</div>
          <div style="font-size:20px;font-weight:800;color:#16a34a;margin-bottom:10px;">Quick Win</div>
          <ul style="font-size:13px;color:#374151;padding-left:16px;margin:0 0 14px;line-height:1.9;">
            <li>Konversi <b>{overpack_n_all:,} item</b> dari Wood ke Carton Box</li>
            <li>Fokus: dimensi &lt;1.500 mm, berat &lt;50 kg</li>
            <li>Sosialisasi panduan kemasan ke tim operasional</li>
          </ul>
          <div style="background:#f0fdf4;border-radius:8px;padding:10px 12px;
          font-size:12px;color:#166534;">
            &#127919; <b>Target:</b> Kurangi <b>{overpack_n_all:,} item</b> Wood kecil
            &rarr; hemat <b>~{overpack_pct_all:.0f}%</b> dari total Wood
          </div>
        </div>""", unsafe_allow_html=True)
        with st.expander(f"Lihat {min(15, len(overpack_items))} item prioritas konversi"):
            st.dataframe(
                overpack_items.head(15).rename(columns={'Item': 'Item Name', 'Lines': 'Shipments'}),
                use_container_width=True, hide_index=True)

    with a2:
        st.markdown(f"""
        <div style="background:white;border-radius:14px;padding:20px;
        box-shadow:0 2px 12px rgba(0,0,0,.08);border-top:5px solid #2563eb;">
          <div style="font-size:12px;font-weight:700;color:#2563eb;text-transform:uppercase;
          letter-spacing:.6px;margin-bottom:8px;">&#128204; JANGKA MENENGAH &mdash; 3 sampai 12 Bulan</div>
          <div style="font-size:20px;font-weight:800;color:#2563eb;margin-bottom:10px;">Strategic</div>
          <ul style="font-size:13px;color:#374151;padding-left:16px;margin:0 0 14px;line-height:1.9;">
            <li>Kolaborasi dengan <b>{top_cust_name}</b> &amp; customer besar lain</li>
            <li>Review item ukuran S/M (1,5&ndash;4 m) per kasus</li>
            <li>Pilot program kemasan alternatif untuk &ldquo;gray area&rdquo;</li>
          </ul>
          <div style="background:#eff6ff;border-radius:8px;padding:10px 12px;
          font-size:12px;color:#1e3a8a;">
            &#127919; <b>Target:</b> Review <b>{n_sm:,} item</b> kategori S &amp; M
            yang masih menggunakan Wood
          </div>
        </div>""", unsafe_allow_html=True)
        cust_disp = cust_wood_pct[['Name', 'Lines', 'Pct_Wood']].copy()
        cust_disp.columns = ['Pelanggan', 'Total Pengiriman', '% Pakai Kayu']
        cust_disp['% Pakai Kayu'] = cust_disp['% Pakai Kayu'].apply(lambda v: f"{v:.1f}%")
        with st.expander("Lihat daftar customer & tingkat penggunaan kayu"):
            st.dataframe(cust_disp, use_container_width=True, hide_index=True)

    with a3:
        st.markdown(f"""
        <div style="background:white;border-radius:14px;padding:20px;
        box-shadow:0 2px 12px rgba(0,0,0,.08);border-top:5px solid #7C4A19;">
          <div style="font-size:12px;font-weight:700;color:#7C4A19;text-transform:uppercase;
          letter-spacing:.6px;margin-bottom:8px;">&#127807; JANGKA PANJANG &mdash; 12+ Bulan</div>
          <div style="font-size:20px;font-weight:800;color:#7C4A19;margin-bottom:10px;">Transformational</div>
          <ul style="font-size:13px;color:#374151;padding-left:16px;margin:0 0 14px;line-height:1.9;">
            <li>Desain ulang kemasan untuk item ukuran M/L</li>
            <li>Sertifikasi kemasan berkelanjutan untuk Pledge 6</li>
            <li>Eksplorasi material inovatif: recycled plastic foil, reusable packaging</li>
          </ul>
          <div style="background:#fdf6f0;border-radius:8px;padding:10px 12px;
          font-size:12px;color:#7C4A19;">
            &#127919; <b>Target:</b> Kurangi Wood dari <b>{wood_pct_all:.0f}%</b>
            &rarr; <b>&lt;60%</b> dalam 2 tahun
          </div>
        </div>""", unsafe_allow_html=True)
        with st.expander("Lihat item berdasarkan volume penggunaan kayu"):
            st.dataframe(
                items_wood_rank.head(15).rename(columns={'Item': 'Item Name'}),
                use_container_width=True, hide_index=True)

    st.markdown("<br>", unsafe_allow_html=True)

    # ─────────────────────────────────────────────────────────────────────────
    # BAGIAN 8 — DOWNLOAD OVERVIEW (HTML / OFFLINE)
    # ─────────────────────────────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("### Download Overview — Offline Report")

    def build_overview_html():
        from datetime import datetime
        generated_at = datetime.now().strftime('%d %B %Y, %H:%M WIB')

        # ── Helpers ───────────────────────────────────────────────────────────
        def fig_to_json(fig):
            return fig.to_json()

        # ═══════════════════════════════════════════════
        # 1. DONUT — komposisi material
        # ═══════════════════════════════════════════════
        mat_dist_h = df_all['Material_Label'].value_counts().reset_index()
        mat_dist_h.columns = ['Material', 'Count']
        fig_donut_h = go.Figure(go.Pie(
            labels=mat_dist_h['Material'], values=mat_dist_h['Count'], hole=0.58,
            marker=dict(colors=[MATERIAL_COLORS.get(m,'#aaa') for m in mat_dist_h['Material']],
                        line=dict(color='white', width=2)),
            textinfo='percent+label', textfont=dict(size=12),
            hovertemplate='<b>%{label}</b><br>%{value:,} item (%{percent})<extra></extra>'
        ))
        fig_donut_h.add_annotation(
            text=f"<b>{wood_pct_all:.0f}%</b><br><span style='font-size:10px'>Wood</span>",
            x=0.5, y=0.5, font=dict(size=17, color='#7C4A19'), showarrow=False)
        fig_donut_h.update_layout(
            title=f"Komposisi Kemasan — {ov_label}", paper_bgcolor='white',
            height=360, margin=dict(l=10,r=10,t=45,b=10), showlegend=True,
            legend=dict(orientation='v', x=1.02, y=0.5, font=dict(size=11)))

        # ═══════════════════════════════════════════════
        # 2. TREND LINE — monthly all years
        # ═══════════════════════════════════════════════
        mn_ov = df_all.groupby(['YearMonth','Material_Label']).size().reset_index(name='Count')
        mn_ov['Label'] = mn_ov['YearMonth'].apply(fmt_month_str)
        mn_ov = mn_ov.sort_values('YearMonth')
        fig_trend_h = go.Figure()
        for mat in mat_opts:
            sub = mn_ov[mn_ov['Material_Label']==mat]
            if sub.empty: continue
            fig_trend_h.add_trace(go.Scatter(
                x=sub['Label'], y=sub['Count'], name=mat,
                mode='lines+markers',
                line=dict(color=MATERIAL_COLORS.get(mat,'#888'), width=2.2),
                marker=dict(size=4)))
        fig_trend_h.update_layout(
            title=f"Monthly Trend by Material — {ov_label}",
            paper_bgcolor='white', plot_bgcolor='white', height=340,
            margin=dict(l=20,r=20,t=45,b=60),
            xaxis=dict(showgrid=False, tickangle=-45, tickfont=dict(size=8), title='Month'),
            yaxis=dict(showgrid=True, gridcolor='#f1f5f9', title='Shipments'),
            legend=dict(orientation='h', y=1.12, x=1, xanchor='right'),
            hovermode='x unified')

        # ═══════════════════════════════════════════════
        # 3. PER-MATERIAL MINI TRENDS (4 charts)
        # ═══════════════════════════════════════════════
        mini_figs = {}
        for mat in mat_opts:
            sub = mn_ov[mn_ov['Material_Label']==mat].copy()
            color = MATERIAL_COLORS.get(mat,'#888')
            light = MATERIAL_LIGHT.get(mat,'#eee')
            fig_m = go.Figure()
            fig_m.add_trace(go.Scatter(
                x=sub['Label'], y=sub['Count'], name=mat,
                mode='lines+markers', fill='tozeroy',
                fillcolor=light, line=dict(color=color, width=2.5),
                marker=dict(size=5, color=color),
                hovertemplate='<b>%{x}</b><br>%{y:,} lines<extra></extra>'
            ))
            avg_v = sub['Count'].mean() if not sub.empty else 0
            fig_m.add_hline(y=avg_v, line_dash='dot', line_color='#94a3b8',
                            annotation_text=f"Avg:{avg_v:,.0f}",
                            annotation_position='bottom right',
                            annotation_font_size=10)
            fig_m.update_layout(
                title=dict(text=f"{mat} — Monthly Trend ({ov_label})", font=dict(size=12,color=color)),
                paper_bgcolor='white', plot_bgcolor='white', height=240,
                margin=dict(l=15,r=15,t=40,b=55),
                xaxis=dict(showgrid=False, tickangle=-45, tickfont=dict(size=7.5), title='Month'),
                yaxis=dict(showgrid=True, gridcolor='#f1f5f9', title='Shipments', rangemode='tozero'),
                showlegend=False, hovermode='x unified')
            mini_figs[mat] = fig_to_json(fig_m)

        # ═══════════════════════════════════════════════
        # 4. CORRELATION ANALYSIS — build data
        # ═══════════════════════════════════════════════
        _wood_monthly_h = (
            df_raw[df_raw['Material_Clean']=='W']
            .groupby('YearMonth').size()
            .reset_index(name='WoodLines')
            .sort_values('YearMonth')
        )
        _wood_monthly_h['Label'] = _wood_monthly_h['YearMonth'].apply(fmt_month_str)
        _wood_ref_h = pd.merge(_wood_monthly_h, REF_DF, on='YearMonth', how='outer').sort_values('YearMonth')
        _wood_ref_h['WoodLines']    = _wood_ref_h['WoodLines'].fillna(0)
        _wood_ref_h['DisplayLabel'] = _wood_ref_h['Label_x'].fillna(_wood_ref_h['Label_y'])

        if ov_view == "Yearly" and ov_year != "All":
            _wrp_h = _wood_ref_h[_wood_ref_h['YearMonth'].str.startswith(str(ov_year))].copy()
        else:
            _wrp_h = _wood_ref_h.copy()
        _wrp_h = _wrp_h.dropna(subset=['WoodLines','MetalUsageKG','ProductionUnit']).copy()
        _wrp_h = _wrp_h[_wrp_h['WoodLines']>0].reset_index(drop=True)

        _corr_m_h = _wrp_h['WoodLines'].corr(_wrp_h['MetalUsageKG'])  if len(_wrp_h)>=3 else 0
        _corr_p_h = _wrp_h['WoodLines'].corr(_wrp_h['ProductionUnit']) if len(_wrp_h)>=3 else 0

        def _interp_r_h(r):
            a = abs(r); d = "positif" if r>=0 else "negatif"
            if a>=0.7:   s,b="STRONG","🟢"
            elif a>=0.4: s,b="MODERATE","🟡"
            else:         s,b="WEAK","🔴"
            return d,s,b

        _dir_m_h,_str_m_h,_badge_m_h = _interp_r_h(_corr_m_h)
        _dir_p_h,_str_p_h,_badge_p_h = _interp_r_h(_corr_p_h)

        # direction agreement
        _da_h = _wrp_h.copy()
        _da_h['dW'] = _da_h['WoodLines'].diff()
        _da_h['dM'] = _da_h['MetalUsageKG'].diff()
        _da_h['dP'] = _da_h['ProductionUnit'].diff()
        _da_h = _da_h.dropna(subset=['dW','dM','dP'])
        _agree_m_h = int((_da_h['dW'].apply(np.sign)==_da_h['dM'].apply(np.sign)).sum()) if len(_da_h) else 0
        _agree_p_h = int((_da_h['dW'].apply(np.sign)==_da_h['dP'].apply(np.sign)).sum()) if len(_da_h) else 0
        _total_h   = len(_da_h)
        _pct_m_h   = _agree_m_h/_total_h*100 if _total_h else 0
        _pct_p_h   = _agree_p_h/_total_h*100 if _total_h else 0

        # avg metal/unit
        _avg_metal_h = _wrp_h['MetalUsageKG'].mean()   if len(_wrp_h) else 0
        _avg_prod_h  = _wrp_h['ProductionUnit'].mean()  if len(_wrp_h) else 0
        _avg_wood_h  = _wrp_h['WoodLines'].mean()       if len(_wrp_h) else 0
        _metal_per_unit_h = _avg_metal_h/_avg_prod_h if _avg_prod_h else 0

        # KPI card colors
        def _kpi_color(pct): return '#16a34a' if pct>=60 else ('#d97706' if pct>=45 else '#dc2626')

        _cm_col_h = _kpi_color(_pct_m_h)
        _cp_col_h = _kpi_color(_pct_p_h)

        # ── Corr: Raw Values chart ────────────────────────────────────────
        fig_raw_h = make_subplots(specs=[[{"secondary_y":True}]])
        fig_raw_h.add_trace(go.Bar(
            x=_wrp_h['DisplayLabel'], y=_wrp_h['WoodLines'],
            name='Wood Shipments', marker_color='#7C4A19', opacity=0.70,
            hovertemplate='<b>%{x}</b><br>Wood: <b>%{y:,} shipments</b><extra></extra>'
        ), secondary_y=False)
        fig_raw_h.add_trace(go.Scatter(
            x=_wrp_h['DisplayLabel'], y=_wrp_h['MetalUsageKG'],
            name='Metal Usage (KG)', mode='lines+markers',
            line=dict(color='#2563eb',width=2.5), marker=dict(size=5),
            hovertemplate='<b>%{x}</b><br>Metal: <b>%{y:,.0f} KG</b><extra></extra>'
        ), secondary_y=True)
        fig_raw_h.add_trace(go.Scatter(
            x=_wrp_h['DisplayLabel'], y=_wrp_h['ProductionUnit'],
            name='Production (Unit)', mode='lines+markers',
            line=dict(color='#16a34a',width=2.5,dash='dot'),
            marker=dict(size=5,symbol='diamond'),
            hovertemplate='<b>%{x}</b><br>Production: <b>%{y:,} units</b><extra></extra>'
        ), secondary_y=True)
        fig_raw_h.update_yaxes(title_text='Wood Packaging Lines', showgrid=True, gridcolor='#f1f5f9', secondary_y=False)
        fig_raw_h.update_yaxes(title_text='Metal (KG) / Produksi (Unit)', showgrid=False, secondary_y=True)
        fig_raw_h.update_layout(
            title="Raw Values — Wood Shipments vs Metal Usage vs Production",
            paper_bgcolor='white', plot_bgcolor='white', height=360,
            margin=dict(l=20,r=20,t=50,b=70),
            xaxis=dict(showgrid=False, tickangle=-45, tickfont=dict(size=8.5), title='Month'),
            legend=dict(orientation='h', y=1.12, x=0.5, xanchor='center', font=dict(size=11)),
            hovermode='x unified', barmode='overlay')

        # ── Corr: Movement Index chart ───────────────────────────────────
        if len(_wrp_h)>=2:
            _idx_h = _wrp_h.copy().reset_index(drop=True)
            bw = _idx_h['WoodLines'].iloc[0] or 1
            bm = _idx_h['MetalUsageKG'].iloc[0] or 1
            bp = _idx_h['ProductionUnit'].iloc[0] or 1
            _idx_h['idx_wood']  = _idx_h['WoodLines']     /bw*100
            _idx_h['idx_metal'] = _idx_h['MetalUsageKG']  /bm*100
            _idx_h['idx_prod']  = _idx_h['ProductionUnit']/bp*100
            fig_idx_h = go.Figure()
            fig_idx_h.add_trace(go.Scatter(
                x=_idx_h['DisplayLabel'], y=_idx_h['idx_wood'],
                name='Wood Shipments', mode='lines+markers',
                line=dict(color='#7C4A19',width=3), marker=dict(size=7,color='#7C4A19')))
            fig_idx_h.add_trace(go.Scatter(
                x=_idx_h['DisplayLabel'], y=_idx_h['idx_metal'],
                name='Metal Usage', mode='lines+markers',
                line=dict(color='#2563eb',width=2.5), marker=dict(size=5,color='#2563eb')))
            fig_idx_h.add_trace(go.Scatter(
                x=_idx_h['DisplayLabel'], y=_idx_h['idx_prod'],
                name='Production Output', mode='lines+markers',
                line=dict(color='#16a34a',width=2.5,dash='dot'),
                marker=dict(size=5,symbol='diamond',color='#16a34a')))
            fig_idx_h.add_hline(y=100, line_dash='dot', line_color='#cbd5e1',
                                annotation_text="Baseline (100)",
                                annotation_position='bottom left',
                                annotation_font=dict(size=10,color='#94a3b8'))
            fig_idx_h.update_layout(
                title="Movement Index (Base = 100 at first month)",
                paper_bgcolor='white', plot_bgcolor='white', height=360,
                margin=dict(l=20,r=20,t=50,b=70),
                xaxis=dict(showgrid=False, tickangle=-45, tickfont=dict(size=8.5), title='Month'),
                yaxis=dict(showgrid=True, gridcolor='#f1f5f9', title='Indeks (bulan pertama=100)'),
                legend=dict(orientation='h', y=1.12, x=0.5, xanchor='center', font=dict(size=11)),
                hovermode='x unified')
            idx_json_h = fig_to_json(fig_idx_h)
        else:
            idx_json_h = 'null'

        # ── Corr: Scatter plots (quadrant) ───────────────────────────────
        scatter_figs = {}
        if len(_wrp_h)>=3:
            _da2_h = _wrp_h.copy().reset_index(drop=True)
            _da2_h['pct_W'] = _da2_h['WoodLines'].pct_change()*100
            _da2_h['pct_M'] = _da2_h['MetalUsageKG'].pct_change()*100
            _da2_h['pct_P'] = _da2_h['ProductionUnit'].pct_change()*100
            _da2_h = _da2_h.dropna(subset=['pct_W','pct_M','pct_P']).reset_index(drop=True)

            _col_wm = ['#16a34a' if np.sign(w)==np.sign(m) else '#dc2626'
                       for w,m in zip(_da2_h['pct_W'],_da2_h['pct_M'])]
            _col_wp = ['#16a34a' if np.sign(w)==np.sign(p) else '#dc2626'
                       for w,p in zip(_da2_h['pct_W'],_da2_h['pct_P'])]
            _col_mp_h = ['#16a34a' if np.sign(m)==np.sign(p) else '#dc2626'
                         for m,p in zip(_da2_h['pct_M'],_da2_h['pct_P'])]

            _agree_mp_h = sum(1 for c in _col_mp_h if c=='#16a34a')
            _pct_mp_h   = _agree_mp_h/len(_col_mp_h)*100 if _col_mp_h else 0

            for (key, x_col, y_col, x_lbl, y_lbl, colors, x_title, y_title) in [
                ('wm', 'pct_W','pct_M', 'Wood','Metal', _col_wm,
                 'Wood — change from prev month (%)','Metal Usage — change from prev month (%)'),
                ('wp', 'pct_W','pct_P', 'Wood','Production', _col_wp,
                 'Wood — change from prev month (%)','Production — change from prev month (%)'),
                ('mp', 'pct_M','pct_P', 'Metal','Production', _col_mp_h,
                 'Metal Usage — change from prev month (%)','Production Unit — change from prev month (%)'),
            ]:
                xmax = max(abs(_da2_h[x_col]).max(),5)*1.3
                ymax = max(abs(_da2_h[y_col]).max(),5)*1.3
                fig_sc = go.Figure()
                for qx,qy,qc in [
                    ([0,xmax],[0,ymax],'rgba(220,252,231,0.55)'),
                    ([-xmax,0],[-ymax,0],'rgba(220,252,231,0.55)'),
                    ([-xmax,0],[0,ymax],'rgba(254,226,226,0.55)'),
                    ([0,xmax],[-ymax,0],'rgba(254,226,226,0.55)'),
                ]:
                    fig_sc.add_shape(type='rect',x0=qx[0],x1=qx[1],y0=qy[0],y1=qy[1],
                                     fillcolor=qc,line_width=0,layer='below')
                fig_sc.add_hline(y=0, line_color='#94a3b8', line_width=1.5)
                fig_sc.add_vline(x=0, line_color='#94a3b8', line_width=1.5)
                for ax,ay,txt,tc in [
                    (xmax*0.5, ymax*0.75, f'Same Dir\n{x_lbl}↑ {y_lbl}↑','#166534'),
                    (-xmax*0.5,-ymax*0.75,f'Same Dir\n{x_lbl}↓ {y_lbl}↓','#166534'),
                    (-xmax*0.5, ymax*0.75,f'Opposite\n{x_lbl}↓ {y_lbl}↑','#991b1b'),
                    (xmax*0.5,-ymax*0.75, f'Opposite\n{x_lbl}↑ {y_lbl}↓','#991b1b'),
                ]:
                    fig_sc.add_annotation(x=ax,y=ay,text=txt,showarrow=False,
                        font=dict(size=9,color=tc),bgcolor='rgba(255,255,255,0.8)',
                        borderpad=3,align='center')
                fig_sc.add_trace(go.Scatter(
                    x=list(_da2_h[x_col]), y=list(_da2_h[y_col]),
                    mode='markers+text',
                    marker=dict(color=colors, size=11, line=dict(color='white',width=1.5)),
                    text=_da2_h['DisplayLabel'].str.replace(' ','<br>'),
                    textposition='top center', textfont=dict(size=7,color='#475569'),
                    hovertemplate='<b>%{text}</b><extra></extra>'
                ))
                fig_sc.update_layout(
                    title=dict(text=f'{x_lbl} vs {y_lbl} — Monthly Direction',
                               font=dict(size=12), x=0.5, xanchor='center'),
                    paper_bgcolor='white', plot_bgcolor='white', height=360,
                    margin=dict(l=20,r=20,t=50,b=55),
                    xaxis=dict(showgrid=False,zeroline=False,title=x_title,
                               range=[-xmax,xmax],ticksuffix='%'),
                    yaxis=dict(showgrid=False,zeroline=False,title=y_title,
                               range=[-ymax,ymax],ticksuffix='%'),
                    showlegend=False)
                scatter_figs[key] = fig_to_json(fig_sc)
        else:
            _da2_h = pd.DataFrame()
            scatter_figs = {'wm':'null','wp':'null','mp':'null'}
            _pct_mp_h = 0
            _agree_mp_h = 0

        # verdict texts
        def _verdict_h(pct, sub1, sub2):
            if pct>=60: return f"✅ {sub1} <b>cenderung mengikuti</b> pergerakan {sub2}."
            elif pct>=45: return f"⚠️ {sub1} <b>tidak konsisten mengikuti</b> {sub2}."
            else: return f"🚨 {sub1} <b>bergerak relatif independen</b> dari {sub2}."

        verd_m_h  = _verdict_h(_pct_m_h,  "Wood Shipments", "Metal Usage")
        verd_p_h  = _verdict_h(_pct_p_h,  "Wood Shipments", "Production Output")
        verd_mp_h = _verdict_h(_pct_mp_h, "Metal Usage",    "Production Unit")

        def _badge_color(pct):
            return '#16a34a' if pct>=60 else ('#d97706' if pct>=45 else '#dc2626')
        def _badge_bg(pct):
            return '#f0fdf4' if pct>=60 else ('#fffbeb' if pct>=45 else '#fef2f2')
        def _emoji_badge(pct):
            return '🟢' if pct>=70 else ('🟡' if pct>=45 else '🔴')

        # ═══════════════════════════════════════════════
        # 5. DIM PROFILE TABLE
        # ═══════════════════════════════════════════════
        dim_rows_html = ""
        for cat in DIM_LABELS:
            sub_cat = df_all[df_all['Dim_Category']==cat]
            if sub_cat.empty: continue
            n_t=len(sub_cat)
            w_n2=int((sub_cat['Material_Clean']=='W').sum())
            cb_n2=int((sub_cat['Material_Clean']=='B').sum())
            w_p=w_n2/n_t*100 if n_t else 0
            cb_p=cb_n2/n_t*100 if n_t else 0
            avg_w=sub_cat[sub_cat['Net Weight']>0]['Net Weight'].mean()
            avg_v=sub_cat[sub_cat['Volume_m3']>0]['Volume_m3'].mean()
            if w_p>90: sts='🔴 Wajib Kayu'
            elif w_p>=40: sts='🟡 Perlu Dikaji'
            else: sts='🟢 Potensial Diganti'
            dim_rows_html += f"<tr><td>{cat}</td><td>{n_t:,}</td><td>{w_n2:,} ({w_p:.1f}%)</td><td>{cb_n2:,} ({cb_p:.1f}%)</td><td>{'%.1f'%avg_w if not pd.isna(avg_w) else '—'} kg</td><td>{sts}</td></tr>"

        # ═══════════════════════════════════════════════
        # 6. OVERPACKAGED CHART & TABLE
        # ═══════════════════════════════════════════════
        top10_h = overpack_items.head(10) if not overpack_items.empty else pd.DataFrame(columns=['Item','Lines'])
        fig_over_h = go.Figure(go.Bar(
            x=top10_h['Lines'].tolist() if not top10_h.empty else [],
            y=top10_h['Item'].tolist()  if not top10_h.empty else [],
            orientation='h',
            marker=dict(color=top10_h['Lines'].tolist() if not top10_h.empty else [],
                        colorscale=[[0,'#FEF2F2'],[1,'#dc2626']], showscale=False),
            text=(top10_h['Lines'].apply(lambda v:f"{v:,} lines").tolist() if not top10_h.empty else []),
            textposition='outside',
            hovertemplate='<b>%{y}</b><br>%{x:,} overpackaged<extra></extra>'
        ))
        fig_over_h.update_layout(
            title=f'Top 10 Item Overpackaged — {ov_label}',
            paper_bgcolor='white', plot_bgcolor='white', height=360,
            margin=dict(l=10,r=80,t=50,b=20),
            xaxis=dict(showgrid=True,gridcolor='#f1f5f9',title='Shipments'),
            yaxis=dict(showgrid=False,autorange='reversed',tickfont=dict(size=10)))

        over_rows_html = ""
        for _,row in overpack_items.head(20).iterrows():
            over_rows_html += f"<tr><td>{row['Item']}</td><td>{row['Lines']:,}</td></tr>"

        # ═══════════════════════════════════════════════
        # 7. CUSTOMER CHART & TABLE
        # ═══════════════════════════════════════════════
        if not cust_wood_pct.empty:
            fig_cust_h = go.Figure(go.Bar(
                x=cust_wood_pct['Name'].tolist(), y=cust_wood_pct['Pct_Wood'].tolist(),
                marker=dict(color=cust_wood_pct['Pct_Wood'].tolist(),
                            colorscale=[[0,'#FEF9F5'],[1,'#7C4A19']], showscale=False),
                text=cust_wood_pct['Pct_Wood'].apply(lambda v:f'{v:.1f}%').tolist(),
                textposition='outside',
                hovertemplate='<b>%{x}</b><br>%{y:.1f}% Wood<extra></extra>'
            ))
            fig_cust_h.update_layout(
                title=f'% Penggunaan Wood per Customer (Top 10) — {ov_label}',
                paper_bgcolor='white', plot_bgcolor='white', height=340,
                margin=dict(l=20,r=20,t=50,b=80),
                xaxis=dict(showgrid=False,tickangle=-30,tickfont=dict(size=9)),
                yaxis=dict(showgrid=True,gridcolor='#f1f5f9',title='% Wood',range=[0,115]))
            cust_json_h = fig_to_json(fig_cust_h)
            cust_rows_html = ""
            for _,row in cust_wood_pct.iterrows():
                p=row['Pct_Wood']
                cat_c='🔴 Prioritas' if p>90 else ('🟡 Perlu Review' if p>=50 else '🟢 Relatif Baik')
                cust_rows_html += f"<tr><td>{row['Name']}</td><td>{int(row['Lines']):,}</td><td>{p:.1f}%</td><td>{cat_c}</td></tr>"
        else:
            cust_json_h = 'null'
            cust_rows_html = "<tr><td colspan='4'>Tidak ada data customer.</td></tr>"

        # ═══════════════════════════════════════════════
        # 8. FORECAST TABLE
        # ═══════════════════════════════════════════════
        fc_rows_html = ""
        fc_trend_str = "—"
        if ov_view != "Monthly":
            mfc, fcd, fcp, _ = forecast_xgboost('Wood', periods=12)
            if mfc is not None and fcd is not None:
                avg_h_fc = float(mfc['Count'].mean())
                for _,frow in fcd.iterrows():
                    vs=frow['Forecast']-avg_h_fc
                    vs_s=f"+{vs:,.0f}" if vs>=0 else f"{vs:,.0f}"
                    fc_rows_html += f"<tr><td>{fmt_month_str(frow['YearMonth'])}</td><td>{frow['Forecast']:,.0f}</td><td>{vs_s}</td></tr>"
                fc_trend_str = f"{'▲' if fcp>0 else '▼'} {abs(fcp):.1f}% terhadap rata-rata historis"

        # ═══════════════════════════════════════════════
        # 9. YEARLY SUMMARY TABLE
        # ═══════════════════════════════════════════════
        yearly_summary_html = ""
        if ov_view == "Yearly" and ov_year == "All":
            yr_totals_h = {yr: int((df_raw['Year']==yr).sum()) for yr in year_opts}
            yearly_summary_html = "<table><thead><tr><th>Kemasan</th>"
            for yr in year_opts:
                yearly_summary_html += f"<th>{yr}</th>"
            yearly_summary_html += "<th>Total</th><th>%</th></tr></thead><tbody>"
            total_raw_h = len(df_raw)
            for mat in mat_opts:
                yearly_summary_html += f"<tr><td>{mat}</td>"
                tot_mat=0
                for yr in year_opts:
                    cnt2=int(((df_raw['Material_Label']==mat)&(df_raw['Year']==yr)).sum())
                    yr_tot=yr_totals_h[yr]
                    pct2=cnt2/yr_tot*100 if yr_tot else 0
                    yearly_summary_html += f"<td>{cnt2:,} ({pct2:.0f}%)</td>"
                    tot_mat+=cnt2
                yearly_summary_html += f"<td>{tot_mat:,}</td><td>{tot_mat/total_raw_h*100:.1f}%</td></tr>"
            yearly_summary_html += "</tbody></table>"

        # ══════════════════════════════════════════════
        # COLLECT ALL JSON
        # ══════════════════════════════════════════════
        donut_json  = fig_to_json(fig_donut_h)
        trend_json  = fig_to_json(fig_trend_h)
        raw_json    = fig_to_json(fig_raw_h)
        over_json   = fig_to_json(fig_over_h)

        sc_wm_json = scatter_figs.get('wm','null')
        sc_wp_json = scatter_figs.get('wp','null')
        sc_mp_json = scatter_figs.get('mp','null')

        def _card_color_css(pct):
            c=_badge_color(pct); bg=_badge_bg(pct)
            return f'background:{bg};border-top:4px solid {c};'

        # ══════════════════════════════════════════════
        # BUILD HTML
        # ══════════════════════════════════════════════
        html = f"""<!DOCTYPE html>
<html lang="id">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1.0">
<title>Overview Report — Packaging Analysis PT Güntner</title>
<script src="https://cdn.plot.ly/plotly-2.35.2.min.js"></script>
<style>
*{{box-sizing:border-box;margin:0;padding:0;}}
body{{font-family:'Segoe UI',Arial,sans-serif;background:#f1f5f9;color:#1e293b;font-size:14px;line-height:1.5;}}
.header{{background:white;border-bottom:2px solid #e5e7eb;padding:18px 40px;
  display:flex;align-items:center;gap:20px;position:sticky;top:0;z-index:100;}}
.badge{{background:linear-gradient(135deg,#166534,#16a34a);color:white;
  padding:7px 20px;border-radius:20px;font-size:12px;font-weight:700;margin-left:auto;letter-spacing:.3px;}}
.container{{max-width:1200px;margin:0 auto;padding:28px 24px;}}
.meta-bar{{background:white;border-radius:10px;padding:12px 20px;margin-bottom:24px;
  display:flex;align-items:center;gap:8px;font-size:13px;color:#64748b;border:1px solid #e5e7eb;}}
.meta-bar b{{color:#1e293b;}}
.meta-sep{{color:#cbd5e1;}}
.section{{margin-bottom:36px;}}
.section-title{{font-size:15px;font-weight:700;color:#0f172a;margin-bottom:16px;
  padding-bottom:10px;border-bottom:2px solid #e5e7eb;letter-spacing:.1px;}}
.subsection-title{{font-size:13px;font-weight:700;color:#374151;margin-bottom:10px;text-transform:uppercase;letter-spacing:.5px;}}
.kpi-grid{{display:grid;grid-template-columns:repeat(4,1fr);gap:14px;margin-bottom:18px;}}
.kpi-card{{background:white;border-radius:12px;padding:18px 20px;border:1px solid #e5e7eb;}}
.kpi-label{{font-size:10.5px;font-weight:700;color:#94a3b8;letter-spacing:1px;
  text-transform:uppercase;margin-bottom:6px;}}
.kpi-value{{font-size:28px;font-weight:800;line-height:1.1;}}
.kpi-sub{{font-size:11.5px;color:#64748b;margin-top:5px;}}
.progress-box{{background:white;border-radius:12px;padding:20px 24px;
  border:1px solid #e5e7eb;margin-bottom:18px;}}
.progress-bar-bg{{background:#f1f5f9;border-radius:8px;height:20px;overflow:hidden;position:relative;margin:12px 0 8px;}}
.progress-bar-fill{{background:linear-gradient(90deg,#166534,#4ade80);height:100%;border-radius:8px;transition:width .6s ease;}}
.progress-label{{position:absolute;top:3px;left:50%;transform:translateX(-50%);
  font-size:11.5px;font-weight:700;color:#166534;white-space:nowrap;}}
.progress-meta{{display:flex;justify-content:space-between;align-items:center;margin-bottom:4px;flex-wrap:wrap;gap:8px;}}
.progress-notes{{font-size:12px;color:#64748b;margin-top:8px;display:flex;gap:16px;flex-wrap:wrap;}}
.grid-2{{display:grid;grid-template-columns:1fr 1.8fr;gap:16px;margin-bottom:16px;}}
.grid-2eq{{display:grid;grid-template-columns:1fr 1fr;gap:16px;margin-bottom:16px;}}
.grid-3{{display:grid;grid-template-columns:repeat(3,1fr);gap:14px;margin-bottom:16px;}}
.grid-4{{display:grid;grid-template-columns:repeat(4,1fr);gap:14px;margin-bottom:16px;}}
.chart-box{{background:white;border-radius:12px;padding:16px;border:1px solid #e5e7eb;}}
.note-box{{background:#f8fafc;border:1px solid #e2e8f0;padding:10px 16px;border-radius:8px;
  font-size:12px;color:#64748b;margin-bottom:12px;line-height:1.7;}}
.info-green{{background:#f0fdf4;border-left:3px solid #16a34a;padding:12px 18px;
  border-radius:0 8px 8px 0;margin:12px 0;font-size:13px;line-height:1.7;color:#166534;}}
.info-yellow{{background:#fffbeb;border-left:3px solid #d97706;padding:12px 18px;
  border-radius:0 8px 8px 0;margin:12px 0;font-size:13px;line-height:1.7;color:#92400e;}}
.info-red{{background:#fef2f2;border-left:3px solid #dc2626;padding:12px 18px;
  border-radius:0 8px 8px 0;margin:12px 0;font-size:13px;line-height:1.7;color:#991b1b;}}
table{{width:100%;border-collapse:collapse;font-size:13px;}}
thead tr{{background:#f8fafc;}}
th{{padding:10px 14px;text-align:left;color:#374151;font-weight:700;font-size:12px;
  letter-spacing:.3px;border-bottom:2px solid #e5e7eb;}}
td{{padding:9px 14px;border-bottom:1px solid #f1f5f9;color:#374151;}}
tr:last-child td{{border-bottom:none;}}
tr:hover td{{background:#fafbfc;}}
.tag-red{{background:#fef2f2;color:#dc2626;padding:3px 10px;border-radius:6px;font-weight:700;font-size:12px;}}
.tag-yellow{{background:#fffbeb;color:#d97706;padding:3px 10px;border-radius:6px;font-weight:700;font-size:12px;}}
.tag-green{{background:#f0fdf4;color:#16a34a;padding:3px 10px;border-radius:6px;font-weight:700;font-size:12px;}}
.roadmap-grid{{display:grid;grid-template-columns:repeat(3,1fr);gap:16px;margin-top:14px;}}
.roadmap-card{{background:white;border-radius:12px;padding:20px;border:1px solid #e5e7eb;}}
.roadmap-phase{{font-size:10.5px;font-weight:700;text-transform:uppercase;letter-spacing:.7px;margin-bottom:6px;}}
.roadmap-title{{font-size:18px;font-weight:800;margin-bottom:10px;}}
.roadmap-card ul{{font-size:13px;color:#374151;padding-left:16px;line-height:2;margin:8px 0 14px;}}
.roadmap-target{{border-radius:8px;padding:10px 14px;font-size:12.5px;line-height:1.6;}}
.corr-kpi{{background:white;border-radius:12px;padding:16px 18px;border:1px solid #e5e7eb;text-align:center;}}
.corr-kpi-label{{font-size:10.5px;color:#94a3b8;font-weight:700;text-transform:uppercase;
  letter-spacing:.8px;margin-bottom:6px;}}
.corr-kpi-value{{font-size:30px;font-weight:800;}}
.corr-kpi-sub{{font-size:12px;color:#64748b;margin-top:5px;line-height:1.5;}}
.verdict-box{{padding:14px 18px;border-radius:0 10px 10px 0;font-size:13px;line-height:1.7;}}
.footer{{text-align:center;color:#94a3b8;padding:24px;background:white;
  border-radius:12px;margin-top:36px;border:1px solid #e5e7eb;}}
@media(max-width:900px){{
  .kpi-grid,.grid-4{{grid-template-columns:repeat(2,1fr);}}
  .grid-2,.grid-2eq,.grid-3,.roadmap-grid{{grid-template-columns:1fr;}}
  .meta-bar{{flex-wrap:wrap;}}
}}
@media print{{
  .header{{position:static;}}
  .chart-box,.section{{page-break-inside:avoid;}}
}}
</style>
</head>
<body>

<div class="header">
  <div style="width:48px;height:48px;display:flex;align-items:center;justify-content:center;">
    <svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 64 64" width="48" height="48">
      <path d="M32 56C32 56 10 44 10 28C10 16 20 8 32 8C32 8 14 22 20 38C24 48 32 56 32 56Z" fill="#166534" opacity="0.9"/>
      <path d="M32 56C32 56 54 44 54 28C54 16 44 8 32 8C32 8 50 22 44 38C40 48 32 56 32 56Z" fill="#16a34a" opacity="0.85"/>
    </svg>
  </div>
  <div>
    <div style="font-size:22px;font-weight:800;color:#111827;letter-spacing:-.2px;">Packaging Analysis — Overview Report</div>
    <div style="font-size:12.5px;color:#9ca3af;margin-top:3px;">Supporting Pledge 6 · Make all our packaging fully sustainable · PT Güntner Indonesia</div>
  </div>
  <div class="badge">Pledge 6 Tracker</div>
</div>

<div class="container">
<div class="meta-bar">
  <span>Periode</span>
  <b>{ov_label}</b>
  <span class="meta-sep">·</span>
  <span>{total_l:,} pengiriman</span>
  <span class="meta-sep">·</span>
  <span>Dibuat</span>
  <b>{generated_at}</b>
</div>

<!-- SECTION 1: KPI SCORECARD -->
<div class="section">
  <div class="section-title">KPI Scorecard</div>
  <div class="kpi-grid">
    <div class="kpi-card" style="border-top:3px solid #2563eb;">
      <div class="kpi-label">Total Shipments</div>
      <div class="kpi-value" style="color:#2563eb;">{total_l:,}</div>
      <div class="kpi-sub">Periode: {ov_label}</div>
    </div>
    <div class="kpi-card" style="border-top:3px solid #7C4A19;">
      <div class="kpi-label">Wood Shipments</div>
      <div class="kpi-value" style="color:#7C4A19;">{wood_n:,}
        <span style="font-size:15px;font-weight:600;opacity:.65;">({wood_pct_all:.1f}%)</span></div>
      <div class="kpi-sub">dari {total_l:,} total pengiriman</div>
    </div>
    <div class="kpi-card" style="border-top:3px solid #d97706;">
      <div class="kpi-label">Switchable to Carton</div>
      <div class="kpi-value" style="color:#d97706;">{overpack_n_all:,}
        <span style="font-size:15px;font-weight:600;opacity:.65;">({overpack_pct_all:.1f}%)</span></div>
      <div class="kpi-sub">dimensi ≤2m &amp; berat ≤75 kg</div>
    </div>
    <div class="kpi-card" style="border-top:3px solid #16a34a;">
      <div class="kpi-label">Sustainable Packaging</div>
      <div class="kpi-value" style="color:#16a34a;">{non_wood_n:,}
        <span style="font-size:15px;font-weight:600;opacity:.65;">({non_wood_pct_all:.1f}%)</span></div>
      <div class="kpi-sub">Carton Box + Cardboard + Carton</div>
    </div>
  </div>

  <div class="progress-box">
    <div class="progress-meta">
      <div style="font-size:14px;font-weight:700;color:#166534;">Pledge 6 Progress — Target: 0% Wood</div>
      <div style="font-size:12px;color:#64748b;">
        Saat ini: <b style="color:#dc2626;">{wood_pct_all:.1f}% Wood</b> &nbsp;·&nbsp;
        Target: <b style="color:#16a34a;">0% Wood</b>
      </div>
    </div>
    <div class="progress-bar-bg">
      <div class="progress-bar-fill" style="width:{progress_pct:.1f}%;"></div>
      <div class="progress-label">{progress_pct:.1f}% tercapai</div>
    </div>
    <div class="progress-notes">
      <span>Quick win: konversi <b>{overpack_n_all:,} item kecil</b> ke Carton</span>
      <span>· tambah <b>~{overpack_pct_all:.1f}%</b> kemajuan</span>
      <span>· sisa gap: <b>{wood_pct_all:.1f}%</b></span>
    </div>
  </div>
</div>

<!-- SECTION 2: KOMPOSISI & TREN -->
<div class="section">
  <div class="section-title">Packaging Composition &amp; Trend</div>
  <div class="grid-2">
    <div class="chart-box"><div id="c_donut" style="height:340px;"></div></div>
    <div class="chart-box"><div id="c_trend" style="height:340px;"></div></div>
  </div>
  <div class="info-green"><b>Insight ({ov_label}):</b> Wood mendominasi <b>{wood_pct_all:.1f}%</b>
    dari total shipment. Peluang terbesar ada di kategori ukuran kecil
    (XS: {xs_pct_val:.0f}% Wood, S: {s_pct_val:.0f}% Wood).</div>
</div>

<!-- YEARLY SUMMARY TABLE -->
{f'<div class="section"><div class="section-title">Volume Pengiriman per Tahun</div><div class="chart-box">{yearly_summary_html}</div></div>' if yearly_summary_html else ''}

<!-- SECTION 2B: MONTHLY TREND PER MATERIAL -->
<div class="section">
  <div class="section-title">Monthly Trend per Packaging Type</div>
  <div class="grid-2eq">
    <div class="chart-box"><div id="c_mini_wood" style="height:240px;"></div></div>
    <div class="chart-box"><div id="c_mini_cb" style="height:240px;"></div></div>
  </div>
  <div class="grid-2eq">
    <div class="chart-box"><div id="c_mini_cbd" style="height:240px;"></div></div>
    <div class="chart-box"><div id="c_mini_ct" style="height:240px;"></div></div>
  </div>
</div>

<!-- SECTION 2C: CORRELATION ANALYSIS -->
<div class="section">
  <div class="section-title">Correlation Analysis: Wood Shipments vs Metal Usage &amp; Production</div>

  <div class="subsection-title">Trend Summary</div>
  <div class="grid-4" style="margin-bottom:18px;">
    <div class="corr-kpi" style="{_card_color_css(_pct_m_h)}">
      <div class="corr-kpi-label">Wood &amp; Metal Trend</div>
      <div class="corr-kpi-value" style="color:{_badge_color(_pct_m_h)};">{_pct_m_h:.0f}%</div>
      <div class="corr-kpi-sub">{_agree_m_h} dari {_total_h} bulan<br>bergerak searah</div>
    </div>
    <div class="corr-kpi" style="{_card_color_css(_pct_p_h)}">
      <div class="corr-kpi-label">Wood &amp; Production Trend</div>
      <div class="corr-kpi-value" style="color:{_badge_color(_pct_p_h)};">{_pct_p_h:.0f}%</div>
      <div class="corr-kpi-sub">{_agree_p_h} dari {_total_h} bulan<br>bergerak searah</div>
    </div>
    <div class="corr-kpi" style="border-top:3px solid #2563eb;">
      <div class="corr-kpi-label">Avg Metal / Unit</div>
      <div class="corr-kpi-value" style="color:#2563eb;">{_metal_per_unit_h:,.1f} KG</div>
      <div class="corr-kpi-sub">per unit produksi<br>avg {_avg_metal_h/1000:,.0f}t metal · {_avg_prod_h:,.0f} unit/bln</div>
    </div>
    <div class="corr-kpi" style="border-top:3px solid #7C4A19;">
      <div class="corr-kpi-label">Avg Wood / Month</div>
      <div class="corr-kpi-value" style="color:#7C4A19;">{_avg_wood_h:,.0f}</div>
      <div class="corr-kpi-sub">pengiriman per bulan<br>dari {len(_wrp_h)} bulan data</div>
    </div>
  </div>

  <!-- Correlation r values -->
  <div class="grid-2eq" style="margin-bottom:18px;">
    <div class="chart-box" style="padding:18px 22px;">
      <div style="font-size:12px;font-weight:700;color:#64748b;margin-bottom:6px;text-transform:uppercase;letter-spacing:.5px;">
        Wood ↔ Metal Usage — Pearson r
      </div>
      <div style="font-size:40px;font-weight:800;color:#2563eb;line-height:1;">{_corr_m_h:.3f}</div>
      <div style="font-size:12px;color:#64748b;margin-top:8px;">
        Korelasi <b>{_str_m_h}</b> ({_dir_m_h}) · {_pct_m_h:.0f}% bulan bergerak searah
      </div>
    </div>
    <div class="chart-box" style="padding:18px 22px;">
      <div style="font-size:12px;font-weight:700;color:#64748b;margin-bottom:6px;text-transform:uppercase;letter-spacing:.5px;">
        Wood ↔ Production Output — Pearson r
      </div>
      <div style="font-size:40px;font-weight:800;color:#16a34a;line-height:1;">{_corr_p_h:.3f}</div>
      <div style="font-size:12px;color:#64748b;margin-top:8px;">
        Korelasi <b>{_str_p_h}</b> ({_dir_p_h}) · {_pct_p_h:.0f}% bulan bergerak searah
      </div>
    </div>
  </div>

  <!-- Raw Values Chart -->
  <div class="subsection-title">Raw Values (Individual Scale)</div>
  <div class="note-box">
    Sumbu kiri (coklat): Wood Packaging Shipments ·
    Sumbu kanan (biru &amp; hijau): Metal Usage (KG) dan Production Output (Unit)
  </div>
  <div class="chart-box" style="margin-bottom:16px;"><div id="c_raw" style="height:360px;"></div></div>

  <!-- Movement Index Chart -->
  <div class="subsection-title">Movement Index (Base = 100)</div>
  <div class="note-box">
    Nilai asli disetarakan ke skala yang sama — bulan pertama = 100.
    Jika ketiga garis bergerak bersamaan, Wood Shipments mengikuti pola Production &amp; Metal Usage.
  </div>
  <div class="chart-box" style="margin-bottom:16px;"><div id="c_idx" style="height:360px;"></div></div>

  <!-- Scatter Charts -->
  <div class="subsection-title">Monthly Direction — Wood vs Metal &amp; Production</div>
  <div class="note-box">
    Setiap titik mewakili satu bulan.
    <span style="color:#166534;font-weight:600;">Hijau</span>: bergerak searah &nbsp;·&nbsp;
    <span style="color:#991b1b;font-weight:600;">Merah</span>: berlawanan arah
  </div>
  <div class="grid-3" style="margin-bottom:16px;">
    <div class="chart-box"><div id="c_sc_wm" style="height:340px;"></div></div>
    <div class="chart-box"><div id="c_sc_wp" style="height:340px;"></div></div>
    <div class="chart-box"><div id="c_sc_mp" style="height:340px;"></div></div>
  </div>

  <!-- Verdict Summary -->
  <div class="grid-3">
    <div class="verdict-box" style="background:{_badge_bg(_pct_m_h)};border-left:3px solid {_badge_color(_pct_m_h)};">
      <b>Wood ↔ Metal Usage</b><br>
      Searah dalam <b>{_agree_m_h}/{_total_h} bulan ({_pct_m_h:.0f}%)</b><br>
      <span style="color:#64748b;">{verd_m_h}</span>
    </div>
    <div class="verdict-box" style="background:{_badge_bg(_pct_p_h)};border-left:3px solid {_badge_color(_pct_p_h)};">
      <b>Wood ↔ Production Output</b><br>
      Searah dalam <b>{_agree_p_h}/{_total_h} bulan ({_pct_p_h:.0f}%)</b><br>
      <span style="color:#64748b;">{verd_p_h}</span>
    </div>
    <div class="verdict-box" style="background:{_badge_bg(_pct_mp_h)};border-left:3px solid {_badge_color(_pct_mp_h)};">
      <b>Metal Usage ↔ Production Unit</b><br>
      Searah dalam <b>{_agree_mp_h}/{len(_col_mp_h) if len(scatter_figs)>0 and 'wm' in scatter_figs and scatter_figs['wm']!='null' else _total_h} bulan ({_pct_mp_h:.0f}%)</b><br>
      <span style="color:#64748b;">{verd_mp_h}</span>
    </div>
  </div>
</div>

<!-- SECTION 3: DIMENSI & BERAT -->
<div class="section">
  <div class="section-title">Dimensi &amp; Berat — Pilihan Kemasan</div>
  <div class="chart-box" style="margin-bottom:14px;">
    <div class="subsection-title" style="margin-bottom:12px;">Profil Kemasan per Kategori Ukuran</div>
    <table>
      <thead><tr><th>Kategori</th><th>Jumlah Items</th><th>Wood</th><th>Carton Box</th><th>Rata-rata Berat</th><th>Status</th></tr></thead>
      <tbody>{dim_rows_html}</tbody>
    </table>
  </div>
  <div class="info-yellow">
    <b>Temuan Kritis:</b> Item kategori M (2–4m), L (4–7m), dan XL (&gt;7m)
    menggunakan Wood hampir 100% karena alasan struktural.
    Peluang penggantian terfokus pada XS dan S —
    XS: {xs_pct_val:.0f}% Wood, S: {s_pct_val:.0f}% Wood pada periode {ov_label}.
  </div>
  <div class="chart-box" style="margin-top:14px;">
    <div class="subsection-title" style="margin-bottom:12px;">Panduan Pemilihan Kemasan</div>
    <table>
      <thead><tr><th>Ukuran &amp; Berat Item</th><th>Kemasan yang Tepat</th><th>Alasan</th></tr></thead>
      <tbody>
        <tr><td>Panjang/Lebar/Tinggi &gt; 4 m <b>atau</b> Berat &gt; 200 kg</td>
          <td><span class="tag-red">Wajib Kayu</span></td>
          <td style="color:#64748b;">Item besar/berat butuh palet kayu untuk kekuatan struktural</td></tr>
        <tr><td>Ukuran 1.5–4 m <b>dan</b> Berat 50–200 kg</td>
          <td><span class="tag-yellow">Perlu Dikaji</span></td>
          <td style="color:#64748b;">Tergantung bentuk item — harus dievaluasi satu per satu</td></tr>
        <tr><td>Semua sisi &lt; 1.5 m <b>dan</b> Berat &lt; 50 kg</td>
          <td><span class="tag-green">Ganti Karton</span></td>
          <td style="color:#64748b;">Item kecil &amp; ringan cukup dilindungi karton</td></tr>
      </tbody>
    </table>
  </div>
</div>

<!-- SECTION 4: PELUANG KONVERSI -->
<div class="section">
  <div class="section-title">Peluang Konversi: Wood → Material Alternatif</div>
  <div class="grid-3" style="margin-bottom:14px;">
    <div class="kpi-card" style="border-top:3px solid #dc2626;">
      <div class="kpi-label">Item Wood Overpackaged</div>
      <div class="kpi-value" style="color:#dc2626;">{overpack_n_all:,}</div>
      <div class="kpi-sub">Dimensi &lt;1.5m &amp; berat &lt;50kg — tapi pakai kayu</div>
    </div>
    <div class="kpi-card" style="border-top:3px solid #d97706;">
      <div class="kpi-label">% dari Total Wood</div>
      <div class="kpi-value" style="color:#d97706;">{overpack_pct_all:.1f}%</div>
      <div class="kpi-sub">Quick win langsung untuk Pledge 6</div>
    </div>
    <div class="kpi-card" style="border-top:3px solid #16a34a;">
      <div class="kpi-label">Target Konversi</div>
      <div class="kpi-value" style="color:#16a34a;">{overpack_n_all:,}</div>
      <div class="kpi-sub">→ hemat ~{overpack_pct_all:.1f}% Wood</div>
    </div>
  </div>
  <div class="chart-box" style="margin-bottom:14px;"><div id="c_over" style="height:360px;"></div></div>
  <div class="info-red">
    <b>Quick Win ({ov_label}):</b> Terdapat
    <b>{overpack_n_all:,} item ({overpack_pct_all:.1f}% dari total Wood)</b>
    berukuran kecil (&lt;1.500 mm) dan ringan (&lt;50 kg) namun masih pakai kayu.
    Mengganti ke Carton Box adalah langkah tercepat untuk Pledge 6.
  </div>
  <div class="chart-box" style="margin-top:14px;">
    <div class="subsection-title" style="margin-bottom:12px;">20 Item Prioritas Konversi</div>
    <table><thead><tr><th>Item Name</th><th>Shipments</th></tr></thead>
    <tbody>{over_rows_html}</tbody></table>
  </div>
</div>

<!-- SECTION 5: CUSTOMER -->
<div class="section">
  <div class="section-title">Customer &amp; Penggunaan Wood</div>
  <div class="chart-box" style="margin-bottom:14px;"><div id="c_cust" style="height:340px;"></div></div>
  <div class="chart-box" style="margin-bottom:14px;">
    <table><thead><tr><th>Pelanggan</th><th>Total Pengiriman</th><th>% Pakai Kayu</th><th>Kategori</th></tr></thead>
    <tbody>{cust_rows_html}</tbody></table>
  </div>
  <div class="info-yellow">
    <b>Insight ({ov_label}):</b> Customer <b>{top_cust_name}</b> adalah pengguna Wood terbesar.
    Customer dengan % Wood tinggi adalah kandidat utama kolaborasi Pledge 6.
  </div>
</div>

<!-- SECTION 6: FORECAST -->
<div class="section">
  <div class="section-title">Proyeksi Penggunaan Wood — 12 Bulan ke Depan (XGBoost)</div>
  {"<div class='info-yellow'>Proyeksi tidak tersedia untuk tampilan Monthly.</div>" if ov_view=='Monthly' else f"""
  <div class="chart-box">
    <table><thead><tr><th>Bulan</th><th>Prediksi Lines</th><th>vs Rata-rata Historis</th></tr></thead>
    <tbody>{fc_rows_html}</tbody></table>
    <div style="margin-top:14px;font-size:13px;color:#64748b;border-top:1px solid #f1f5f9;padding-top:12px;">
      Tren Forecast: <b style="color:#1e293b;">{fc_trend_str}</b>
    </div>
  </div>"""}
</div>

<!-- SECTION 7: ROADMAP -->
<div class="section">
  <div class="section-title">Roadmap Aksi menuju Pledge 6</div>
  <div class="roadmap-grid">
    <div class="roadmap-card" style="border-top:3px solid #16a34a;">
      <div class="roadmap-phase" style="color:#16a34a;">0 – 3 Bulan</div>
      <div class="roadmap-title" style="color:#16a34a;">Quick Win</div>
      <ul>
        <li>Konversi <b>{overpack_n_all:,} item</b> dari Wood ke Carton Box</li>
        <li>Fokus: dimensi &lt;1.500 mm, berat &lt;50 kg</li>
        <li>Sosialisasi panduan kemasan ke tim operasional</li>
      </ul>
      <div class="roadmap-target" style="background:#f0fdf4;color:#166534;">
        <b>Target:</b> Kurangi <b>{overpack_n_all:,} item</b> Wood kecil · hemat <b>~{overpack_pct_all:.0f}%</b>
      </div>
    </div>
    <div class="roadmap-card" style="border-top:3px solid #2563eb;">
      <div class="roadmap-phase" style="color:#2563eb;">3 – 12 Bulan</div>
      <div class="roadmap-title" style="color:#2563eb;">Strategic</div>
      <ul>
        <li>Kolaborasi dengan <b>{top_cust_name}</b> &amp; customer besar lain</li>
        <li>Review item ukuran S/M (1.5–4 m) per kasus</li>
        <li>Pilot program kemasan alternatif untuk "gray area"</li>
      </ul>
      <div class="roadmap-target" style="background:#eff6ff;color:#1e3a8a;">
        <b>Target:</b> Review <b>{n_sm:,} item</b> kategori S &amp; M yang masih pakai Wood
      </div>
    </div>
    <div class="roadmap-card" style="border-top:3px solid #7C4A19;">
      <div class="roadmap-phase" style="color:#7C4A19;">12+ Bulan</div>
      <div class="roadmap-title" style="color:#7C4A19;">Transformational</div>
      <ul>
        <li>Desain ulang kemasan untuk item ukuran M/L</li>
        <li>Sertifikasi kemasan berkelanjutan untuk Pledge 6</li>
        <li>Eksplorasi material inovatif: recycled plastic foil, reusable packaging</li>
      </ul>
      <div class="roadmap-target" style="background:#fdf6f0;color:#7C4A19;">
        <b>Target:</b> Kurangi Wood dari <b>{wood_pct_all:.0f}%</b> → <b>&lt;60%</b> dalam 2 tahun
      </div>
    </div>
  </div>
</div>

<!-- FOOTER -->
<div class="footer">
  <p style="font-weight:800;font-size:14px;color:#166534;margin-bottom:6px;">
    Packaging Analysis Dashboard · Overview Report</p>
  <p style="font-size:13px;">PT Güntner Indonesia &nbsp;·&nbsp; Supporting Pledge 6: Make all our packaging fully sustainable</p>
  <p style="font-size:11px;color:#cbd5e1;margin-top:8px;">
    {generated_at} &nbsp;·&nbsp; Streamlit · Plotly · XGBoost</p>
</div>

</div><!-- /container -->

<script>
var donutD  = {donut_json};
var trendD  = {trend_json};
var miniW   = {mini_figs.get('Wood','null')};
var miniCB  = {mini_figs.get('Carton Box','null')};
var miniCBD = {mini_figs.get('Cardboard','null')};
var miniCT  = {mini_figs.get('Carton','null')};
var rawD    = {raw_json};
var idxD    = {idx_json_h};
var scWM    = {sc_wm_json};
var scWP    = {sc_wp_json};
var scMP    = {sc_mp_json};
var overD   = {over_json};
var custD   = {cust_json_h};

var cfg = {{responsive:true, displayModeBar:false}};

Plotly.newPlot('c_donut',  donutD.data,  donutD.layout,  cfg);
Plotly.newPlot('c_trend',  trendD.data,  trendD.layout,  cfg);
if(miniW)  Plotly.newPlot('c_mini_wood', miniW.data,  miniW.layout,  cfg);
if(miniCB) Plotly.newPlot('c_mini_cb',  miniCB.data, miniCB.layout, cfg);
if(miniCBD)Plotly.newPlot('c_mini_cbd', miniCBD.data,miniCBD.layout,cfg);
if(miniCT) Plotly.newPlot('c_mini_ct',  miniCT.data, miniCT.layout, cfg);
if(rawD)   Plotly.newPlot('c_raw', rawD.data, rawD.layout, cfg);
if(idxD)   Plotly.newPlot('c_idx', idxD.data, idxD.layout, cfg);
if(scWM && scWM!='null') Plotly.newPlot('c_sc_wm', scWM.data, scWM.layout, cfg);
if(scWP && scWP!='null') Plotly.newPlot('c_sc_wp', scWP.data, scWP.layout, cfg);
if(scMP && scMP!='null') Plotly.newPlot('c_sc_mp', scMP.data, scMP.layout, cfg);
if(overD)  Plotly.newPlot('c_over', overD.data, overD.layout, cfg);
if(custD && custD!='null') Plotly.newPlot('c_cust', custD.data, custD.layout, cfg);
else document.getElementById('c_cust').innerHTML='<div style="padding:20px;color:#64748b;text-align:center;">Tidak ada data customer.</div>';
</script>
</body>
</html>"""
        return html

    # ── Tombol Download ───────────────────────────────────────────────────────
    _, dl_center, _ = st.columns([2, 2, 2])
    with dl_center:
        try:
            html_content = build_overview_html()
            st.download_button(
                label="Download Overview (HTML)",
                data=html_content.encode('utf-8'),
                file_name=f"overview_packaging_guntner_{ov_label.replace(' ','_').replace('/','_')}.html",
                mime="text/html",
                use_container_width=True,
                type="primary"
            )
        except Exception as e:
            st.error(f"Gagal membuat HTML: {e}")

    st.markdown("<br>", unsafe_allow_html=True)



# ═════════════════════════════════════════════════════════════════════════════
# TAB 2 — DISTRIBUTION ANALYSIS
# ═════════════════════════════════════════════════════════════════════════════
with tab2:
    c_title2, c_pt2, c_sel2 = st.columns([3, 1, 1])
    with c_title2:
        st.markdown("## Distribution Analysis")
    with c_pt2:
        dist_view = st.selectbox("Period Type", ['Yearly', 'Monthly'], key='dist_view')
    with c_sel2:
        if dist_view == 'Monthly':
            all_yms2 = sorted(df_raw['YearMonth'].unique(), reverse=True)
            ym_labels2 = [fmt_month_str(ym) for ym in all_yms2]
            ym_map2 = dict(zip(ym_labels2, all_yms2))
            dist_month_label = st.selectbox("Select Month", ym_labels2, key='dist_month')
            dist_ym_sel = ym_map2[dist_month_label]
            dist_year = str(dist_ym_sel[:4])
        else:
            dist_year = st.selectbox("Select Year", ['All'] + [str(y) for y in sorted(year_opts, reverse=True)], key='dist_year')

    if dist_view == 'Monthly':
        df2 = df_raw[df_raw['YearMonth'] == dist_ym_sel].copy()
    else:
        df2 = apply_year_filter(df_raw, 'All' if dist_year == 'All' else dist_year)

    # 2.1 Distribusi Material
    with st.expander("2.1 — Distribusi Material", expanded=True):
        c1, c2, c3, c4 = st.columns(4)
        for col, mat in zip([c1,c2,c3,c4], mat_opts):
            cnt  = int((df2['Material_Label'] == mat).sum())
            pct2 = cnt / len(df2) * 100 if len(df2) else 0
            col.metric(mat, f"{cnt:,}", f"{pct2:.1f}%")

        col_a, col_b = st.columns(2)
        with col_a:
            m2 = df2['Material_Label'].value_counts().reset_index()
            m2.columns = ['Material','Count']
            fig_pie2 = px.pie(m2, values='Count', names='Material',
                color='Material', color_discrete_map=MATERIAL_COLORS,
                title='Proporsi Material')
            fig_pie2.update_layout(paper_bgcolor='white', height=300, margin=dict(l=10,r=10,t=40,b=10))
            st.plotly_chart(fig_pie2, use_container_width=True, config={'displayModeBar':False})

        with col_b:
            # Always show monthly line — filter by year if Yearly mode
            mn_dist = df_raw.groupby(['YearMonth','Material_Label']).size().reset_index(name='n')
            mn_dist['Label'] = mn_dist['YearMonth'].apply(fmt_month_str)
            mn_dist = mn_dist.sort_values('YearMonth')
            if dist_view == 'Yearly' and dist_year != 'All':
                mn_dist = mn_dist[mn_dist['YearMonth'].str.startswith(dist_year)]
                dist_chart_title = f"Monthly Trend {dist_year}"
            elif dist_view == 'Monthly':
                dist_chart_title = f"Monthly Trend — Highlighted: {fmt_month_str(dist_ym_sel)}"
            else:
                dist_chart_title = "Monthly Trend 2023–2026"

            fig_dist_line = go.Figure()
            for mat in mat_opts:
                sub = mn_dist[mn_dist['Material_Label'] == mat]
                fig_dist_line.add_trace(go.Scatter(
                    x=sub['Label'], y=sub['n'], name=mat,
                    mode='lines+markers',
                    line=dict(color=MATERIAL_COLORS.get(mat), width=2),
                    marker=dict(size=4, color=MATERIAL_COLORS.get(mat))
                ))
            if dist_view == 'Monthly':
                sel_lbl2 = fmt_month_str(dist_ym_sel)
                for mat in mat_opts:
                    sub_sel2 = mn_dist[(mn_dist['Material_Label'] == mat) & (mn_dist['Label'] == sel_lbl2)]
                    if not sub_sel2.empty:
                        fig_dist_line.add_trace(go.Scatter(
                            x=sub_sel2['Label'], y=sub_sel2['n'],
                            mode='markers', showlegend=False,
                            marker=dict(size=14, color='#2563eb', symbol='circle', line=dict(color='white', width=2)),
                            hovertemplate=f'<b>{sel_lbl2} (dipilih)</b><br>{mat}: %{{y:,}} lines<extra></extra>'
                        ))
            fig_dist_line.update_layout(title=dist_chart_title,
                paper_bgcolor='white', plot_bgcolor='white', height=300,
                margin=dict(l=20,r=20,t=40,b=55),
                xaxis=dict(showgrid=False, tickangle=-45, tickfont=dict(size=9), title='Month'),
                yaxis=dict(showgrid=True, gridcolor='#f1f5f9', title='Lines'),
                legend=dict(orientation='h',y=1.12,x=1,xanchor='right'),
                hovermode='x unified')
            st.plotly_chart(fig_dist_line, use_container_width=True, config={'displayModeBar':False})
            note2 = f"Titik biru = {fmt_month_str(dist_ym_sel)}" if dist_view == 'Monthly' else ""
            st.markdown(axis_note("Month", "Shipments per month per material", note2), unsafe_allow_html=True)

    # 2.2 Distribusi Dimensi
    with st.expander("2.2 — Distribusi Dimensi (Panjang / Lebar / Tinggi / Volume)", expanded=True):
        st.markdown(
            '<div class="tooltip-box">'
            '<b>Mengapa Dimensi Kritis?</b><br>'
            'Panjang (Length) > 4.000 mm → item umumnya memerlukan palet kayu sebagai penopang struktural.<br>'
            'Lebar & Tinggi besar → item bervolume besar rentan benturan, butuh proteksi kayu.<br>'
            'Volume > 0,5 m³ → hampir selalu menggunakan kayu karena beban distribusi besar; karton tidak mampu menahan.'
            '</div>', unsafe_allow_html=True
        )

        dim_cols   = ['Length', 'Width', 'Height', 'Volume_m3']
        dim_titles = ['Distribusi Panjang', 'Distribusi Lebar', 'Distribusi Tinggi', 'Distribusi Volume']
        dim_xlab   = ['Panjang (mm)', 'Lebar (mm)', 'Tinggi (mm)', 'Volume (m³)']
        dim_desc   = [
            'Semakin panjang item, semakin tinggi kebutuhan palet kayu. Batas kritis: 4.000 mm.',
            'Lebar besar mengindikasikan kebutuhan perlindungan sisi yang hanya bisa dipenuhi kayu.',
            'Tinggi besar → risiko tipping; kayu dibutuhkan untuk stabilisasi dan bracing.',
            'Volume > 0,5 m³ menunjukkan item besar yang hampir pasti memerlukan kayu.'
        ]

        fig_dims = make_subplots(rows=2, cols=2, subplot_titles=dim_titles,
                                  horizontal_spacing=0.1, vertical_spacing=0.18)
        for idx, (dcol, xt) in enumerate(zip(dim_cols, dim_xlab)):
            row, col_ = (idx // 2) + 1, (idx % 2) + 1
            for mat in mat_opts:
                sub = df2[df2['Material_Label'] == mat][dcol]
                sub = sub[sub > 0]
                if sub.empty: continue
                # Compute bin width so we can show range in tooltip
                nbins = 40
                fig_dims.add_trace(go.Histogram(
                    x=sub, name=mat, marker_color=MATERIAL_COLORS[mat],
                    opacity=0.7, nbinsx=nbins, showlegend=(idx == 0),
                    hovertemplate=(
                        f'<b>{mat}</b><br>'
                        'Rentang: %{x}<br>'
                        'Jumlah item: %{y:,}<extra></extra>'
                    )
                ), row=row, col=col_)
            fig_dims.update_xaxes(title_text=xt, showgrid=False, row=row, col=col_)
            fig_dims.update_yaxes(title_text='Frekuensi (jumlah item)', showgrid=True,
                                   gridcolor='#f1f5f9', row=row, col=col_)

        fig_dims.update_layout(barmode='overlay', paper_bgcolor='white', plot_bgcolor='white',
            height=560, margin=dict(l=20,r=20,t=60,b=20),
            font=dict(size=11),
            legend=dict(orientation='h', y=-0.06, x=0.5, xanchor='center'))
        st.plotly_chart(fig_dims, use_container_width=True, config={'displayModeBar':False})

        st.markdown(
            '<div class="axis-note">'
            '<b>Cara membaca grafik di atas:</b> Sumbu X = nilai dimensi (ukuran item), '
            'Sumbu Y = Frekuensi (berapa banyak item yang memiliki ukuran tersebut). '
            'Puncak grafik = ukuran yang paling umum. Warna mewakili material packaging yang digunakan.'
            '</div>', unsafe_allow_html=True
        )

        # Insight per dimensi
        for i, desc in enumerate(dim_desc):
            st.markdown(f'<div style="font-size:12.5px;color:#475569;margin-bottom:4px;">• <b>{dim_titles[i]}:</b> {desc}</div>', unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)

        # Tabel insight — semua 4 material
        st.markdown("**Tabel Detail per Kategori Dimensi**")
        dim_insight_rows = []
        for cat in DIM_LABELS:
            sub_cat = df2[df2['Dim_Category'] == cat]
            if sub_cat.empty: continue
            w_pct   = (sub_cat['Material_Clean']=='W').mean() * 100
            cb_pct  = (sub_cat['Material_Clean']=='B').mean() * 100
            cbd_pct = (sub_cat['Material_Clean']=='Cb').mean() * 100
            ct_pct  = (sub_cat['Material_Clean']=='Ct').mean() * 100
            dim_insight_rows.append({
                'Kategori Ukuran': cat,
                'Jumlah Items': f"{len(sub_cat):,}",
                '% Wood': f"{w_pct:.1f}%",
                '% Carton Box': f"{cb_pct:.1f}%",
                '% Cardboard': f"{cbd_pct:.1f}%",
                '% Carton': f"{ct_pct:.1f}%",
                'Rata-rata Volume (m³)': f"{sub_cat[sub_cat['Volume_m3']>0]['Volume_m3'].mean():.4f}",
                'Rata-rata Berat (kg)': f"{sub_cat[sub_cat['Net Weight']>0]['Net Weight'].mean():.1f}",
            })
        st.dataframe(pd.DataFrame(dim_insight_rows), use_container_width=True, hide_index=True)

        # Visualisasi: grouped bar % semua material per kategori (menggantikan tabel)
        st.markdown("**% Penggunaan Setiap Material per Kategori Ukuran**")
        st.markdown('<div class="axis-note">Dari grafik ini terlihat pola yang jelas: semakin besar ukuran item, semakin mendominasi Wood. Carton Box hanya signifikan di kategori kecil (XS dan S).</div>', unsafe_allow_html=True)

        viz_rows = []
        for cat in DIM_LABELS:
            sub_cat = df2[df2['Dim_Category'] == cat]
            if sub_cat.empty: continue
            total_cat = len(sub_cat)
            for mat, code in [('Wood','W'),('Carton Box','B'),('Cardboard','Cb'),('Carton','Ct')]:
                pct_val = (sub_cat['Material_Clean'] == code).mean() * 100
                viz_rows.append({'Kategori': cat, 'Material': mat, '% Penggunaan': pct_val})
        viz_df = pd.DataFrame(viz_rows)

        fig_viz_cat = go.Figure()
        for mat in mat_opts:
            sub_v = viz_df[viz_df['Material'] == mat]
            # Sort by DIM_LABELS order
            sub_v = sub_v.set_index('Kategori').reindex(DIM_LABELS).reset_index().dropna()
            fig_viz_cat.add_trace(go.Bar(
                name=mat,
                x=sub_v['Kategori'],
                y=sub_v['% Penggunaan'],
                marker_color=MATERIAL_COLORS.get(mat, '#aaa'),
                text=sub_v['% Penggunaan'].apply(lambda v: f'{v:.1f}%' if v > 1 else ''),
                textposition='outside',
                hovertemplate=f'<b>{mat}</b><br>Kategori: %{{x}}<br>%{{y:.1f}}% dari items<extra></extra>'
            ))
        fig_viz_cat.update_layout(
            barmode='group',
            paper_bgcolor='white', plot_bgcolor='white', height=360,
            margin=dict(l=20, r=20, t=20, b=20),
            xaxis=dict(showgrid=False, categoryorder='array', categoryarray=DIM_LABELS),
            yaxis=dict(showgrid=True, gridcolor='#f1f5f9', title='% Penggunaan Material', range=[0, 120]),
            legend=dict(orientation='h', y=1.08, x=1, xanchor='right')
        )
        st.plotly_chart(fig_viz_cat, use_container_width=True, config={'displayModeBar': False})
        st.markdown(axis_note(
            "Kategori Ukuran (XS = terkecil, XL = terbesar)",
            "Persentase item dalam kategori itu yang menggunakan material tertentu",
            "Contoh: bar Wood di kategori XL hampir 100% artinya hampir semua item besar wajib pakai Wood"
        ), unsafe_allow_html=True)

    # 2.3 Distribusi Berat
    with st.expander("2.3 — Distribusi Berat (Net Weight & Gross Weight)", expanded=True):
        st.markdown(
            '<div class="tooltip-box">Berat item berkorelasi kuat dengan pilihan material: '
            'item berat (>50 kg) hampir selalu membutuhkan palet kayu untuk penanganan forklift. '
            'Item ringan (<50 kg) berpotensi diganti ke Carton Box.</div>', unsafe_allow_html=True
        )
        col_e, col_f = st.columns(2)
        with col_e:
            # Histogram berat per material (lebih mudah dipahami dari box plot)
            fig_bw_hist = go.Figure()
            for mat in mat_opts:
                sub = df2[df2['Material_Label'] == mat]['Net Weight']
                sub = sub[(sub > 0) & (sub < 2000)]
                if sub.empty: continue
                fig_bw_hist.add_trace(go.Histogram(
                    x=sub, name=mat, marker_color=MATERIAL_COLORS[mat],
                    opacity=0.7, nbinsx=40,
                    hovertemplate=f'<b>{mat}</b><br>Berat: %{{x:.0f}} kg — Frekuensi: %{{y:,}}<extra></extra>'
                ))
            fig_bw_hist.add_vline(x=50, line_dash='dot', line_color='#dc2626',
                                   annotation_text='Batas 50 kg', annotation_position='top right')
            fig_bw_hist.update_layout(
                barmode='overlay', title='Distribusi Net Weight per Material',
                paper_bgcolor='white', plot_bgcolor='white', height=360,
                margin=dict(l=20,r=20,t=50,b=20),
                xaxis=dict(title='Net Weight (kg)', showgrid=False),
                yaxis=dict(showgrid=True, gridcolor='#f1f5f9', title='Frekuensi (jumlah item)'),
                legend=dict(orientation='h',y=1.12,x=1,xanchor='right'))
            st.plotly_chart(fig_bw_hist, use_container_width=True, config={'displayModeBar':False})
            st.markdown(axis_note("Net Weight (kg) — berat bersih item tanpa kemasan",
                "Frekuensi — berapa banyak item dengan berat tersebut",
                "Garis merah putus-putus = batas 50 kg; di bawahnya berpotensi diganti ke Carton"), unsafe_allow_html=True)

        with col_f:
            scat_wt_src = df2[(df2['Net Weight']>0)&(df2['Gross Wight']>0)&(df2['Net Weight']<3000)]
            scat_wt = scat_wt_src.sample(min(2000, len(scat_wt_src)), random_state=1)
            fig_wg = px.scatter(scat_wt, x='Net Weight', y='Gross Wight',
                color='Material_Label', color_discrete_map=MATERIAL_COLORS, opacity=0.5,
                labels={'Net Weight':'Net Weight (kg)','Gross Wight':'Gross Weight (kg)'},
                title='Net Weight vs Gross Weight per Material')
            fig_wg.update_layout(paper_bgcolor='white', plot_bgcolor='white', height=360,
                margin=dict(l=20,r=20,t=50,b=20),
                legend=dict(orientation='h',y=1.12,x=1,xanchor='right'))
            st.plotly_chart(fig_wg, use_container_width=True, config={'displayModeBar':False})
            st.markdown(axis_note("Net Weight (kg) — berat bersih item",
                "Gross Weight (kg) — berat total termasuk kemasan",
                "Selisih Gross–Net menunjukkan bobot kemasan; Wood packaging lebih berat"), unsafe_allow_html=True)

        wt_stats = df2.groupby('Material_Label')['Net Weight'].agg(
            Jumlah='count', Rata2='mean', Median='median', Std='std', Min='min', Max='max'
        ).round(2).reset_index()
        st.dataframe(wt_stats, use_container_width=True, hide_index=True)

    # 2.4 Top Item per Material
    with st.expander("2.4 — Top Item per Material", expanded=True):
        mat_sel_dd = st.selectbox("Pilih Material:", mat_opts, index=0, key='item_material_dd')
        mat_code_map = {'Wood':'W','Carton Box':'B','Cardboard':'Cb','Carton':'Ct'}
        item_filt = df2[df2['Material_Clean'] == mat_code_map[mat_sel_dd]]
        top_items_dd = (item_filt.groupby('Item').size()
                        .reset_index(name='Count').sort_values('Count', ascending=False).head(20))

        col_g, col_h = st.columns([1.2, 1])
        with col_g:
            fig_item_dd = go.Figure(go.Bar(
                x=top_items_dd['Count'], y=top_items_dd['Item'], orientation='h',
                marker_color=MATERIAL_COLORS.get(mat_sel_dd, '#64748b'),
                text=top_items_dd['Count'].apply(lambda x: f'{x:,}'), textposition='outside',
                hovertemplate='<b>%{y}</b><br>%{x:,} lines<extra></extra>'
            ))
            fig_item_dd.update_layout(title=f'Top 20 Item — {mat_sel_dd}',
                paper_bgcolor='white', plot_bgcolor='white', height=500,
                margin=dict(l=10,r=50,t=50,b=10),
                xaxis=dict(showgrid=True, gridcolor='#f1f5f9', title='Shipments'),
                yaxis=dict(showgrid=False, autorange='reversed', tickfont=dict(size=10)))
            st.plotly_chart(fig_item_dd, use_container_width=True, config={'displayModeBar':False})
            st.markdown(axis_note("Item / Product Unit", f"Shipments using {mat_sel_dd}"), unsafe_allow_html=True)

        with col_h:
            item_stats = (item_filt[item_filt['Item'].isin(top_items_dd['Item'])]
                .groupby('Item').agg(Count=('Item','count'),
                    Avg_L=('Length','mean'), Avg_W=('Width','mean'),
                    Avg_H=('Height','mean'), Avg_Vol=('Volume_m3','mean'),
                    Avg_Wt=('Net Weight','mean')).round(1).reset_index()
                .sort_values('Count',ascending=False))
            st.markdown(f"**Detail Statistik — {mat_sel_dd}**")
            st.dataframe(item_stats.rename(columns={
                'Avg_L':'Avg L (mm)','Avg_W':'Avg W (mm)',
                'Avg_H':'Avg H (mm)','Avg_Vol':'Avg Vol (m³)','Avg_Wt':'Avg Wt (kg)'
            }), use_container_width=True, hide_index=True, height=460)

    # 2.5 Customer
    with st.expander("2.5 — Distribusi per Customer (Top 15)", expanded=True):
        top_cust = (df2.groupby('Name').size().reset_index(name='Count')
                    .sort_values('Count', ascending=False).head(15))
        fig_cust = go.Figure(go.Bar(
            x=top_cust['Count'], y=top_cust['Name'], orientation='h',
            marker=dict(color=top_cust['Count'],
                        colorscale=[[0,'#dbeafe'],[1,'#1d4ed8']], showscale=False),
            text=top_cust['Count'].apply(lambda x: f'{x:,}'), textposition='outside',
            hovertemplate='<b>%{y}</b><br>%{x:,} lines<extra></extra>'
        ))
        fig_cust.update_layout(title='Top 15 Customer berdasarkan Volume Shipment',
            paper_bgcolor='white', plot_bgcolor='white', height=460,
            margin=dict(l=20,r=50,t=50,b=20),
            xaxis=dict(showgrid=True, gridcolor='#f1f5f9', title='Shipments'),
            yaxis=dict(showgrid=False, autorange='reversed', tickfont=dict(size=10.5)))
        st.plotly_chart(fig_cust, use_container_width=True, config={'displayModeBar':False})
        st.markdown(axis_note("Shipments (total from that customer)", "Customer Name"), unsafe_allow_html=True)

    # 2.6 Temporal
    with st.expander("2.6 — Distribusi Temporal", expanded=True):
        if dist_view == 'Monthly':
            monthly_mat2 = (df_raw.groupby(['YearMonth','Material_Label']).size().reset_index(name='Count'))
            monthly_mat2['Label'] = monthly_mat2['YearMonth'].apply(fmt_month_str)
            monthly_mat2 = monthly_mat2.sort_values('YearMonth')
            fig_line2 = go.Figure()
            for mat in mat_opts:
                sub = monthly_mat2[monthly_mat2['Material_Label'] == mat]
                fig_line2.add_trace(go.Scatter(x=sub['Label'], y=sub['Count'], name=mat,
                    mode='lines+markers', line=dict(color=MATERIAL_COLORS[mat], width=2.5),
                    marker=dict(size=5)))
            sel_lbl_26 = fmt_month_str(dist_ym_sel)
            # Highlight selected month with big dot markers
            for mat in mat_opts:
                sub_sel26 = monthly_mat2[(monthly_mat2['Material_Label'] == mat) & (monthly_mat2['Label'] == sel_lbl_26)]
                if not sub_sel26.empty:
                    fig_line2.add_trace(go.Scatter(
                        x=sub_sel26['Label'], y=sub_sel26['Count'],
                        mode='markers', showlegend=False,
                        marker=dict(size=14, color='#2563eb', symbol='circle', line=dict(color='white', width=2)),
                        hovertemplate=f'<b>{sel_lbl_26} (dipilih)</b><br>{mat}: %{{y:,}} lines<extra></extra>'
                    ))
            fig_line2.update_layout(title=f'Monthly Trend — Bulan dipilih: {sel_lbl_26}',
                paper_bgcolor='white', plot_bgcolor='white', height=380,
                margin=dict(l=20,r=20,t=50,b=60),
                xaxis=dict(showgrid=False, tickangle=-45, tickfont=dict(size=9), title='Month'),
                yaxis=dict(showgrid=True, gridcolor='#f1f5f9', title='Shipments'),
                legend=dict(orientation='h',y=1.12,x=1,xanchor='right'))
            st.plotly_chart(fig_line2, use_container_width=True, config={'displayModeBar':False})
            st.markdown(axis_note("Month", "Shipments", f"Garis biru = {sel_lbl_26}"), unsafe_allow_html=True)
        else:
            yr_mat2 = df_raw.groupby(['Year','Material_Label']).size().reset_index(name='Count')
            if dist_year != 'All':
                yr_mat2 = yr_mat2[yr_mat2['Year'] == int(dist_year)]
            fig_yr2 = go.Figure()
            for mat in mat_opts:
                sub = yr_mat2[yr_mat2['Material_Label'] == mat]
                fig_yr2.add_trace(go.Scatter(x=sub['Year'].astype(str), y=sub['Count'], name=mat,
                    mode='lines+markers', line=dict(color=MATERIAL_COLORS[mat], width=2.5),
                    marker=dict(size=7)))
            fig_yr2.update_layout(title='Yearly Trend — All Materials',
                paper_bgcolor='white', plot_bgcolor='white', height=380,
                margin=dict(l=20,r=20,t=50,b=20),
                xaxis=dict(showgrid=False, title='Year'),
                yaxis=dict(showgrid=True, gridcolor='#f1f5f9', title='Shipments'),
                legend=dict(orientation='h',y=1.12,x=1,xanchor='right'))
            st.plotly_chart(fig_yr2, use_container_width=True, config={'displayModeBar':False})
            st.markdown(axis_note("Year", "Shipments"), unsafe_allow_html=True)


# ═════════════════════════════════════════════════════════════════════════════
# TAB 3 — DEEP DIVE ANALYSIS
# ═════════════════════════════════════════════════════════════════════════════
with tab3:
    c_t3, c_pt3, c_sel3 = st.columns([3, 1, 1])
    with c_t3:
        st.markdown("## Deep Dive Analysis")
    with c_pt3:
        dd_view = st.selectbox("Period Type", ['Yearly', 'Monthly'], key='dd_view')
    with c_sel3:
        if dd_view == 'Monthly':
            all_yms3 = sorted(df_raw['YearMonth'].unique(), reverse=True)
            ym_labels3 = [fmt_month_str(ym) for ym in all_yms3]
            ym_map3 = dict(zip(ym_labels3, all_yms3))
            dd_month_label = st.selectbox("Select Month", ym_labels3, key='dd_month')
            dd_ym_sel = ym_map3[dd_month_label]
            dd_year = str(dd_ym_sel[:4])
        else:
            dd_year = st.selectbox("Select Year", ['All'] + [str(y) for y in sorted(year_opts, reverse=True)], key='dd_year')

    if dd_view == 'Monthly':
        df3 = df_raw[df_raw['YearMonth'] == dd_ym_sel].copy()
    else:
        df3 = apply_year_filter(df_raw, 'All' if dd_year == 'All' else dd_year)
    wood_n3 = int((df3['Is_Wood']).sum())

    # 3.1 Dimensi vs Material — analisis detail
    with st.expander("3.1 — Analisis Dimensi & Berat → Pilihan Material (Ketetapan Umum)", expanded=True):
        st.markdown(
            '<div class="tooltip-box">'
            '<b>Ketetapan Umum Pilihan Packaging:</b><br>'
            'Analisis ini menetapkan aturan berbasis data tentang kapan Wood diperlukan vs kapan Carton Box cukup. '
            'Tujuannya adalah memberikan panduan yang dapat dioperasionalkan oleh tim packaging.'
            '</div>', unsafe_allow_html=True
        )

        # Scatter: Max Dimension vs Net Weight, warna = material
        scatter3 = df3[(df3['Max_Dimension']>0) & (df3['Net Weight']>0)].copy()
        scatter3 = scatter3.sample(min(4000, len(scatter3)), random_state=42)
        fig_sc3 = px.scatter(scatter3, x='Max_Dimension', y='Net Weight',
            color='Material_Label', color_discrete_map=MATERIAL_COLORS,
            opacity=0.55, size_max=10,
            labels={'Max_Dimension':'Dimensi Terbesar (mm)','Net Weight':'Berat Bersih (kg)',
                    'Material_Label':'Material'},
            title='Hubungan Dimensi & Berat terhadap Pilihan Material')
        fig_sc3.add_vline(x=1500, line_dash='dash', line_color='#d97706',
            annotation_text='<1.500 mm → Karton Feasible', annotation_position='top left',
            annotation_font_size=11)
        fig_sc3.add_vline(x=4000, line_dash='dash', line_color='#dc2626',
            annotation_text='>4.000 mm → Wajib Wood', annotation_position='top right',
            annotation_font_size=11)
        fig_sc3.add_hline(y=50, line_dash='dash', line_color='#2563eb',
            annotation_text='<50 kg → Carton feasible', annotation_position='right',
            annotation_font_size=11)
        fig_sc3.update_layout(paper_bgcolor='white', plot_bgcolor='white', height=430,
            margin=dict(l=20,r=20,t=50,b=20))
        st.plotly_chart(fig_sc3, use_container_width=True, config={'displayModeBar':False})
        st.markdown(axis_note(
            "Dimensi Terbesar (mm) — nilai terbesar antara Panjang, Lebar, dan Tinggi item",
            "Berat Bersih / Net Weight (kg) — berat item tanpa kemasan",
            "Setiap titik = 1 shipment. Warna = material yang digunakan"
        ), unsafe_allow_html=True)

        st.markdown(
            '<div class="critical-box">'
            '<b>Ketetapan Umum Berbasis Data:</b><br>'
            '🔴 <b>Wajib Wood:</b> Dimensi terbesar > 4.000 mm ATAU berat > 200 kg<br>'
            '🟡 <b>Gray Area:</b> Dimensi 1.500–4.000 mm DAN berat 50–200 kg → perlu evaluasi per kasus<br>'
            '🟢 <b>Bisa diganti Carton Box:</b> Semua dimensi < 1.500 mm DAN berat < 50 kg'
            '</div>', unsafe_allow_html=True
        )

        # Bar chart % Wood per segment
        st.markdown("**Persentase Wood berdasarkan Kombinasi Dimensi & Berat**")
        seg_rows = []
        segs = [
            ('Kecil Ringan\n(<1.5m & <50kg)',    (df3['Max_Dimension'] < 1500) & (df3['Net Weight'] < 50)   & (df3['Net Weight']>0)),
            ('Kecil Berat\n(<1.5m & ≥50kg)',     (df3['Max_Dimension'] < 1500) & (df3['Net Weight'] >= 50)  & (df3['Net Weight']>0)),
            ('Sedang Ringan\n(1.5–4m & <50kg)',  (df3['Max_Dimension'].between(1500,4000)) & (df3['Net Weight'] < 50)  & (df3['Net Weight']>0)),
            ('Sedang Berat\n(1.5–4m & ≥50kg)',  (df3['Max_Dimension'].between(1500,4000)) & (df3['Net Weight'] >= 50) & (df3['Net Weight']>0)),
            ('Besar\n(>4m)',                     df3['Max_Dimension'] > 4000),
        ]
        for label, mask in segs:
            sub_s = df3[mask]
            if sub_s.empty: continue
            w_pct_s = sub_s['Is_Wood'].mean() * 100
            seg_rows.append({'Segmen': label, '% Wood': w_pct_s, 'Jumlah Items': len(sub_s)})
        seg_df = pd.DataFrame(seg_rows)

        fig_seg = go.Figure(go.Bar(
            x=seg_df['Segmen'], y=seg_df['% Wood'],
            marker=dict(color=seg_df['% Wood'],
                        colorscale=[[0,'#FEF9F5'],[0.5,'#E8A838'],[1,'#6B2D0E']], showscale=False),
            text=seg_df.apply(lambda r: f"{r['% Wood']:.1f}%\n({r['Jumlah Items']:,} items)", axis=1),
            textposition='outside',
            hovertemplate='<b>%{x}</b><br>% Wood: %{y:.1f}%<extra></extra>'
        ))
        fig_seg.update_layout(paper_bgcolor='white', plot_bgcolor='white', height=360,
            margin=dict(l=20,r=20,t=20,b=20),
            xaxis=dict(showgrid=False),
            yaxis=dict(showgrid=True, gridcolor='#f1f5f9', title='% Penggunaan Wood', range=[0,115]))
        st.plotly_chart(fig_seg, use_container_width=True, config={'displayModeBar':False})
        st.markdown(axis_note(
            "Segmen item berdasarkan kombinasi ukuran & berat",
            "Persentase item dalam segmen tersebut yang menggunakan Wood",
            "Angka dalam kurung = jumlah item di segmen tersebut"
        ), unsafe_allow_html=True)

        # Tabel — semua 4 material
        st.markdown("**Tabel Detail per Kategori Dimensi**")
        dim_rows3 = []
        for cat in DIM_LABELS:
            sub_cat = df3[df3['Dim_Category'] == cat]
            if sub_cat.empty: continue
            w_pct3   = sub_cat['Is_Wood'].mean() * 100
            cb_pct3  = (sub_cat['Material_Clean']=='B').mean() * 100
            cbd_pct3 = (sub_cat['Material_Clean']=='Cb').mean() * 100
            ct_pct3  = (sub_cat['Material_Clean']=='Ct').mean() * 100
            avg_wt3  = sub_cat[sub_cat['Net Weight']>0]['Net Weight'].mean()
            avg_vol3 = sub_cat[sub_cat['Volume_m3']>0]['Volume_m3'].mean()
            dim_rows3.append({
                'Kategori': cat,
                'Jumlah Items': f"{len(sub_cat):,}",
                '% Wood': f"{w_pct3:.1f}%",
                '% Carton Box': f"{cb_pct3:.1f}%",
                '% Cardboard': f"{cbd_pct3:.1f}%",
                '% Carton': f"{ct_pct3:.1f}%",
                'Rata-rata Berat (kg)': f"{avg_wt3:.1f}",
                'Rata-rata Volume (m³)': f"{avg_vol3:.4f}",
                'Rekomendasi': ('Wajib Wood' if w_pct3 > 90 else 'Evaluasi per kasus' if w_pct3 > 40 else 'Pertimbangkan Carton Box')
            })
        st.dataframe(pd.DataFrame(dim_rows3), use_container_width=True, hide_index=True)

    # 3.2 Overpackaged
    with st.expander("3.2 — Peluang Penggantian Wood ke Material Alternatif", expanded=True):
        overpack_df = df3[df3['Overpack_Candidate']].copy()
        overpack_n  = len(overpack_df)
        overpack_pct = overpack_n / wood_n3 * 100 if wood_n3 else 0

        o1, o2, o3 = st.columns(3)
        o1.metric("Item Wood 'Overpackaged'", f"{overpack_n:,}", "Kecil (<1.5m & <50kg) tapi pakai kayu")
        o2.metric("% dari Total Wood", f"{overpack_pct:.1f}%")
        o3.metric("Potensi Penggantian", f"{overpack_n:,} lines", "Kandidat diganti ke Carton Box")

        st.markdown(
            f'<div class="insight-box">Terdapat <b>{overpack_n:,} lines ({overpack_pct:.1f}%)</b> '
            f'dari total Wood yang dimensi kecil (<1.500 mm) dan berat ringan (<50 kg). '
            'Item-item ini berpotensi diganti dengan Carton Box atau Cardboard — '
            'lebih ramah lingkungan dan lebih hemat biaya.</div>', unsafe_allow_html=True
        )

        col_k, col_l = st.columns(2)
        with col_k:
            wood_all = df3[df3['Is_Wood'] & (df3['Max_Dimension'] > 0)]
            fig_wood_dist = go.Figure()
            fig_wood_dist.add_trace(go.Histogram(
                x=wood_all[~wood_all['Overpack_Candidate']]['Max_Dimension'],
                name='Wood — Normal', marker_color='#7C4A19', opacity=0.8, nbinsx=50))
            fig_wood_dist.add_trace(go.Histogram(
                x=wood_all[wood_all['Overpack_Candidate']]['Max_Dimension'],
                name='Wood — Overpackaged', marker_color='#ef4444', opacity=0.8, nbinsx=50))
            fig_wood_dist.add_vline(x=1500, line_dash='dot', line_color='#dc2626',
                                     annotation_text='Batas 1.5 m')
            fig_wood_dist.update_layout(barmode='overlay',
                title='Distribusi Dimensi — Wood Normal vs Overpackaged',
                paper_bgcolor='white', plot_bgcolor='white', height=360,
                margin=dict(l=20,r=20,t=50,b=20),
                xaxis=dict(title='Dimensi Terbesar (mm)', showgrid=False),
                yaxis=dict(showgrid=True, gridcolor='#f1f5f9', title='Frekuensi (jumlah item)'),
                legend=dict(orientation='h',y=1.12,x=1,xanchor='right'))
            st.plotly_chart(fig_wood_dist, use_container_width=True, config={'displayModeBar':False})
            st.markdown(axis_note("Dimensi Terbesar (mm)", "Frekuensi (jumlah item)",
                "Merah = item Wood yang ukurannya kecil dan berpotensi diganti ke Carton"), unsafe_allow_html=True)

        with col_l:
            if not overpack_df.empty:
                top_over = (overpack_df.groupby('Item').size()
                            .reset_index(name='Count').sort_values('Count', ascending=False).head(15))
                fig_over = go.Figure(go.Bar(
                    x=top_over['Count'], y=top_over['Item'], orientation='h',
                    marker_color='#ef4444', opacity=0.85,
                    text=top_over['Count'], textposition='outside',
                    hovertemplate='<b>%{y}</b><br>%{x} lines<extra></extra>'
                ))
                fig_over.update_layout(title='Top 15 Item Overpackaged (Wood → bisa diganti)',
                    paper_bgcolor='white', plot_bgcolor='white', height=360,
                    margin=dict(l=10,r=40,t=50,b=10),
                    xaxis=dict(showgrid=True, gridcolor='#f1f5f9', title='Shipments'),
                    yaxis=dict(showgrid=False, autorange='reversed', tickfont=dict(size=10)))
                st.plotly_chart(fig_over, use_container_width=True, config={'displayModeBar':False})
                st.markdown(axis_note("Item / Product Unit", "Jumlah Lines overpackaged"), unsafe_allow_html=True)

    # 3.3 Customer
    with st.expander("3.3 — Pola Customer × Material", expanded=True):
        top15_cust = (df3.groupby('Name').size().reset_index(name='n')
                      .sort_values('n', ascending=False).head(15)['Name'].tolist())
        cust_mat_piv = (df3[df3['Name'].isin(top15_cust)]
                        .groupby(['Name','Material_Label']).size().unstack(fill_value=0))
        for mat in mat_opts:
            if mat not in cust_mat_piv.columns: cust_mat_piv[mat] = 0
        cust_mat_piv = cust_mat_piv[mat_opts]

        cust_wood_pct = (df3[df3['Name'].isin(top15_cust)].groupby('Name')
                         .apply(lambda x: x['Is_Wood'].mean() * 100)
                         .reset_index(name='Wood_Pct').sort_values('Wood_Pct', ascending=False))
        fig_wcust = go.Figure(go.Bar(
            x=cust_wood_pct['Name'], y=cust_wood_pct['Wood_Pct'],
            marker=dict(color=cust_wood_pct['Wood_Pct'],
                        colorscale=[[0,'#FEF9F5'],[1,'#7C4A19']], showscale=False),
            text=cust_wood_pct['Wood_Pct'].apply(lambda x: f'{x:.1f}%'), textposition='outside',
            hovertemplate='<b>%{x}</b><br>%{y:.1f}% menggunakan Wood<extra></extra>'
        ))
        fig_wcust.update_layout(title='% Penggunaan Wood per Customer (Top 15)',
            paper_bgcolor='white', plot_bgcolor='white', height=360,
            margin=dict(l=20,r=20,t=50,b=80),
            xaxis=dict(showgrid=False, tickangle=-30, tickfont=dict(size=9.5), title='Customer'),
            yaxis=dict(showgrid=True, gridcolor='#f1f5f9', title='% Wood', range=[0,115]))
        st.plotly_chart(fig_wcust, use_container_width=True, config={'displayModeBar':False})
        st.markdown(axis_note("Nama Customer", "Persentase shipment customer tersebut yang menggunakan Wood",
            "Customer dengan % Wood tinggi adalah prioritas kolaborasi untuk Pledge 6"), unsafe_allow_html=True)

        st.markdown('<div class="insight-box">Customer dengan % Wood tinggi adalah kandidat utama '
            'untuk program daur ulang kemasan (Pledge 6).</div>', unsafe_allow_html=True)


# ═════════════════════════════════════════════════════════════════════════════
# TAB 4 — PREDICTION
# ═════════════════════════════════════════════════════════════════════════════
with tab4:
    c_t4, c_pt4, c_sel4 = st.columns([3, 1, 1])
    with c_t4:
        st.markdown("## Prediction Analysis — XGBoost Forecasting")
    with c_pt4:
        pred_view = st.selectbox("Period Type", ['Yearly', 'Monthly'], key='pred_view')
    with c_sel4:
        if pred_view == 'Monthly':
            all_yms4 = sorted(df_raw['YearMonth'].unique(), reverse=True)
            ym_labels4 = [fmt_month_str(ym) for ym in all_yms4]
            ym_map4 = dict(zip(ym_labels4, all_yms4))
            pred_month_label = st.selectbox("Select Month", ym_labels4, key='pred_month')
            pred_ym_sel = ym_map4[pred_month_label]
            pred_year = str(pred_ym_sel[:4])
        else:
            pred_year = st.selectbox("Select Year", ['All'] + [str(y) for y in sorted(year_opts, reverse=True)], key='pred_year')

    st.markdown(
        '<div class="tooltip-box"><b>Algoritma: XGBoost</b> — Dipilih karena efektif untuk time series pendek, '
        'menangkap pola musiman, dan robust terhadap outlier. '
        'Fitur: lag 1–12 bulan, rolling mean, komponen Fourier (sin/cos bulanan).</div>',
        unsafe_allow_html=True
    )

    # Historical
    monthly_all = df_raw.groupby(['YearMonth','Material_Label']).size().reset_index(name='Count')
    monthly_all['Label'] = monthly_all['YearMonth'].apply(fmt_month_str)
    monthly_all = monthly_all.sort_values('YearMonth')
    if pred_view == 'Monthly':
        monthly_all = monthly_all[monthly_all['YearMonth'] == pred_ym_sel]
    elif pred_year != 'All':
        monthly_all = monthly_all[monthly_all['YearMonth'].str.startswith(pred_year)]

    fig_hist_all = go.Figure()
    for mat in mat_opts:
        sub = monthly_all[monthly_all['Material_Label'] == mat]
        fig_hist_all.add_trace(go.Scatter(x=sub['Label'], y=sub['Count'], name=mat,
            mode='lines+markers', line=dict(color=MATERIAL_COLORS[mat], width=2.5),
            marker=dict(size=5)))
    fig_hist_all.update_layout(title='Historical Usage — All Materials',
        paper_bgcolor='white', plot_bgcolor='white', height=360,
        margin=dict(l=20,r=20,t=30,b=60),
        xaxis=dict(showgrid=False, tickangle=-45, tickfont=dict(size=9), title='Month'),
        yaxis=dict(showgrid=True, gridcolor='#f1f5f9', title='Shipments'),
        legend=dict(orientation='h', y=1.05, x=1, xanchor='right'),
        hovermode='x unified')
    st.plotly_chart(fig_hist_all, use_container_width=True, config={'displayModeBar':False})
    st.markdown(axis_note("Month (chronological)", "Shipments per month per material"), unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("### 12-Month Forecast")

    with st.spinner("Menghitung prediksi XGBoost…"):
        fc_results = {}
        for mat in mat_opts:
            fc_results[mat] = forecast_xgboost(mat, periods=12)

    fc_row1 = st.columns(2)
    fc_row2 = st.columns(2)
    fc_cols = [fc_row1[0], fc_row1[1], fc_row2[0], fc_row2[1]]

    # Warna prediksi per material (berbeda dari warna historis)
    PRED_COLORS = {
        'Wood':       '#f59e0b',   # amber
        'Carton Box': '#3b82f6',   # blue
        'Cardboard':  '#8b5cf6',   # purple
        'Carton':     '#10b981',   # emerald
    }

    for col_fc, mat in zip(fc_cols, mat_opts):
        hist_df, pred_df, pct, model_obj = fc_results[mat]
        with col_fc:
            if hist_df is None:
                st.warning(f"Data tidak cukup untuk forecast {mat}.")
                continue
            badge_html = trend_badge(pct, "vs avg historis")
            direction  = "Diprediksi NAIK" if pct > 0 else "Diprediksi TURUN"
            d_color    = "#dc2626" if (mat == 'Wood' and pct > 0) else "#16a34a"
            st.markdown(
                f'<div style="background:{MATERIAL_LIGHT.get(mat,"#f8fafc")};'
                f'border-radius:12px;padding:14px;margin-bottom:6px;">'
                f'<div style="font-size:14px;font-weight:700;color:{MATERIAL_COLORS.get(mat,"#333")};">{mat}</div>'
                f'<div style="margin-top:5px;">{badge_html}</div>'
                f'<div style="font-size:12px;color:{d_color};font-weight:600;margin-top:3px;">{direction}</div>'
                f'</div>', unsafe_allow_html=True
            )
            base_c = MATERIAL_COLORS.get(mat, '#64748b')
            pred_c = PRED_COLORS.get(mat, '#f59e0b')

            # ── Sambungkan ujung historis ke awal prediksi (bridge point) ──────
            last_hist_ym    = hist_df['YearMonth'].iloc[-1]
            last_hist_count = float(hist_df['Count'].iloc[-1])
            bridge_x = [fmt_month_str(last_hist_ym)] + pred_df['YearMonth'].apply(fmt_month_str).tolist()
            bridge_y = [last_hist_count] + pred_df['Forecast'].tolist()

            x_hist = hist_df['YearMonth'].apply(fmt_month_str).tolist()
            avg_v  = float(hist_df['Count'].mean())

            fig_fc = go.Figure()

            # Historis — garis solid penuh
            fig_fc.add_trace(go.Scatter(
                x=x_hist, y=hist_df['Count'].tolist(),
                name='Aktual (Historis)',
                mode='lines+markers',
                line=dict(color=base_c, width=2.5),
                marker=dict(size=4, color=base_c),
                hovertemplate='<b>%{x}</b><br>Aktual: %{y:,}<extra></extra>'
            ))

            # Prediksi — disambung dari titik terakhir historis, garis putus-putus
            fig_fc.add_trace(go.Scatter(
                x=bridge_x, y=bridge_y,
                name='Prediksi (12 bln)',
                mode='lines+markers',
                line=dict(color=pred_c, width=2.5, dash='dot'),
                marker=dict(size=6, color=pred_c, symbol='diamond'),
                hovertemplate='<b>%{x}</b><br>Prediksi: %{y:,.0f}<extra></extra>'
            ))

            # Garis rata-rata historis
            fig_fc.add_hline(y=avg_v, line_dash='dash', line_color='#94a3b8',
                             annotation_text=f'Avg historis {avg_v:.0f}',
                             annotation_font_size=9, annotation_font_color='#94a3b8')

            # Vertical separator: pakai shapes dict (add_vline tidak support categorical x)
            split_label = fmt_month_str(last_hist_ym)
            all_x = x_hist + bridge_x[1:]
            split_pos = all_x.index(split_label) if split_label in all_x else len(x_hist) - 1

            fig_fc.update_layout(
                paper_bgcolor='white', plot_bgcolor='white', height=300,
                margin=dict(l=10, r=10, t=20, b=55),
                xaxis=dict(showgrid=False, tickangle=-45, tickfont=dict(size=8), title='Month'),
                yaxis=dict(showgrid=True, gridcolor='#f1f5f9', title='Lines', tickfont=dict(size=9)),
                legend=dict(orientation='h', y=1.10, x=1, xanchor='right', font=dict(size=9)),
                hovermode='x unified',
                shapes=[dict(
                    type='line', xref='x', yref='paper',
                    x0=split_pos, x1=split_pos, y0=0, y1=1,
                    line=dict(color='#64748b', width=1.2, dash='dot')
                )],
                annotations=[dict(
                    xref='x', yref='paper',
                    x=split_pos + 0.3, y=0.97,
                    text='← Historis  |  Prediksi →',
                    showarrow=False,
                    font=dict(size=8, color='#64748b'),
                    xanchor='left',
                    bgcolor='rgba(255,255,255,0.8)',
                    bordercolor='#e2e8f0', borderwidth=1
                )]
            )

    st.markdown(
        '<div class="axis-note">'
        '📊 <b>Cara membaca grafik:</b> '
        'Garis solid = data historis aktual dari Excel. '
        'Garis putus-putus (warna berbeda) = prediksi XGBoost 12 bulan ke depan. '
        'Kedua garis tersambung di titik bulan terakhir data historis. '
        'Garis abu-abu = rata-rata historis sebagai baseline perbandingan.'
        '</div>', unsafe_allow_html=True
    )

    # ── Tabel Prediksi dengan 3 Bulan Historis Sebelumnya ────────────────────
    st.markdown("---")
    st.markdown("### 12-Month Forecast Detail")
    st.markdown(
        '<div class="tooltip-box">'
        'Tabel di bawah menampilkan <b>3 bulan historis terakhir</b> (sebagai konteks) '
        'dilanjutkan dengan <b>12 bulan prediksi ke depan</b>. '
        'Baris historis ditandai dengan latar abu-abu.'
        '</div>', unsafe_allow_html=True
    )

    # Ambil 3 bulan historis terakhir per material
    hist_lookback_rows = []
    for mat in mat_opts:
        hist_df_lb, _, _, _ = fc_results[mat]
        if hist_df_lb is None: continue
        last3 = hist_df_lb.tail(3)
        for _, row in last3.iterrows():
            hist_lookback_rows.append({
                'Tipe': '📋 Historis',
                'Material': mat,
                'YearMonth': row['YearMonth'],
                'Month': fmt_month_str(row['YearMonth']),
                'Lines': int(row['Count']),
            })

    pred_all_rows = []
    for mat in mat_opts:
        _, pred_df_p, _, _ = fc_results[mat]
        if pred_df_p is None: continue
        for _, row in pred_df_p.iterrows():
            pred_all_rows.append({
                'Tipe': '🔮 Prediksi',
                'Material': mat,
                'YearMonth': row['YearMonth'],
                'Month': fmt_month_str(row['YearMonth']),
                'Lines': int(round(row['Forecast'])),
            })

    all_table_rows = hist_lookback_rows + pred_all_rows

    if all_table_rows:
        combined_df = pd.DataFrame(all_table_rows)
        # Pivot per material
        pivot_comb = combined_df.pivot_table(
            index=['Tipe', 'YearMonth', 'Month'],
            columns='Material', values='Lines', aggfunc='sum'
        ).reset_index()
        pivot_comb = pivot_comb.sort_values('YearMonth').reset_index(drop=True)

        for mat in mat_opts:
            if mat not in pivot_comb.columns:
                pivot_comb[mat] = 0
        for mat in mat_opts:
            pivot_comb[mat] = pivot_comb[mat].fillna(0).astype(int)

        pivot_comb['Total'] = pivot_comb[[m for m in mat_opts if m in pivot_comb.columns]].sum(axis=1).astype(int)
        pivot_comb = pivot_comb.drop(columns='YearMonth')

        col_order = ['Tipe', 'Month'] + [m for m in mat_opts if m in pivot_comb.columns] + ['Total']
        pivot_comb = pivot_comb[col_order]

        # Rename kolom Tipe supaya lebih jelas
        pivot_comb = pivot_comb.rename(columns={'Month': 'Bulan', 'Tipe': 'Status'})

        st.dataframe(
            pivot_comb,
            use_container_width=True,
            hide_index=True,
            column_config={
                'Status': st.column_config.TextColumn('Status', width='medium'),
                'Bulan': st.column_config.TextColumn('Bulan', width='medium'),
            }
        )

        # Download hanya baris prediksi
        pred_only = pivot_comb[pivot_comb['Status'] == '🔮 Prediksi'].copy()
        st.download_button("Download Tabel Prediksi (CSV)",
            pred_only.to_csv(index=False).encode('utf-8'),
            'forecast_packaging.csv', 'text/csv')

    # Interpretasi
    st.markdown("---")
    st.markdown("### Interpretasi Prediksi")
    interp_cols = st.columns(2)
    for idx, mat in enumerate(mat_opts):
        _, _, pct_i, _ = fc_results[mat]
        if pct_i is None: continue
        with interp_cols[idx % 2]:
            if mat == 'Wood':
                if pct_i > 5:
                    msg, note, box = f"Penggunaan kayu diprediksi <b>naik {pct_i:.1f}%</b>.", "Perlu tindakan segera — beban lingkungan meningkat.", "critical-box"
                elif pct_i < -5:
                    msg, note, box = f"Penggunaan kayu diprediksi <b>turun {abs(pct_i):.1f}%</b>.", "Tren positif — mengarah ke Pledge 6.", "tooltip-box"
                else:
                    msg, note, box = f"Penggunaan kayu relatif <b>stabil ({pct_i:+.1f}%)</b>.", "Masih perlu intervensi untuk mencapai Pledge 6.", "insight-box"
            else:
                if pct_i > 0:
                    msg, note, box = f"Penggunaan {mat} diprediksi <b>naik {pct_i:.1f}%</b>.", "Tren positif — substitusi material lebih berkelanjutan.", "tooltip-box"
                else:
                    msg, note, box = f"Penggunaan {mat} diprediksi <b>turun {abs(pct_i):.1f}%</b>.", "Pantau penyebab penurunan.", "insight-box"
            st.markdown(f'<div class="{box}"><b>{mat}</b><br>{msg}<br><br>{note}</div>', unsafe_allow_html=True)


# ═════════════════════════════════════════════════════════════════════════════
# TAB 5 — RECOMMENDATION (GEMINI AI)
# ═════════════════════════════════════════════════════════════════════════════
with tab5:
    c_t5, c_pt5, c_sel5 = st.columns([3, 1, 1])
    with c_t5:
        st.markdown("## AI Recommendation — Powered by Gemini")
    with c_pt5:
        rec_view = st.selectbox("Period Type", ['Yearly', 'Monthly'], key='rec_view')
    with c_sel5:
        if rec_view == 'Monthly':
            all_yms5 = sorted(df_raw['YearMonth'].unique(), reverse=True)
            ym_labels5 = [fmt_month_str(ym) for ym in all_yms5]
            ym_map5 = dict(zip(ym_labels5, all_yms5))
            rec_month_label = st.selectbox("Select Month", ym_labels5, key='rec_month')
            rec_ym_sel = ym_map5[rec_month_label]
            rec_year_sel = str(rec_ym_sel[:4])
        else:
            rec_year_sel = st.selectbox("Select Year", ['All'] + [str(y) for y in sorted(year_opts, reverse=True)], key='rec_year')

    st.markdown(
        '<div class="tooltip-box">Gunakan AI untuk mendapatkan rekomendasi konkret berbasis data aktual packaging PT Güntner. '
        'Rekomendasi mencakup 3 tingkatan: <b>Quick Win (0–3 bulan)</b>, <b>Strategic (3–12 bulan)</b>, dan <b>Transformational (12+ bulan)</b>.</div>',
        unsafe_allow_html=True
    )

    inp1, inp2, inp3 = st.columns(3)
    with inp1:
        focus_mat_rec = st.selectbox("Material Fokus",
            ['All Materials', 'Wood', 'Carton Box', 'Cardboard', 'Carton'], index=0, key='rec_mat')
    with inp2:
        sustainability_goal = st.selectbox("Target Sustainability",
            ['Reduce Wood Usage', 'Circular Packaging', 'Zero-Waste Packaging', 'Customer Collaboration'],
            index=0, key='rec_goal')
    with inp3:
        context_extra = st.text_input("Konteks Tambahan (opsional)",
            placeholder="Contoh: target zero-wood 2028…", key='rec_context')

    # Compute stats
    rec_df = df_raw.copy()
    if rec_view == 'Monthly':
        rec_df = rec_df[rec_df['YearMonth'] == rec_ym_sel]
    elif rec_year_sel != 'All':
        rec_df = rec_df[rec_df['Year'] == int(rec_year_sel)]

    rec_total   = len(rec_df)
    rec_wood    = int((rec_df['Material_Clean'] == 'W').sum())
    rec_woodpct = rec_wood / rec_total * 100 if rec_total else 0
    rec_cb      = int((rec_df['Material_Clean'] == 'B').sum())
    rec_cbpct   = rec_cb / rec_total * 100 if rec_total else 0
    rec_cbd     = int((rec_df['Material_Clean'] == 'Cb').sum())
    rec_ct      = int((rec_df['Material_Clean'] == 'Ct').sum())
    avg_wood_dim_rec    = rec_df[rec_df['Material_Clean']=='W']['Max_Dimension'].mean()
    avg_nonwood_dim_rec = rec_df[rec_df['Material_Clean']!='W']['Max_Dimension'].mean()
    top_wood_items_rec  = rec_df[rec_df['Material_Clean']=='W']['Unit'].value_counts().head(5).to_dict()
    yearly_wood_rec     = rec_df[rec_df['Material_Clean']=='W'].groupby('Year').size().to_dict()
    overpack_count      = int(rec_df['Overpack_Candidate'].sum())
    overpack_pct2       = overpack_count / rec_wood * 100 if rec_wood else 0
    _, _, wood_fc_pct, _ = (fc_results if 'fc_results' in dir() else {}).get('Wood', (None, None, None, None)) if 'fc_results' in dir() else (None, None, None, None)
    wood_trend_str = (f"MENINGKAT {wood_fc_pct:.1f}%" if (wood_fc_pct or 0) > 0 else f"MENURUN {abs(wood_fc_pct or 0):.1f}%") if wood_fc_pct is not None else "Tidak tersedia"
    top_wood_cust = rec_df[rec_df['Is_Wood']].groupby('Name').size().sort_values(ascending=False).head(3).to_dict()

    rec_period_label = fmt_month_str(rec_ym_sel) if rec_view == 'Monthly' else (rec_year_sel if rec_year_sel != 'All' else '2023–2026')

    GEMINI_PROMPT = f"""
Kamu adalah pakar sustainability dan konsultan packaging industri manufaktur untuk PT Güntner Indonesia.

## DATA AKTUAL PACKAGING PT GÜNTNER ({rec_period_label})
- Total shipment lines: {rec_total:,}
- Wood: {rec_wood:,} lines ({rec_woodpct:.1f}%)
- Carton Box: {rec_cb:,} lines ({rec_cbpct:.1f}%)
- Cardboard: {rec_cbd:,} lines | Carton: {rec_ct:,} lines
- Rata-rata dimensi item Wood: {avg_wood_dim_rec:.0f} mm
- Rata-rata dimensi item Non-Wood: {avg_nonwood_dim_rec:.0f} mm
- Item Wood kecil (<1.5m, <50kg) yang bisa diganti: {overpack_count:,} lines ({overpack_pct2:.1f}% dari Wood)
- Tren Wood 12 bulan ke depan: {wood_trend_str}

## TREN TAHUNAN WOOD: {yearly_wood_rec}
## TOP 5 ITEM WOOD: {top_wood_items_rec}
## TOP 3 CUSTOMER PENGGUNA WOOD: {top_wood_cust}
## TARGET: {sustainability_goal}
## KONTEKS: {context_extra if context_extra.strip() else 'Tidak ada.'}

## PLEDGE 6 PT GÜNTNER:
"Make all our packaging fully sustainable. Most of our packaging currently comes from renewable sources,
so as well as looking at more eco-friendly materials such as recycled plastic foil, we are also looking
at ways in which all of our packaging can be reused, repurposed and/or recycled."

---
Berikan **3 rekomendasi strategis dan konkret** berdasarkan data di atas.

Format PERSIS untuk setiap rekomendasi:

---
### REKOMENDASI 1: [JUDUL] — Tingkat: MUDAH (Quick Win · 0–3 bulan)

**Latar Belakang Berbasis Data:**
[2–3 kalimat menggunakan angka spesifik]

**Deskripsi Rekomendasi:**
[Penjelasan detail]

**Langkah Implementasi:**
1. [Langkah 1]
2. [Langkah 2]
3. [Langkah 3]
4. [Langkah 4]

**Dampak pada Pledge 6:**
[Estimasi % pengurangan]

**Keuntungan Bisnis:**
[Estimasi penghematan]

**KPI Keberhasilan:**
- KPI 1: [Metrik & target]
- KPI 2: [Metrik & target]

---
### REKOMENDASI 2: [JUDUL] — Tingkat: SEDANG (Strategic · 3–12 bulan)
[Struktur sama]

---
### REKOMENDASI 3: [JUDUL] — Tingkat: SULIT (Transformational · 12+ bulan)
[Struktur sama]

---
### PENUTUP
[Paragraf penutup inspiratif]

---
Gunakan bahasa Indonesia profesional. WAJIB gunakan angka spesifik dari data di atas.
"""

    col_btn1, col_btn2, _ = st.columns([1.5, 1, 5])
    with col_btn1:
        gen_button = st.button("Generate Rekomendasi AI", type="primary", use_container_width=True)
    with col_btn2:
        if st.button("Reset", use_container_width=True):
            if 'rec_result' in st.session_state:
                del st.session_state['rec_result']

    # Data snapshot
    with st.expander("Data Snapshot yang Akan Dianalisis AI", expanded=True):
        s1, s2, s3, s4 = st.columns(4)
        s1.metric("Total Lines", f"{rec_total:,}")
        s2.metric("Wood Lines", f"{rec_wood:,}", f"{rec_woodpct:.1f}%")
        s3.metric("Overpackaged Wood", f"{overpack_count:,}", f"{overpack_pct2:.1f}% dari Wood")
        s4.metric("Wood Trend 12m", wood_trend_str)

    if gen_button:
        try:
            import google.generativeai as genai
            GEMINI_API_KEY = st.secrets["GEMINI_API_KEY"]
            genai.configure(api_key=GEMINI_API_KEY)
            gemini_model = genai.GenerativeModel('gemini-2.5-flash')
            with st.spinner("Gemini sedang menganalisis data dan menyusun rekomendasi…"):
                response = gemini_model.generate_content(GEMINI_PROMPT)
                st.session_state['rec_result'] = response.text
        except ImportError:
            st.error("google-generativeai belum terinstall. Jalankan: pip install google-generativeai")
        except KeyError:
            st.error("GEMINI_API_KEY tidak ditemukan di secrets.")
        except Exception as e:
            st.error(f"Error: {e}")

    if 'rec_result' in st.session_state:
        rec_text = st.session_state['rec_result']
        st.markdown("---")
        st.markdown("## Hasil Rekomendasi AI")

        sections = rec_text.split("---")
        for section in sections:
            section = section.strip()
            if not section: continue
            if 'MUDAH' in section or 'REKOMENDASI 1' in section:
                bg_col, border_col = '#f0fdf4', '#16a34a'
            elif 'SEDANG' in section or 'REKOMENDASI 2' in section:
                bg_col, border_col = '#fffbeb', '#d97706'
            elif 'SULIT' in section or 'REKOMENDASI 3' in section:
                bg_col, border_col = '#fef2f2', '#dc2626'
            elif 'PENUTUP' in section:
                bg_col, border_col = '#f0f9ff', '#0284c7'
            else:
                bg_col, border_col = '#f8fafc', '#64748b'

            with st.container():
                st.markdown(
                    f'<div style="background:{bg_col};border-left:5px solid {border_col};'
                    f'border-radius:12px;padding:2px 20px;margin-bottom:16px;">',
                    unsafe_allow_html=True)
                st.markdown(section)
                st.markdown('</div>', unsafe_allow_html=True)

        # Download DOCX
        st.markdown("---")
        dl1, dl2 = st.columns(2)

        def make_docx(content: str, title: str = "AI Recommendation Report") -> bytes:
            try:
                from docx import Document
                from docx.shared import Pt
                from docx.enum.text import WD_ALIGN_PARAGRAPH
            except ImportError:
                import subprocess, sys
                subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx', '-q'])
                from docx import Document
                from docx.shared import Pt
                from docx.enum.text import WD_ALIGN_PARAGRAPH
            doc = Document()
            # Title
            h = doc.add_heading(title, 0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Meta
            doc.add_paragraph(f"PT Güntner Indonesia | Supporting Pledge 6")
            doc.add_paragraph(f"Periode: {rec_year_sel} | Material: {focus_mat_rec} | Target: {sustainability_goal}")
            doc.add_paragraph("")
            # Content — parse markdown headings crudely
            for line in content.split('\n'):
                line = line.strip()
                if not line: doc.add_paragraph("")
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=2)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=1)
                elif line.startswith('**') and line.endswith('**'):
                    p = doc.add_paragraph()
                    run = p.add_run(line.strip('*'))
                    run.bold = True
                elif line.startswith('- '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.startswith(('1.','2.','3.','4.')):
                    doc.add_paragraph(line, style='List Number')
                else:
                    doc.add_paragraph(line)
            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()

        with dl1:
            docx_bytes = make_docx(rec_text, "AI Recommendation — PT Güntner")
            st.download_button("Download Rekomendasi (DOCX)",
                docx_bytes,
                'packaging_recommendations_guntner.docx',
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                use_container_width=True)
        with dl2:
            full_text = (
                f"PACKAGING RECOMMENDATION REPORT — PT Güntner Indonesia\n{'='*60}\n"
                f"Periode: {rec_year_sel} | Material: {focus_mat_rec} | Target: {sustainability_goal}\n"
                f"{'='*60}\n\nDATA SNAPSHOT\n"
                f"Total Lines: {rec_total:,}\nWood: {rec_wood:,} ({rec_woodpct:.1f}%)\n"
                f"Carton Box: {rec_cb:,} ({rec_cbpct:.1f}%)\n"
                f"Overpackaged: {overpack_count:,} ({overpack_pct2:.1f}% dari Wood)\n"
                f"Wood Trend: {wood_trend_str}\n\n{'='*60}\n\n{rec_text}"
            )
            full_docx = make_docx(full_text, "Full Packaging Report — PT Güntner")
            st.download_button("Download Full Report (DOCX)",
                full_docx,
                'full_packaging_report_guntner.docx',
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                use_container_width=True)

# FOOTER
st.markdown("---")
st.markdown("""
<div style='text-align:center;color:#64748b;padding:1.5rem 1rem;background:white;border-radius:.75rem;'>
  <p style='font-weight:700;font-size:1rem;margin-bottom:.3rem;color:#166534;'>Packaging Analysis Dashboard</p>
  <p style='font-size:.83rem;'>PT Güntner Indonesia &nbsp;|&nbsp; Supporting Pledge 6: Make all our packaging fully sustainable</p>
  <p style='font-size:.73rem;color:#9ca3af;margin-top:.2rem;'>Powered by XGBoost Forecasting · Gemini AI · Streamlit</p>
</div>
""", unsafe_allow_html=True)